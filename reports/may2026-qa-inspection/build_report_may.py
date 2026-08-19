#!/usr/bin/env python3
"""
May 2026 QA monthly report — V0 SNAPSHOT (截至 2026-05-04).

This is a partial-month snapshot (4 days of data) for QA leadership.
The final May monthly report will be regenerated 6/1+ on full-month data.

Structural rules from April v2 baseline apply, with snapshot adjustments:
  - Cover subtitle: Mid-Month Snapshot · 2026-05-01 to 2026-05-04
  - 文档信息 状态: V0 snapshot稿
  - ⚠ 数据说明 leads with snapshot caveat (covers X 次巡检; final 6 月初)
  - §1.1 / §1.3 partial-data flags: 覆盖率 7/13 = 53.8%
  - §2.1 thresholds vs currently-inspected (N=7) with footnote
  - §4.2 cross-month: April corrected baseline (含申诉调整)
  - §4.4 SKIP: each May store has only 1 inspection
  - §7.3 SKIP: no May store has ≥2 self-checks
  - §7.5 trend: 5月 row 4-day partial + status 🟡
  - §7.7 front-loaded with remaining-27-days targets

10-module strict taxonomy (Other → flag and STOP). Footer: ── 报告结束 ──, no auto meta-line.
"""
import csv, datetime as dt
from collections import defaultdict, Counter
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ============================================================================
# RUN PARAMETERS
# ============================================================================
TARGET_MONTH    = "05"
TARGET_MONTH_CN = "2026年05月"
TARGET_YM       = "2026-05"
DOC_ID          = "LCNA-QA-2026-005-snapshot"
TODAY           = "2026-05-05"
SNAPSHOT_FROM   = "2026-05-01"
SNAPSHOT_TO     = "2026-05-04"
SNAPSHOT_DAYS   = 4
DAYS_REMAINING  = 27
APR_COHORT_N    = 13   # April baseline cohort denominator for /13 narrative

DATA_DIR        = Path("/app/reports/may2026-qa-inspection")
APR_DIR         = Path("/app/reports/april2026-qa-inspection")
OUT_DIR         = DATA_DIR / "output"
OUT_DIR.mkdir(exist_ok=True)
OUT_PATH        = OUT_DIR / f"QA门店巡检月度分析报告_{TARGET_MONTH_CN}_截至4日_snapshot.docx"

# Color palette (matches April baseline)
NAVY        = "1F4E79"; BLUE        = "2E75B6"; TEXT        = "2C3E50"
BORDER      = "BFBFBF"; ALT_FILL    = "F2F2F2"
CRITICAL_FG = "C0392B"; WARN_FG     = "E67E22"; HEALTHY_FG  = "27AE60"; INFO_FG = "0365C0"
CALLOUT_YELLOW = "FFF2CC"; CALLOUT_GREEN = "E2EFDA"
CALLOUT_BLUE   = "D6E4F0"; CALLOUT_RED   = "FCE4EC"

# ============================================================================
# 10-MODULE STRICT TAXONOMY
# ============================================================================
MODULES_10 = ["清洁卫生","过程控制","设施","证照文件","职场安全",
              "虫害防控","温控有效期","员工健康卫生","设备维护","供应商"]
MODULE_SHORT = {"清洁卫生":"清洁","过程控制":"过程","设施":"设施","证照文件":"证照",
                "职场安全":"职安","虫害防控":"虫害","温控有效期":"温控",
                "员工健康卫生":"员工","设备维护":"设备","供应商":"供应","其他":"其他"}

EN2CN = {
    "Cleaning and Sanitation": "清洁卫生",
    "Process Control": "过程控制",
    "Facility": "设施",
    "Document Record": "证照文件",
    "Workplace Safety": "职场安全",
    "Pests Control": "虫害防控",
    "Temperature Control / Expiration Date Management.": "温控有效期",
    "Employees’ Health and Personal Hygiene": "员工健康卫生",   # U+2019 right single quote
    "Employees' Health and Personal Hygiene": "员工健康卫生",   # ASCII apostrophe variant
    "Maintenance of Equipment": "设备维护",
    "Approved Supplier": "供应商",
}
def cn_module(it):
    return EN2CN.get(it["module_name"], "其他")

# Treat "(无描述)" as empty
EMPTY_DESC_MARKERS = {"(无描述)", "(无描述）", "（无描述）", "（无描述)", ""}
def real_desc(s):
    s = (s or "").strip()
    if s in EMPTY_DESC_MARKERS: return ""
    return s

# ============================================================================
# DATA LOADING
# ============================================================================
def load_csv(p):
    with open(p, encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))

print(f"[load] target={TARGET_MONTH_CN} (snapshot 截至 {SNAPSHOT_TO})  doc_id={DOC_ID}")
SUMMARY = load_csv(DATA_DIR / "may2026_inspection_summary.csv")
ITEMS   = load_csv(DATA_DIR / "may2026_inspection_items.csv")
STORE_MASTER  = load_csv(DATA_DIR / "may2026_store_master.csv")
INSP_TREND    = load_csv(DATA_DIR / "may2026_inspector_trend.csv")
TREND_SUMMARY = load_csv(DATA_DIR / "jan_to_may2026_trend_summary.csv")
APR_SUM = load_csv(DATA_DIR / "april2026_inspection_summary.csv")
APR_ITM = load_csv(DATA_DIR / "april2026_inspection_items.csv")

# Type normalization for May summary
def norm_summary(rows):
    for r in rows:
        for k in ("adjusted_total_score","original_total_score","adjusted_total_deduction",
                  "original_total_deduction","is_appealed","item_count",
                  "s_count","m_count","g_count","l_count","inspection_id"):
            try: r[k] = int(r[k]) if r[k] not in ("","None",None) else 0
            except: r[k] = 0
        # appeal_status derivation (no headers JSON for May; use score delta)
        if r["is_appealed"] == 0:
            r["appeal_status"] = "none"
        elif r["adjusted_total_score"] != r["original_total_score"]:
            r["appeal_status"] = "approved"
        else:
            r["appeal_status"] = "denied"   # default; pending requires headers JSON
norm_summary(SUMMARY)
norm_summary(APR_SUM)

for r in ITEMS:
    r["deduction_points"] = int(r["deduction_points"])
    r["is_appealed_finding"] = int(r["is_appealed_finding"])
    r["inspection_id"] = int(r["inspection_id"])
    r["cn_module"] = cn_module(r)
    r["clean_desc"] = real_desc(r.get("issue_description",""))

print(f"[load] summary={len(SUMMARY)} items={len(ITEMS)} stores={len(STORE_MASTER)}")
print(f"[load] apr_summary={len(APR_SUM)} apr_items={len(APR_ITM)}")

# Module mapping anomalies — STOP if found
OTHER_FLAGS = Counter(it["module_name"] for it in ITEMS if it["cn_module"] == "其他")
if OTHER_FLAGS:
    raise SystemExit(f"[FAIL] Module mapping has 'Other' bucket items: {dict(OTHER_FLAGS)}")
print(f"[mapping] ✅ all module names map to canonical 10")

# ============================================================================
# DERIVED
# ============================================================================
PRIORITY = {"QA审计": 0, "区经检查": 1, "门店自检": 2}

by_store = defaultdict(list)
for r in SUMMARY: by_store[r["store_code"]].append(r)
def main_inspection_for_store(rows):
    if not rows: return None
    return sorted(rows, key=lambda r:(PRIORITY.get(r["inspection_type"], 99),
                                      -dt.date.fromisoformat(r["inspection_date"]).toordinal()))[0]
MAIN_INSP = {sc: main_inspection_for_store(rows) for sc, rows in by_store.items()}
MAIN_IIDS = {m["inspection_id"] for m in MAIN_INSP.values() if m}
ACTIVE_STORE_CODES = sorted(by_store.keys())
N_STORES_INSPECTED = len(ACTIVE_STORE_CODES)   # = 7

items_by_iid = defaultdict(list)
for it in ITEMS: items_by_iid[it["inspection_id"]].append(it)

MAIN_ITEMS = [it for it in ITEMS if it["inspection_id"] in MAIN_IIDS]
print(f"[derive] main_inspections={len(MAIN_INSP)} main_items={len(MAIN_ITEMS)} all_items={len(ITEMS)}")

def aggregate(items):
    sev = Counter(); ded = Counter()
    mod_sev = defaultdict(Counter); mod_ded = defaultdict(Counter)
    mod_stores = defaultdict(set); store_mod_ded = defaultdict(lambda: defaultdict(int))
    for it in items:
        sev[it["severity"]] += 1
        ded[it["severity"]] += it["deduction_points"]
        mod_sev[it["cn_module"]][it["severity"]] += 1
        mod_ded[it["cn_module"]][it["severity"]] += it["deduction_points"]
        mod_stores[it["cn_module"]].add(it["store_code"])
        store_mod_ded[it["store_code"]][it["cn_module"]] += it["deduction_points"]
    return sev, ded, mod_sev, mod_ded, mod_stores, store_mod_ded

MAIN_SEV, MAIN_DED, MAIN_MOD_SEV, MAIN_MOD_DED, MAIN_MOD_STORES, MAIN_STORE_MOD = aggregate(MAIN_ITEMS)
ALL_SEV, ALL_DED, ALL_MOD_SEV, ALL_MOD_DED, ALL_MOD_STORES, _ = aggregate(ITEMS)

INSP_BY_TYPE = defaultdict(list)
for r in SUMMARY: INSP_BY_TYPE[r["inspection_type"]].append(r)
N_BY_TYPE = {t: len(INSP_BY_TYPE[t]) for t in ("门店自检","QA审计","区经检查")}
N_TOTAL = sum(N_BY_TYPE.values())

valid_main_scores = [m["adjusted_total_score"] for m in MAIN_INSP.values()]
avg_score_main = round(sum(valid_main_scores)/len(valid_main_scores), 1)

APPEALED         = [r for r in SUMMARY if r["is_appealed"] == 1]
APPEAL_APPROVED  = [r for r in APPEALED if r["appeal_status"] == "approved"]
APPEAL_DENIED    = [r for r in APPEALED if r["appeal_status"] == "denied"]
APPEAL_PENDING   = [r for r in APPEALED if r["appeal_status"] == "pending"]
print(f"[appeal] total={len(APPEALED)} approved={len(APPEAL_APPROVED)} denied={len(APPEAL_DENIED)} pending={len(APPEAL_PENDING)}")

# April baseline 主巡检 per store (for §1.3 / §4.2 cross-month)
apr_by_store = defaultdict(list)
for r in APR_SUM: apr_by_store[r["store_code"]].append(r)
APR_MAIN = {sc: main_inspection_for_store(rs) for sc, rs in apr_by_store.items()}
def apr_main_score_for(sc):
    m = APR_MAIN.get(sc)
    return m["adjusted_total_score"] if m else None

# Recently-opened stores (for callout)
RECENTLY_OPENED_NO_MAY = []
for s in STORE_MASTER:
    if s["status"] != "active" or s["store_code"].startswith("US999") or s["store_code"] == "US00000":
        continue
    od = s["open_date"]
    if od and "2026-04" <= od <= SNAPSHOT_TO and s["inspected_in_may"] == "No":
        RECENTLY_OPENED_NO_MAY.append(s)

# ============================================================================
# DOCX HELPERS  (matches April template exactly)
# ============================================================================
def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)

def set_cell_borders(cell, color_hex=BORDER, sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    bdrs = OxmlElement("w:tcBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"),"single"); b.set(qn("w:sz"),str(sz))
        b.set(qn("w:space"),"0"); b.set(qn("w:color"), color_hex)
        bdrs.append(b)
    tcPr.append(bdrs)

def set_run_font(run, name="Arial", size=10, bold=False, color_hex=TEXT, italic=False):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfont = rpr.find(qn("w:rFonts"))
    if rfont is None:
        rfont = OxmlElement("w:rFonts"); rpr.append(rfont)
    rfont.set(qn("w:ascii"), name); rfont.set(qn("w:hAnsi"), name)
    rfont.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    if color_hex: run.font.color.rgb = RGBColor.from_string(color_hex)

def add_para(doc, text, size=10, bold=False, color_hex=TEXT, align=None,
             space_before=0, space_after=4, indent_left=0, italic=False):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent_left: p.paragraph_format.left_indent = Pt(indent_left)
    if text:
        r = p.add_run(text); set_run_font(r, size=size, bold=bold, color_hex=color_hex, italic=italic)
    return p

def add_heading(doc, text, level=1):
    sizes = {1:16, 2:13, 3:11}; colors = {1:NAVY, 2:BLUE, 3:NAVY}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level==1 else 8)
    p.paragraph_format.space_after  = Pt(6 if level==1 else 4)
    r = p.add_run(text); set_run_font(r, size=sizes[level], bold=True, color_hex=colors[level])

def add_callout(doc, lines, kind="warn", title=None):
    bg, bar = {
        "warn":     (CALLOUT_YELLOW, WARN_FG),
        "good":     (CALLOUT_GREEN,  HEALTHY_FG),
        "info":     (CALLOUT_BLUE,   INFO_FG),
        "critical": (CALLOUT_RED,    CRITICAL_FG),
    }[kind]
    table = doc.add_table(rows=1, cols=1); table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, bg); set_cell_borders(cell, color_hex=bar, sz=8)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    if isinstance(lines, str): lines = [lines]
    if title:
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title); set_run_font(r, size=10, bold=True, color_hex=bar)
        for ln in lines:
            pp = cell.add_paragraph(); pp.paragraph_format.space_after = Pt(2)
            rr = pp.add_run(ln); set_run_font(rr, size=10, color_hex=TEXT)
    else:
        for i, ln in enumerate(lines):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(ln); set_run_font(r, size=10, color_hex=TEXT)
    add_para(doc, "", size=4, space_after=4)

def add_data_table(doc, headers, rows, col_widths=None, align_center_cols=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = "Table Grid"
    if col_widths:
        for ci, w in enumerate(col_widths):
            for c in t.columns[ci].cells: c.width = Cm(w)
    for i, h in enumerate(headers):
        c = t.cell(0, i); set_cell_shading(c, NAVY)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.text = ""; p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        r = p.add_run(h); set_run_font(r, size=10, bold=True, color_hex="FFFFFF")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.cell(ri+1, ci)
            if ri % 2 == 1: set_cell_shading(c, ALT_FILL)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            c.text = ""; p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
            txt, opts = (val if isinstance(val, tuple) else (val, {}))
            if align_center_cols and ci in align_center_cols:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif opts.get("align") == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(txt))
            set_run_font(r, size=opts.get("size", 9),
                         bold=opts.get("bold", False),
                         color_hex=opts.get("color", TEXT))

def add_page_field(run, instr):
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"),"begin")
    instr_text = OxmlElement("w:instrText"); instr_text.text = instr
    fld_end   = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"),"end")
    run._r.append(fld_begin); run._r.append(instr_text); run._r.append(fld_end)

def setup_page_footer(section, doc_id, today):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""
    r1 = p.add_run(f"{doc_id} | 编制：曾翔宇 | 日期：{today} | 第 ")
    set_run_font(r1, size=9, color_hex="666666")
    r2 = p.add_run(""); set_run_font(r2, size=9, color_hex="666666")
    add_page_field(r2, "PAGE")
    r3 = p.add_run(" 页 / 共 "); set_run_font(r3, size=9, color_hex="666666")
    r4 = p.add_run(""); set_run_font(r4, size=9, color_hex="666666")
    add_page_field(r4, "NUMPAGES")
    r5 = p.add_run(" 页"); set_run_font(r5, size=9, color_hex="666666")

# ============================================================================
# DOCUMENT
# ============================================================================
doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.59); section.page_height = Cm(27.94)
section.top_margin = Cm(1.27); section.bottom_margin = Cm(1.27)
section.left_margin = Cm(1.59); section.right_margin = Cm(1.59)
setup_page_footer(section, DOC_ID, TODAY)

style = doc.styles["Normal"]
style.font.name = "Arial"; style.font.size = Pt(10)

# ---------- COVER ----------
add_para(doc, "瑞幸咖啡北美", size=22, bold=True, color_hex=NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=8)
add_para(doc, "QA门店巡检月度分析报告", size=18, bold=True, color_hex=NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
# Snapshot adjustment #1 — replace English subtitle
add_para(doc, f"Mid-Month Snapshot · {SNAPSHOT_FROM} to {SNAPSHOT_TO}", size=12, color_hex=BLUE, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
add_para(doc, TARGET_MONTH_CN, size=20, bold=True, color_hex=NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=80)
add_para(doc, "质量保障部 / 基础设施部", size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para(doc, "编制：曾翔宇    日期：" + TODAY, size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para(doc, f"报告编号：{DOC_ID} | 状态：V0 snapshot稿", size=10, color_hex=BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
doc.add_page_break()

# ---------- 文档信息 ----------
add_heading(doc, "文档信息", 1)
total_findings_all  = len(ITEMS)
total_findings_main = len(MAIN_ITEMS)
all_s, all_m, all_g, all_l = (ALL_SEV.get(k,0) for k in ("S","M","G","L"))
main_s, main_m, main_g, main_l = (MAIN_SEV.get(k,0) for k in ("S","M","G","L"))
appeal_text = "本月暂无申诉立案"

doc_info_rows = [
    ["报告编号", DOC_ID],
    ["报告周期", TARGET_MONTH_CN],
    ["数据范围", f"{SNAPSHOT_FROM} 至 {SNAPSHOT_TO}（共 {SNAPSHOT_DAYS} 天）"],
    ["有效门店", f"{N_STORES_INSPECTED} 家已巡检 / 4 月可比基线 {APR_COHORT_N} 家"],
    ["巡检类型", f"门店自检（{N_BY_TYPE['门店自检']}次） + QA审计（{N_BY_TYPE['QA审计']}次） + 区经检查（{N_BY_TYPE['区经检查']}次） = 共 {N_TOTAL} 次（已提交）"],
    ["问题总数", f"全月 {total_findings_all} 个扣分项（S 项 {all_s}、M 项 {all_m}、G 项 {all_g}、L 项 {all_l}）；主巡检 {total_findings_main} 个（S 项 {main_s}、M 项 {main_m}、G 项 {main_g}、L 项 {main_l}）"],
    ["编制人", "曾翔宇"],
    ["部门", "质量保障部 / 基础设施部"],
    ["数据来源", "empapp 门店稽核系统（aws-luckyus-opqualitycontrol-rw / luckyus_opqualitycontrol）"],
    ["状态", f"V0 snapshot稿（截至 {SNAPSHOT_TO} 数据，5 月底完成正式月度版）"],
]
add_data_table(doc, ["项目","内容"], doc_info_rows, col_widths=[3.5, 14.5])
add_para(doc, "", space_after=6)

# ---------- 数据说明 ----------
add_heading(doc, "⚠ 数据说明（V0 snapshot稿）", 1)
data_notes = [
    f"(a) 本报告为 {SNAPSHOT_FROM} 至 {SNAPSHOT_TO} 截至 snapshot，覆盖 {N_TOTAL} 次巡检；正式月度报告将于 6 月初基于完整 5 月数据生成。本报告不应作为 5 月月度结论使用。",
    f"(b) 本月暂无申诉立案；4 月已闭环的 6 起申诉（2 获批 / 3 驳回 / 1 审批中）后续整改情况将在 5 月月度版跟踪。",
    f"(c) 巡检覆盖：5 月已完成 {N_TOTAL} 次有效巡检，{N_STORES_INSPECTED} 家门店覆盖（4 月可比基线 {APR_COHORT_N} 家，覆盖率 {N_STORES_INSPECTED}/{APR_COHORT_N} = {N_STORES_INSPECTED*100/APR_COHORT_N:.1f}%），剩余 {DAYS_REMAINING} 天 5 月节奏待持续观察。",
    f"(d) QA 人员状态：Eamonn Caballar 5 月前 4 日尚未提交 QA 审计（4 月共 12 次），属典型「月初轻、月末重」节奏，需结合月底数据评估；Yu Jiang HR 离职 / 转岗流程仍未关闭（4 月、5 月均无任何巡检）。",
    f"(e) 区经巡检：Daniel Chu 已完成 3 次（覆盖 US00003、US00006、US00025），节奏稳定；Jung Han Liang 5 月前 4 日尚无提交。",
    f"(f) 新开门店待覆盖：US00007（108th & Broadway，4/30 开业）、US00010（154 Bleecker，4/28 开业）、US00019（29th & 3rd，4/11 开业）截至 5/4 仍未完成 5 月首次巡检；US00015（41st & Lexington，4/30 开业）已于 5/1 完成首次自检。",
    "(g) 标准规则：① 主巡检优先级 QA审计 > 区经检查 > 门店自检；同优先级取最近日期（每店 5 月仅 1 次巡检，主巡检与全月口径一致）；② 仅采用已提交单据，草稿已剔除；③ 同人同店同日重复 100 分零扣分自检按误提交规则剔除（5 月前 4 日未发现此类）。",
    "(h) v3 prompt 规则在效：10 模块严格分类、§2.2/§2.3/§4.1/§4.5 主巡检口径锁定、申诉双轨制、未生成自动 footer 元信息行、中文正文（finding 原文除外）。",
]
add_callout(doc, data_notes, kind="warn")

# ---------- 管理摘要 ----------
add_heading(doc, "管理摘要", 1)
add_para(doc,
    f"本月平均分 {avg_score_main} 分（基于各门店主巡检），覆盖 {N_STORES_INSPECTED} 家门店（4 月可比基线 {APR_COHORT_N} 家，覆盖率 {N_STORES_INSPECTED*100/APR_COHORT_N:.1f}%），主巡检 {total_findings_main} 项有效扣分项，全月共 {total_findings_all} 项（5 月每店仅 1 次巡检，主巡检 = 全月）。",
    size=10, space_after=8)

lowest = min(MAIN_INSP.values(), key=lambda r: r["adjusted_total_score"])
highest = max(MAIN_INSP.values(), key=lambda r: r["adjusted_total_score"])

add_callout(doc,
    f"5 月前 4 日体系节奏：4 月已完成「危机—响应」全面恢复（58 次三类齐全），5 月剩余 {DAYS_REMAINING} 天需维持月度全覆盖。当前 {N_TOTAL} 次提交（自检 {N_BY_TYPE['门店自检']} / QA {N_BY_TYPE['QA审计']} / 区经 {N_BY_TYPE['区经检查']}）属月初典型轻量节奏，QA 审计尚未启动需关注。",
    kind="info", title="🟡 月初节奏 (snapshot)")

# Compute key sub-item systemic clusters from May data
all_s_items = [it for it in ITEMS if it["severity"] == "S"]
sub_cluster = Counter(it["module_subcategory"] for it in all_s_items)
top_sub = sub_cluster.most_common(2)
sub_text = "、".join(f"{n}（{c} 起）" for n, c in top_sub) if top_sub else "—"
n_with_s_stores = len(set(it["store_code"] for it in all_s_items))

if all_s_items:
    add_callout(doc,
        f"S 项已现：5 月前 4 日发现 {all_s} 个 S 项（关键项），分布在 {n_with_s_stores} 家门店，子项 {sub_text}。值得注意：4 月 air gap / Sinks and Pipes 系统性 S 项（10 起）尚未根治，5/3 区经检查在 102 Fulton 再次命中同一类问题，需 BD 整改清单优先跟踪。",
        kind="warn", title="⚠ S 项预警")
else:
    add_callout(doc,
        f"S 项暂无：5 月前 4 日各门店主巡检均未触发 S 项，但样本仅 {N_STORES_INSPECTED} 家、覆盖 {SNAPSHOT_DAYS} 天，远不足以判定系统性改善。建议月底再行评估。",
        kind="good", title="✅ S 项 (snapshot)")

# Cross-month decline check
decliners = []
for sc, m in MAIN_INSP.items():
    apr = apr_main_score_for(sc)
    if apr is not None:
        d = m["adjusted_total_score"] - apr
        if d <= -10:
            decliners.append((sc, m, apr, d))
if decliners:
    txt = "、".join(f"{m['store_name']} 4 月 {apr}→5 月 {m['adjusted_total_score']}（{d:+d}）" for sc,m,apr,d in decliners)
    add_callout(doc,
        f"跨月对比预警：{len(decliners)} 家门店主巡检较 4 月（含申诉调整）下滑 ≥10 分：{txt}。注意 5 月主巡检多为门店自检 / 区经检查，与 4 月以 QA 审计为主的口径不完全可比，需结合后续 QA 审计验证。",
        kind="warn", title="⚠ 跨月对比 (snapshot)")

# ============================================================================
# §一、门店整体表现
# ============================================================================
add_heading(doc, "一、门店整体表现", 1)

add_heading(doc, "1.1 本月概览", 2)
n_below_80 = sum(1 for m in MAIN_INSP.values() if m["adjusted_total_score"] < 80)
n_with_s   = sum(1 for m in MAIN_INSP.values() if m["s_count"] > 0)
overview_rows = [[
    f"{highest['store_name']}\n{highest['adjusted_total_score']} 分",
    f"{lowest['store_name']}\n{lowest['adjusted_total_score']} 分",
    f"{n_with_s} 家",
    f"{n_below_80} 家",
]]
add_data_table(doc, ["最高分门店","最低分门店","S 项门店数","<80 分门店数"],
               overview_rows, col_widths=[4.5,4.5,3.0,3.0], align_center_cols=[0,1,2,3])
add_para(doc,
    f"覆盖率（vs 4 月可比基线）：{N_STORES_INSPECTED}/{APR_COHORT_N} = {N_STORES_INSPECTED*100/APR_COHORT_N:.1f}%（剩余 {DAYS_REMAINING} 天待巡检）。",
    size=9, color_hex="666666", space_before=2, space_after=4)

add_heading(doc, "1.2 各门店得分明细（基于主巡检）", 2)
ranked = sorted(MAIN_INSP.values(),
                key=lambda r: (-r["adjusted_total_score"], PRIORITY.get(r["inspection_type"], 99), r["store_code"]))
rows_12 = []
for i, m in enumerate(ranked, 1):
    sc = m["adjusted_total_score"]
    score_color = CRITICAL_FG if sc < 80 else (NAVY if sc >= 85 else TEXT)
    score_bold  = sc < 80 or sc >= 85
    s_color = CRITICAL_FG if m["s_count"] > 0 else TEXT
    rows_12.append([
        str(i),
        (m["store_name"], {"bold": True}),
        m["store_code"],
        (str(sc), {"bold": score_bold, "color": score_color, "align":"center"}),
        m["inspection_type"],
        (str(m["adjusted_total_deduction"]), {"align":"center"}),
        (str(m["s_count"]), {"bold": m["s_count"]>0, "color": s_color, "align":"center"}),
        (str(m["m_count"]), {"align":"center"}),
        (str(m["g_count"]), {"align":"center"}),
        (str(m["l_count"]), {"align":"center"}),
        m["inspector_name"],
    ])
add_data_table(doc, ["#","门店","编号","得分","巡检类型","扣分","S","M","G","L","巡检员"],
               rows_12, col_widths=[0.6,3.0,1.5,1.0,1.7,1.0,0.6,0.6,0.6,0.6,3.0],
               align_center_cols=[0,2,3,4,5,6,7,8,9])
add_para(doc, f"5 月每店仅 1 次巡检（snapshot 4 日窗口），主巡检即为该次提交；本月暂无 ※ 申诉获批门店。",
         size=9, color_hex="666666", space_before=2, space_after=8)

add_heading(doc, "1.3 管理解读", 2)
n_above_85 = sum(1 for m in MAIN_INSP.values() if m["adjusted_total_score"] >= 85)
high_text = "、".join(f"{m['store_name']} {m['adjusted_total_score']} 分" for m in ranked if m["adjusted_total_score"] >= 85)
low_text  = "、".join(f"{m['store_name']} {m['adjusted_total_score']} 分" for m in ranked if m["adjusted_total_score"] < 80)

# Cross-month deltas (May vs April main)
delta_lines = []
for sc, m in MAIN_INSP.items():
    apr = apr_main_score_for(sc)
    if apr is not None:
        d = m["adjusted_total_score"] - apr
        delta_lines.append((d, sc, m["store_name"], apr, m["adjusted_total_score"]))
delta_lines.sort(key=lambda x: x[0])

bullets = [
    f"• 覆盖率：{N_STORES_INSPECTED}/{APR_COHORT_N} = {N_STORES_INSPECTED*100/APR_COHORT_N:.1f}%（剩余 {DAYS_REMAINING} 天待巡检），含 1 家新开店首检（41st & Lexington 4/30 开业，5/1 完成首次自检 90 分）。",
    f"• {n_above_85} 家门店达到 85 分以上（{high_text or '—'}）；{n_below_80} 家低于 80 分（{low_text or '—'}）。",
]
if all_s_items:
    sub_examples = []
    for it in all_s_items:
        sub_examples.append(f"{it['store_name']}（{it['module_subcategory']}）")
    bullets.append(f"• 主巡检 S 项 {all_s} 个分布在 {n_with_s_stores} 家门店：{'、'.join(sub_examples)}；与 4 月 air gap / Sinks and Pipes 系统性短板高度同源，需 BD 整改清单跟踪。")
else:
    bullets.append(f"• 5 月前 4 日主巡检暂未触发 S 项，但样本仅 {N_STORES_INSPECTED} 家、覆盖 {SNAPSHOT_DAYS} 天，待月底完整数据评估。")
# Improvers / decliners
ups = [(d,sc,n,a,b) for d,sc,n,a,b in delta_lines if d >= 5]
downs = [(d,sc,n,a,b) for d,sc,n,a,b in delta_lines if d <= -5]
if ups:
    bullets.append("• 跨月改善（vs 4 月主巡检 含申诉调整）：" + "；".join(f"{n} {a}→{b}（{'+' if d>0 else ''}{d}）" for d,_,n,a,b in ups))
if downs:
    bullets.append("• 跨月下滑（vs 4 月主巡检 含申诉调整）：" + "；".join(f"{n} {a}→{b}（{d:+d}）" for d,_,n,a,b in downs) + "。注意 5 月主巡检多为自检 / 区经，与 4 月 QA 审计口径不完全可比。")
# New stores without baseline
new_stores = [sc for sc in MAIN_INSP if sc not in APR_MAIN]
if new_stores:
    nm = "、".join(f"{MAIN_INSP[sc]['store_name']}（首检 {MAIN_INSP[sc]['adjusted_total_score']} 分）" for sc in new_stores)
    bullets.append(f"• 4 月无可比基线的新开店：{nm}——本店为本月首次巡检覆盖。")
# Self-check active discoveries (none in May 4-day snapshot)
bullets.append(f"• 自检主动发现：{N_BY_TYPE['门店自检']} 次门店自检均披露问题（无 100/100 误提交），平均分 {sum(m['adjusted_total_score'] for m in INSP_BY_TYPE['门店自检'])/max(1,N_BY_TYPE['门店自检']):.1f} 分；Andrew Hu（28th & 6th 自检 82 分，1 M 5 G 3 L）等自检员主动暴露 ppm 异常等问题，可作为标杆。")
bullets.append(f"• 同店跨类型背离分析：5 月每店仅 1 次巡检，无 cross-type 对比基础；待月底巡检密度提升后于 §4.4 / §7.3 补充。")

for b in bullets:
    add_para(doc, b, size=10, indent_left=12, space_after=4)

# ============================================================================
# §二、模块风险分析
# ============================================================================
add_heading(doc, "二、模块风险分析", 1)
add_para(doc,
    f"本月主巡检共发现 {total_findings_main} 个扣分项，分布在 {len([m for m in MODULES_10 if MAIN_MOD_SEV.get(m)])} 个标准模块中（5 月每店仅 1 次巡检，主巡检 = 全月，§3.3 仍按全月口径列示以保持与 4 月报告结构对齐）。",
    space_after=6)

add_heading(doc, "2.1 风险分层（基于主巡检覆盖率）", 2)
def coverage_main(mod):
    return len(MAIN_MOD_STORES.get(mod, set())) / max(1, N_STORES_INSPECTED)
sys_mods = [m for m in MODULES_10 if coverage_main(m) >= 0.5 and MAIN_MOD_SEV.get(m)]
mid_mods = [m for m in MODULES_10 if 0.3 <= coverage_main(m) < 0.5 and MAIN_MOD_SEV.get(m)]
low_mods = [m for m in MODULES_10 if 0 < coverage_main(m) < 0.3 and MAIN_MOD_SEV.get(m)]
add_para(doc, f"🔴 系统性风险（影响≥50% 门店）：{', '.join(sys_mods) or '（无）'}", size=10, indent_left=12, space_after=4)
add_para(doc, f"🟡 中等覆盖面（影响 30-49%）：{', '.join(mid_mods) or '（无）'}", size=10, indent_left=12, space_after=4)
add_para(doc, f"🟢 低覆盖面（<30%）：{', '.join(low_mods) or '（无）'}", size=10, indent_left=12, space_after=8)
add_para(doc, f"※ 覆盖率分母为已巡检 {N_STORES_INSPECTED} 家门店；4 月可比基线 {APR_COHORT_N} 家口径下覆盖率约为当前的一半，需在月底完整数据时重新计算。",
         size=9, color_hex="666666", space_after=4)

add_heading(doc, "2.2 模块排名总览（按扣分排序，主巡检视角）", 2)
mods_with_data = [m for m in MODULES_10 if m in MAIN_MOD_SEV]
mods_ranked = sorted(mods_with_data, key=lambda m: sum(MAIN_MOD_DED.get(m, {}).values()))
module_rows = []
for rank, mod in enumerate(mods_ranked, 1):
    sev_c = MAIN_MOD_SEV[mod]
    ded_total = sum(MAIN_MOD_DED.get(mod, {}).values())
    cnt_total = sum(sev_c.values())
    n_stores  = len(MAIN_MOD_STORES.get(mod, set()))
    cov_pct   = round(n_stores / N_STORES_INSPECTED * 100)
    risks = []
    if cov_pct >= 50: risks.append("⚠ 系统性")
    if sev_c.get("S",0) > 0: risks.append("⚠ 含 S 项")
    if not risks and sev_c.get("M",0) > 0: risks.append("含 M 项")
    if not risks: risks.append("---")
    module_rows.append([str(rank), mod, str(cnt_total), str(ded_total),
                        f"{n_stores}/{N_STORES_INSPECTED}", f"{cov_pct}%",
                        str(sev_c.get("S",0)), str(sev_c.get("M",0)),
                        str(sev_c.get("G",0)), str(sev_c.get("L",0)),
                        " / ".join(risks)])
add_data_table(doc, ["#","模块","问题数","扣分","门店","覆盖率","S","M","G","L","风险"],
               module_rows, col_widths=[0.6,3.4,1.0,1.0,1.0,1.0,0.6,0.6,0.6,0.6,2.6],
               align_center_cols=[0,2,3,4,5,6,7,8,9])
add_para(doc, f"※ 覆盖率分母为已巡检 {N_STORES_INSPECTED} 家门店。",
         size=9, color_hex="666666", space_before=2, space_after=4)

add_heading(doc, "2.3 重点模块详细分析（TOP 5，主巡检视角）", 2)

def quotable_items(items, max_n=20):
    sev_rank = {"S":0,"M":1,"G":2,"L":3}
    out = []
    for it in sorted(items, key=lambda x: (sev_rank.get(x["severity"],9), x["store_code"], x["inspection_date"])):
        if not it["clean_desc"]: continue
        out.append(it)
        if len(out) >= max_n: break
    return out

top_mods = mods_ranked[:5]
for mod in top_mods:
    sev_c = MAIN_MOD_SEV[mod]
    ded_total = sum(MAIN_MOD_DED[mod].values())
    n_stores = len(MAIN_MOD_STORES[mod])
    add_heading(doc, f"{mod} — {sum(sev_c.values())} 个扣分项，{ded_total} 分，影响 {n_stores} 家门店", 3)
    add_para(doc, f"严重级别：S 项 {sev_c.get('S',0)} 个、M 项 {sev_c.get('M',0)} 个、G 项 {sev_c.get('G',0)} 个、L 项 {sev_c.get('L',0)} 个。",
             size=10, space_after=4)
    mod_items_main = [it for it in MAIN_ITEMS if it["cn_module"] == mod]
    quotes = quotable_items(mod_items_main, max_n=20)
    if quotes:
        add_para(doc, "具体问题（引用原始描述，按严重度排序，最多展示 20 条；空描述已跳过）：",
                 size=10, color_hex="666666", space_after=4)
        for it in quotes:
            marker = "⚠ " if it["severity"] == "S" else ""
            desc = it["clean_desc"].replace("\n", " ").replace("\r", " ")[:200]
            add_para(doc,
                f"• {marker}{it['store_name']} ({it['store_code']})｜{it['module_subcategory']}｜{it['severity']} 项 {it['deduction_points']} 分｜{desc}",
                size=9, indent_left=8, space_after=2)
        if len(mod_items_main) > len(quotes):
            n_skipped = len(mod_items_main) - len(quotes)
            add_para(doc, f"… 另有 {n_skipped} 条主巡检发现描述为空（多由门店自检员未填写明细造成），完整明细见原始 CSV。",
                     size=9, color_hex="666666", indent_left=8, space_after=4)
    else:
        add_para(doc, "本子项无可引用的描述（所有发现描述均为空，多由门店自检员未填写明细造成）。",
                 size=10, color_hex="666666", space_after=4)
    if mod == "设施":
        sp_s = [it for it in mod_items_main if it["severity"]=="S" and "Sinks and Pipes" in it["module_subcategory"]]
        if len(sp_s) >= 1:
            add_callout(doc,
                f"4 月 air gap / Sinks and Pipes 系统性 S 项（10 起跨 6 家门店）尚未根治，5 月前 4 日已再次命中（{len(sp_s)} 起）。需 BD 整改清单优先跟踪 5 月剩余 {DAYS_REMAINING} 天的实际整改进展（详见 §6.2 P0）。",
                kind="critical", title="⚠ 关键发现 (snapshot)")

# ============================================================================
# §三、风险等级分布
# ============================================================================
add_heading(doc, "三、风险等级分布", 1)

add_heading(doc, "3.1 整体分布（主巡检）", 2)
def top_modules_for_sev_main(sev, top=4):
    cnt = Counter(it["cn_module"] for it in MAIN_ITEMS if it["severity"] == sev)
    return "、".join(f"{m}（{c}）" for m,c in cnt.most_common(top)) or "—"

main_total = max(1, sum(MAIN_SEV.values()))
sev_rows = [
    ["S 项（关键项）", str(main_s), f"{main_s/main_total*100:.1f}%", "2 天内闭环", top_modules_for_sev_main("S")],
    ["M 项（重要项）", str(main_m), f"{main_m/main_total*100:.1f}%", "7 天内闭环", top_modules_for_sev_main("M")],
    ["G 项（一般项）", str(main_g), f"{main_g/main_total*100:.1f}%", "14 天内闭环", top_modules_for_sev_main("G")],
    ["L 项（轻微项）", str(main_l), f"{main_l/main_total*100:.1f}%", "14 天内闭环", top_modules_for_sev_main("L")],
    ["合计", str(sum(MAIN_SEV.values())), "100%", "---", "---"],
]
add_data_table(doc, ["风险等级","数量","占比","SLA 要求","主要分布模块"], sev_rows,
               col_widths=[2.2,1.0,1.0,1.6,9.0])

add_callout(doc,
    f"主巡检 S {main_s} / M {main_m} / G {main_g} / L {main_l} = {sum(MAIN_SEV.values())} vs 全月 S {all_s} / M {all_m} / G {all_g} / L {all_l} = {total_findings_all}。"
    f" 5 月每店仅 1 次巡检（{SNAPSHOT_DAYS} 天 snapshot），主巡检 = 全月；待月底完整数据时本对比将体现真正差异。",
    kind="info", title="📊 主巡检 vs 全月对比")

add_heading(doc, "3.2 S 项详情（主巡检） ── 必须立即整改", 2)
main_s_items = sorted([it for it in MAIN_ITEMS if it["severity"] == "S"],
                      key=lambda x: (x["store_code"], x["inspection_date"]))
add_para(doc, f"主巡检共发现 {len(main_s_items)} 个 S 项。", size=10, space_after=4)
if main_s_items:
    rows_s = []
    for i, it in enumerate(main_s_items, 1):
        desc_raw = it["clean_desc"].replace("\n"," ").replace("\r"," ")
        desc = desc_raw[:160] if desc_raw else "（描述空白，需 QA 复核补充）"
        rows_s.append([str(i),
                       f"{it['store_name']}\n{it['store_code']}",
                       f"{it['cn_module']}\n{it['module_subcategory']}",
                       desc,
                       (str(it['deduction_points']), {"align":"center"}),
                       it['inspection_type'],
                       it['inspector_name'],
                       it['inspection_date']])
    add_data_table(doc, ["#","门店","模块/子项","问题描述（原文）","扣分","巡检类型","巡检员","日期"],
                   rows_s, col_widths=[0.6,2.2,2.6,5.6,1.0,1.5,2.0,1.4],
                   align_center_cols=[0,4])
else:
    add_para(doc, "本月主巡检无 S 项案例。", size=10, color_hex="666666")

add_heading(doc, "3.3 全月 S 项汇总（含自检与重复巡检）", 2)
add_para(doc, f"全月共 {len(all_s_items)} 个 S 项（5 月每店仅 1 次巡检，主巡检 = 全月），按子项汇总：",
         size=10, space_after=4)
sub_roll = defaultdict(lambda: {"cnt":0, "stores":set(), "samples":[]})
for it in all_s_items:
    k = it["module_subcategory"]
    sub_roll[k]["cnt"] += 1
    sub_roll[k]["stores"].add(it["store_code"])
    if it["clean_desc"] and len(sub_roll[k]["samples"]) < 2:
        sub_roll[k]["samples"].append(it["clean_desc"])
roll_rows = []
for sub, info in sorted(sub_roll.items(), key=lambda kv: -kv[1]["cnt"]):
    sample = (info["samples"][0] if info["samples"] else "（本子项可引用描述均为空）")[:120].replace("\n"," ")
    roll_rows.append([sub, str(info["cnt"]), str(len(info["stores"])), sample])
if roll_rows:
    add_data_table(doc, ["子项","S 项数","门店数","典型问题（截取）"], roll_rows,
                   col_widths=[3.6,1.0,1.0,11.0], align_center_cols=[1,2])
    add_para(doc, "S 项已计入 §6.2 P0 行动（air gap / Sinks and Pipes 跨月持续治理）。",
             size=9, color_hex="666666", space_before=4)
else:
    add_para(doc, "本月 snapshot 4 日内全月范围未触发 S 项。", size=10, color_hex="666666")

add_heading(doc, "3.4 M 项详情（主巡检） ── 7 天内闭环", 2)
main_m_items = sorted([it for it in MAIN_ITEMS if it["severity"] == "M"],
                      key=lambda x: (x["store_code"], x["inspection_date"]))
add_para(doc, f"主巡检共发现 {len(main_m_items)} 个 M 项。", size=10, space_after=4)
if main_m_items:
    rows_m = []
    for i, it in enumerate(main_m_items, 1):
        desc_raw = it["clean_desc"].replace("\n"," ").replace("\r"," ")
        desc = desc_raw[:140] if desc_raw else "（描述空白，需 QA 复核补充）"
        rows_m.append([str(i),
                       f"{it['store_name']}\n{it['store_code']}",
                       f"{it['cn_module']}\n{it['module_subcategory']}",
                       desc,
                       (str(it['deduction_points']), {"align":"center"}),
                       it['inspection_date']])
    add_data_table(doc, ["#","门店","模块/子项","问题描述（原文）","扣分","日期"], rows_m,
                   col_widths=[0.6,2.4,2.8,8.4,1.0,1.4], align_center_cols=[0,4])
else:
    add_para(doc, "本月主巡检无 M 项案例。", size=10, color_hex="666666")

add_heading(doc, "3.5 G 项 / L 项 分布（主巡检）", 2)
g_dist = Counter(it["cn_module"] for it in MAIN_ITEMS if it["severity"] == "G")
l_dist = Counter(it["cn_module"] for it in MAIN_ITEMS if it["severity"] == "L")
add_para(doc, f"G 项（一般项）共 {main_g} 个，主要集中模块：", size=10, bold=True, space_after=2)
for m,c in g_dist.most_common(): add_para(doc, f"  • {m}：{c} 个", size=10, indent_left=8, space_after=2)
add_para(doc, f"L 项（轻微项）共 {main_l} 个，主要集中模块：", size=10, bold=True, space_before=4, space_after=2)
for m,c in l_dist.most_common(): add_para(doc, f"  • {m}：{c} 个", size=10, indent_left=8, space_after=2)

# ============================================================================
# §四、模块与门店关联分析
# ============================================================================
add_heading(doc, "四、模块与门店关联分析", 1)

add_heading(doc, "4.1 门店 × 模块扣分矩阵（主巡检）", 2)
matrix_mods = MODULES_10
mat_headers = ["门店"] + [MODULE_SHORT[m] for m in matrix_mods] + ["合计"]
store_total = {sc: sum(MAIN_STORE_MOD.get(sc, {}).values()) for sc in ACTIVE_STORE_CODES}
ordered = sorted(ACTIVE_STORE_CODES, key=lambda sc: store_total[sc])
mat_rows = []
for sc in ordered:
    name = MAIN_INSP[sc]["store_name"]
    row = [f"{name}\n({sc})"]
    for mod in matrix_mods:
        v = MAIN_STORE_MOD.get(sc, {}).get(mod, 0)
        row.append((str(v) if v else "", {"align":"center"}))
    row.append((str(store_total[sc]), {"bold": True, "align":"center"}))
    mat_rows.append(row)
add_data_table(doc, mat_headers, mat_rows,
               col_widths=[2.6]+[1.15]*10+[1.2], align_center_cols=list(range(1,12)))
add_para(doc, f"※ 本月暂无申诉立案；矩阵覆盖 {N_STORES_INSPECTED} 家已巡检门店（4 月可比基线 {APR_COHORT_N} 家中 6 家未在 5 月前 4 日提交主巡检）。",
         size=9, color_hex="666666", space_before=2, space_after=4)

add_heading(doc, "4.2 最低分门店归因", 2)
low_mod = sorted(MAIN_STORE_MOD.get(lowest["store_code"], {}).items(), key=lambda kv: kv[1])[:3]
low_mod_text = "、".join(f"{m}（{v} 分）" for m,v in low_mod) or "（无扣分项）"
add_para(doc,
    f"最低分主巡检门店：{lowest['store_name']}（{lowest['store_code']}），主巡检得分 {lowest['adjusted_total_score']} 分，扣分 {lowest['adjusted_total_deduction']} 分，集中在：{low_mod_text}。",
    size=10, space_after=4)
add_para(doc,
    f"巡检类型为 {lowest['inspection_type']}（{lowest['inspection_date']}），巡检员 {lowest['inspector_name']}。",
    size=10, space_after=8)

apr_score = apr_main_score_for(lowest["store_code"])
apr_main_row = APR_MAIN.get(lowest["store_code"])
if apr_score is not None:
    delta = lowest["adjusted_total_score"] - apr_score
    direction = "回升" if delta > 0 else ("持平" if delta == 0 else "下滑")
    apr_type = apr_main_row["inspection_type"] if apr_main_row else ""
    add_callout(doc,
        f"对比 4 月主巡检（含申诉调整）：{lowest['store_name']} {apr_score} 分（{apr_type}）→ 5 月 {lowest['adjusted_total_score']} 分（{lowest['inspection_type']}）（{'+' if delta>0 else ''}{delta} 分变动反映 {direction}）。注意 4 月主巡检为 QA 审计、5 月主巡检为 {lowest['inspection_type']}，口径差异需在月底 QA 审计完成后重新校准；同店首次出现 S 项（Sinks and Pipes），与 4 月 air gap 系统性短板同源。",
        kind="critical" if delta <= -10 else "info", title="📈 跨月对比 (snapshot)")
else:
    add_callout(doc, f"本店为本月首次巡检覆盖（4 月主巡检数据缺失，无法跨月比对）。", kind="info", title="📈 跨月对比 (snapshot)")

add_heading(doc, "4.3 申诉调整门店分析", 2)
if APPEALED:
    # placeholder for future appeals
    add_para(doc, f"本月共 {len(APPEALED)} 起申诉立案。", size=10, space_after=6)
else:
    add_para(doc, "本月暂无申诉立案。4 月已闭环的 6 起申诉（2 获批：54th & 8th 69→94、15th & 3rd 71→96；3 驳回：8th & Broadway QA 4/9、102 Fulton QA 4/27、28th & 6th QA 4/27；1 审批中：28th & 6th 区经 4/23）的实际整改进度跟踪将在 5 月月度版（6 月初生成）补充。",
             size=10, color_hex="666666", space_after=4)

# §4.4 SKIP per snapshot adjustment #10
add_heading(doc, "4.4 同店同期评分背离案例", 2)
add_para(doc, f"5 月 4 日截至，多数门店仅 1 次巡检，跨类型背离分析待月底补充。",
         size=10, color_hex="666666", space_after=4)
add_para(doc, f"详细：{N_STORES_INSPECTED} 家已巡检门店中，每家均仅有 1 次提交（自检 / 区经检查互斥），无 cross-type 对比基础。月底 QA 审计完成后将出现自检-QA / 区经-QA 多组对比，届时本节按 4 月口径补充 ≥20 分背离案例。",
         size=10, color_hex="666666", space_after=4)

add_heading(doc, "4.5 模块覆盖面分析（主巡检视角）", 2)
cov_rows = []
for mod in sorted(mods_with_data, key=lambda m: -len(MAIN_MOD_STORES.get(m, set()))):
    n_st = len(MAIN_MOD_STORES.get(mod, set()))
    if n_st == 0: continue
    cov_pct = round(n_st / N_STORES_INSPECTED * 100)
    ded = sum(MAIN_MOD_DED.get(mod, {}).values())
    sev = MAIN_MOD_SEV.get(mod, {})
    risks = []
    if cov_pct >= 50: risks.append("⚠ 系统性")
    if sev.get("S",0) > 0: risks.append("⚠ 含 S 项")
    if not risks and sev.get("M",0) > 0: risks.append("含 M 项")
    if not risks: risks.append("---")
    cov_rows.append([mod, f"{n_st}/{N_STORES_INSPECTED}", f"{cov_pct}%", str(ded), " / ".join(risks)])
add_data_table(doc, ["模块","影响门店","覆盖率","扣分","风险标记"], cov_rows,
               col_widths=[3.6,2.0,1.4,1.4,5.0], align_center_cols=[1,2,3])
add_para(doc, f"※ 覆盖率分母为已巡检 {N_STORES_INSPECTED} 家门店；月底完整数据时将按 {APR_COHORT_N} 家可比基线重新计算。",
         size=9, color_hex="666666", space_before=2, space_after=4)

# ============================================================================
# §五、整改归因与效率
# ============================================================================
add_heading(doc, "五、整改归因与效率", 1)

add_heading(doc, "5.1 关键词归因（主巡检）", 2)
add_para(doc, "⚠ 以下归因基于问题描述关键词自动匹配，仅供参考。实际归因需以整改工单系统数据为准。",
         size=9, color_hex=WARN_FG, space_after=4)
def attribute(text):
    t = (text or "").strip().lower()
    if len(t) < 10: return "未知"
    pipe = ["pipe","sink","leak","airgap","air gap","drain","handwashing sink","plumbing","fixture","light","ceiling","floor","wall"]
    sup = ["license","sign","no smoking","permit","document","expiration date"]
    store_kw = ["clean","sanitize","glove","wiping","cloth","sticker","cup","container","ppm","sanitizer","storage","dirty","residue","spill"]
    if any(k in t for k in pipe): return "机修+营建"
    if any(k in t for k in sup):  return "供应链+行政"
    if any(k in t for k in store_kw): return "门店"
    return "门店"
attr_cnt = Counter(attribute(it["clean_desc"]) for it in MAIN_ITEMS)
attr_total = max(1, sum(attr_cnt.values()))
attr_rows = []
typical = {"门店":"日常清洁、消毒、标签、卫生",
           "机修+营建":"sinks and pipes、air gap、light fixtures",
           "供应链+行政":"License / No smoking sign / 文件记录",
           "未知":"描述模糊或缺失（少于 10 字符；本月多家自检未填写明细）"}
for cat in ["门店","机修+营建","供应链+行政","未知"]:
    c = attr_cnt.get(cat, 0)
    pct = f"{c/attr_total*100:.1f}%"
    attr_rows.append([cat, str(c), pct, typical[cat]])
add_data_table(doc, ["归因类别","数量","占比","典型问题"], attr_rows,
               col_widths=[2.4,1.0,1.0,9.0], align_center_cols=[1,2])
add_para(doc, f"※ snapshot 4 日窗口内多家门店自检（US00001、US00015）未填写问题描述，导致「未知」占比偏高，需 empapp 系统强化必填校验。",
         size=9, color_hex="666666", space_before=2, space_after=4)

add_heading(doc, "5.2 SLA 整改时限标准", 2)
add_data_table(doc, ["风险等级","整改时限","要求"], [
    ["S 项（关键项）", "2 天", "发现后 2 天内完成整改并验证"],
    ["M 项（重要项）", "7 天", "发现后 7 天内完成整改并验证"],
    ["G 项（一般项)", "14 天","发现后 14 天内完成整改并验证"],
    ["L 项（轻微项）", "14 天","发现后 14 天内完成整改并验证"],
], col_widths=[2.5,2.0,9.0], align_center_cols=[1])

add_heading(doc, "5.3 建议整改闭环流程", 2)
flow = [
    "巡检发现问题 → empapp 自动生成整改工单",
    "根据问题类型自动分配责任方（门店 / 机修 / 营建 / 供应链）",
    "责任方在 SLA 时限内完成整改",
    "QA 复核验证整改效果",
    "系统记录闭环时间，计算 SLA 达标率",
]
for i, ln in enumerate(flow, 1):
    add_para(doc, f"{i}. {ln}", size=10, indent_left=12, space_after=3)

# ============================================================================
# §六、建议与下一步行动
# ============================================================================
add_heading(doc, "六、建议与下一步行动", 1)

add_heading(doc, "6.1 本月关键发现", 2)
findings = [
    f"🟡 5 月节奏 (snapshot)：前 4 日完成 {N_TOTAL} 次巡检（自检 {N_BY_TYPE['门店自检']} / QA {N_BY_TYPE['QA审计']} / 区经 {N_BY_TYPE['区经检查']}），覆盖 {N_STORES_INSPECTED}/{APR_COHORT_N} = {N_STORES_INSPECTED*100/APR_COHORT_N:.1f}% 门店；QA 审计尚未启动需关注。",
    f"⚠ Sinks and Pipes 系统性短板再次命中：4 月 air gap / Sinks and Pipes 累计 10 起 S 项（跨 6 家门店）尚未根治，5/3 区经检查在 102 Fulton 又触发 1 起 S 项；需 BD 整改清单优先级升级。",
    f"⚠ 102 Fulton 跨月下滑 16 分：4 月 QA 86 → 5 月 区经 70（同时含 1 S 项 / 1 M 项）；同店首次进入 <80 分区间，需专项辅导。",
    f"✅ 41st & Lexington 新开店首检 90 分：4/30 开业 5/1 即由 Afsana Gu 完成首次自检，得分 90 分（4 G 2 L 0 S），新开店初期巡检节奏与前期良好；待月底 QA 审计验证。",
    f"⚠ 自检描述完整度待提升：8th & Broadway（16 项全空）、41st & Lexington（6 项全空）等门店自检未填写问题描述，影响整改归因；empapp 系统需强化描述必填校验。",
    f"🟡 QA 单点依赖延续：Eamonn Caballar 4 月 12 次 / 5 月前 4 日 0 次（典型月初轻量节奏）；Yu Jiang HR 流程仍未关闭（4 月 0 / 5 月 0）；评估第二位 QA Manager 配置时机。",
    f"✅ 区经巡检节奏稳定：Daniel Chu 已完成 3 次（覆盖 100 Maiden Ln、102 Fulton、221 Grand），与 4 月节奏一致；Jung Han Liang 前 4 日尚无提交，月内需补足。",
]
for ln in findings:
    add_para(doc, "• " + ln, size=10, indent_left=12, space_after=4)

add_heading(doc, "6.2 优先行动项", 2)
prio_rows = [
    ["P0","紧急", f"立即处理 102 Fulton（US00006）的 1 个 S 项（Sinks and Pipes）+ 1 个 M 项，48 小时内闭环；BD 整改清单接续 4 月 air gap 跨门店治理（4 月 10 起未根治）",
     "门店 + QA + BD","48 小时"],
    ["P0【接续】","紧急", f"4 月遗留 S 项跟踪：54th & 8th / 15th & 3rd（申诉获批后 air gap / plumbing 整改）+ 4 月 19 起 S 项中 BD 待闭环部分", "BD + QA","本月内"],
    ["P1","紧急", f"处理本月 {main_m} 个 M 项（28th & 6th sani ppm、37th & Broadway sani ppm、100 Maiden Ln handwashing 等），重点 100 Maiden Ln（员工失手未洗手）", "门店 + QA","7 天"],
    ["P2","高",   f"5 月剩余 {DAYS_REMAINING} 天 BD 整改 / QA 审计 / 区经检查目标：QA 审计预计 12+ 次（覆盖 13 家可比门店）、区经检查预计 12+ 次，月底前完成 5 月全覆盖", "QA 部门 + 运营部","5 月内"],
    ["P3","高",   f"对最低分门店（102 Fulton 70 分、8th & Broadway 77 分）开展专项辅导；41st & Lexington 等新开店纳入新店训练计划", "区域经理","2 周内"],
    ["P4","高",   f"empapp 系统优化：自检描述必填校验（本月 22 项空描述均来自 8th & Broadway / 41st & Lexington 自检）", "基础设施部","2 周内"],
    ["P5","中",   f"评估 QA 团队冗余度：Yu Jiang HR 流程关闭（5 月 0 / 4 月 0）、第二位 QA Manager 配置时机", "QA 部门 + HR","本月内"],
    ["P6","中",   f"新开店覆盖：US00007（4/30）、US00010（4/28）、US00019（4/11）截至 5/4 仍未首检，需 5 月内完成首次 QA / 区经", "运营部 + QA","5 月内"],
    ["P7","中",   f"4 月申诉案例后续：3 起驳回（按原方案闭环验证）+ 1 起审批中（SLA 监控）", "QA 部门","2 周内"],
    ["P8","低",   f"4 月 §7.7 遗留事项【接续】：① 同日重复自检 UI 提示；② S 项整改工单与申诉立案双向同步；③ 申诉审批 SLA 监控", "基础设施部 + QA","本季度"],
]
add_data_table(doc, ["优先级","紧急度","行动项","责任方","时限"], prio_rows,
               col_widths=[1.2,1.0,8.4,3.0,2.0], align_center_cols=[0,1,4])

add_heading(doc, "6.3 模块改善建议（TOP 5）", 2)
suggestions = {
    "清洁卫生": "① 清洁消毒程序每班次执行并记录；② 消毒液浓度（ppm）每日校准（28th & 6th / 37th & Broadway 5 月已连续两次 ppm 异常）；③ 食品加工区域和设备每日深度清洁。",
    "过程控制": "① 食品存储分区标准重新培训（lock box 5 月再次出现未上锁）；② 器具维护和清洁班次检查；③ 物料存储高度要求（6 英寸）每日巡查。",
    "设施":     "① air gap / Sinks and Pipes 跨月持续治理：4 月 10 起 + 5 月 1 起未根治，由 BD 列入跨门店整改清单（最高优先级）；② 油脂阱 / 残渣阱清理纳入月度必检；③ 管道泄漏立即报修。",
    "员工健康卫生": "① 100 Maiden Ln「员工失手未洗手」案例纳入 5 月全员复训；② 个人卫生（指甲、首饰、头发）班前班中检查；③ 洗手程序每周复训。",
    "温控有效期": "① 102 Fulton 过期标签「不可读」案例：开封后标签管理纳入每日开店清单；② FIFO 执行每日检查；③ 过期产品零容忍政策。",
    "虫害防控": "① pest control 每月报告留档；② 灯诱设备完好性每周巡查。",
    "设备维护": "① 设备保养排期每周更新；② 关键设备（搅拌机、咖啡机）每季度专业校准。",
}
for mod in top_mods:
    if mod in suggestions:
        ded = sum(MAIN_MOD_DED.get(mod, {}).values())
        n_st = len(MAIN_MOD_STORES.get(mod, set()))
        add_para(doc, f"{mod}（影响 {n_st} 家门店，主巡检扣分 {ded} 分）：",
                 size=10, bold=True, color_hex=NAVY, space_before=4, space_after=2)
        add_para(doc, suggestions[mod], size=10, indent_left=12, space_after=4)

# ============================================================================
# §七、巡检体系分析
# ============================================================================
add_heading(doc, f"七、三类巡检体系分析（{TARGET_MONTH_CN} 4日截至 snapshot）", 1)

add_callout(doc,
    f"5 月前 4 日体系节奏：4 月已完成体系全面恢复（58 次三类齐全），5 月前 4 日典型月初轻量节奏（{N_TOTAL} 次提交）；QA 审计尚未启动、Jung Han Liang 区经检查尚无提交，需在剩余 {DAYS_REMAINING} 天内补足以维持 4 月已建立的全覆盖节奏。",
    kind="info", title="🟡 月初节奏 (snapshot)")

add_heading(doc, "7.1 巡检概况", 2)
def avg_score_by_type(t, use_adjusted=True):
    rs = INSP_BY_TYPE[t]
    if not rs: return ""
    return round(sum((r["adjusted_total_score"] if use_adjusted else r["original_total_score"]) for r in rs)/len(rs), 1)
def stores_for_type(t): return len({r["store_code"] for r in INSP_BY_TYPE[t]})
def s_for_type(t): return sum(r["s_count"] for r in INSP_BY_TYPE[t])
def m_for_type(t): return sum(r["m_count"] for r in INSP_BY_TYPE[t])

def insp_list_text(t):
    cnt = Counter(r["inspector_name"] for r in INSP_BY_TYPE[t])
    if len(cnt) == 0: return "—"
    if len(cnt) <= 5:
        return "、".join(sorted(cnt.keys()))
    top3 = "、".join(name for name,_ in cnt.most_common(3))
    return f"{len(cnt)} 人（含 {top3}）"

def fmt_avg(t):
    v = avg_score_by_type(t)
    return f"{v} 分" if v != "" else "— (无提交)"

ct_rows = [
    ["巡检次数（已提交）", f"{N_BY_TYPE['门店自检']} 次", f"{N_BY_TYPE['QA审计']} 次", f"{N_BY_TYPE['区经检查']} 次"],
    ["覆盖门店", f"{stores_for_type('门店自检')} 家", f"{stores_for_type('QA审计')} 家", f"{stores_for_type('区经检查')} 家"],
    ["巡检员", insp_list_text("门店自检"), insp_list_text("QA审计"), insp_list_text("区经检查")],
    ["平均得分", fmt_avg("门店自检"), fmt_avg("QA审计"), fmt_avg("区经检查")],
    ["S 项发现", str(s_for_type('门店自检')), str(s_for_type('QA审计')), str(s_for_type('区经检查'))],
    ["M 项发现", str(m_for_type('门店自检')), str(m_for_type('QA审计')), str(m_for_type('区经检查'))],
]
add_data_table(doc, ["维度","门店自检","QA审计","区经检查"], ct_rows,
               col_widths=[2.5,4.5,4.5,4.5], align_center_cols=[1,2,3])
add_para(doc, "※ 5 月 QA 审计 0 次（Eamonn Caballar 月初未启动），平均分无法计算；待月底重新评估。",
         size=9, color_hex="666666", space_before=2, space_after=4)

add_callout(doc,
    f"自检 vs 区经差距：自检平均 {avg_score_by_type('门店自检')} 分（{N_BY_TYPE['门店自检']} 次），区经平均 {avg_score_by_type('区经检查')} 分（{N_BY_TYPE['区经检查']} 次）。差距 {abs(float(avg_score_by_type('区经检查'))-float(avg_score_by_type('门店自检'))):.1f} 分（区经偏严），与 4 月趋势（自检 80.2 / 区经 80.1，差距 0.1）相比反向，但 5 月样本极小（自检 4 + 区经 3），不足以下结论；待月底 QA 审计加入后重新校准。",
    kind="info", title="📊 关键观察 (snapshot)")

add_heading(doc, "7.2 同店三类对比", 2)
type_avg = defaultdict(dict)
for sc, rows in by_store.items():
    for t in ("门店自检","QA审计","区经检查"):
        rs = [r for r in rows if r["inspection_type"] == t]
        if rs:
            type_avg[sc][t] = round(sum(r["adjusted_total_score"] for r in rs)/len(rs), 1)
dt_rows = []
for sc in sorted(type_avg.keys()):
    name = MAIN_INSP[sc]["store_name"]
    s_self = type_avg[sc].get("门店自检", "---")
    s_qa = type_avg[sc].get("QA审计", "---")
    s_area = type_avg[sc].get("区经检查", "---")
    diff_str, conclusion = "---", "—（5 月仅 1 类巡检，无对比）"
    dt_rows.append([f"{name} ({sc})", str(s_self), str(s_qa), str(s_area), diff_str, conclusion])
add_data_table(doc, ["门店","自检均分","QA审计","区经检查","自检 - QA 差","结论"], dt_rows,
               col_widths=[3.6,2.0,2.0,2.0,2.0,4.0], align_center_cols=[1,2,3,4,5])
add_para(doc, f"※ 5 月每店仅 1 次巡检（自检 / 区经互斥），跨类型对比基础尚未形成；待月底 QA 审计完成后于本节按 4 月口径补充。",
         size=9, color_hex="666666", space_before=2, space_after=4)

# §7.3 SKIP per snapshot adjustment #11
add_heading(doc, "7.3 自检评分一致性分析", 2)
add_para(doc, f"5 月 4 日截至，多数门店仅 1 次自检，自检评分一致性分析待月底补充。",
         size=10, color_hex="666666", space_after=4)
add_para(doc, f"详细：{N_BY_TYPE['门店自检']} 次自检覆盖 {stores_for_type('门店自检')} 家不同门店（每家 1 次），无同店重复自检数据；4 月 US00020 同日三次自检（100/100/64 摆动 36 分）案例的 empapp 系统校验上线情况将在月度版跟踪。",
         size=10, color_hex="666666", space_after=4)

add_heading(doc, "7.4 巡检员严格度对比（≥2 次巡检者）", 2)
inspector_stats = defaultdict(lambda: {"cnt":0,"score_sum":0,"ded_sum":0,"item_sum":0,"s":0,"m":0,
                                        "type":Counter(),"role":""})
for r in SUMMARY:
    n = r["inspector_name"]; st = inspector_stats[n]
    st["cnt"] += 1
    st["score_sum"] += r["adjusted_total_score"]
    st["ded_sum"]   += r["adjusted_total_deduction"]
    st["item_sum"]  += r["item_count"]
    st["s"] += r["s_count"]; st["m"] += r["m_count"]
    st["type"][r["inspection_type"]] += 1
    st["role"] = r["inspector_role"]

multi_insp = [(n,st) for n,st in inspector_stats.items() if st["cnt"] >= 2]
single_insp_with_s = [(n,st) for n,st in inspector_stats.items() if st["cnt"] == 1 and st["s"] > 0]

if multi_insp:
    multi_insp.sort(key=lambda kv: kv[1]["score_sum"]/kv[1]["cnt"])
    ins_rows = []
    for n, st in multi_insp:
        avg_s = st["score_sum"]/st["cnt"]
        note = "⚠ 偏严" if avg_s < 70 else ("⚠ 偏宽" if avg_s > 92 else "—")
        ins_rows.append([n, st["role"],
                         st["type"].most_common(1)[0][0],
                         str(st["cnt"]),
                         f"{avg_s:.1f}",
                         note])
    add_data_table(doc, ["巡检员","职位","类型主导","巡检次","平均分","说明"],
                   ins_rows, col_widths=[2.8,3.5,2.0,1.2,1.4,2.0],
                   align_center_cols=[3,4,5])
    add_para(doc, "※ 5 月前 4 日仅 Daniel Chu 满足 ≥2 次门槛；其余巡检员均仅 1 次提交，待月底重新评估。",
             size=9, color_hex="666666", space_before=2, space_after=4)
else:
    add_para(doc, "5 月前 4 日无 ≥2 次巡检的重复巡检员案例。", size=10, color_hex="666666")

if single_insp_with_s:
    add_para(doc, "单次巡检但发现 S 项（重要发现，单独标注）：", size=10, bold=True, space_before=4, space_after=2)
    for n, st in single_insp_with_s:
        t = st["type"].most_common(1)[0][0]
        add_para(doc, f"  • {n}（{st['role']}，{t}，{st['cnt']} 次）：发现 S 项 {st['s']} 个 / 平均分 {st['score_sum']/st['cnt']:.1f}",
                 size=10, indent_left=8, space_after=2)

# Inspector roster diff vs April
add_para(doc, "巡检员名册变动（vs 4 月）：", size=10, bold=True, space_before=6, space_after=2)
roster_lines = []
ret = []  # returned (apr=0 / older months>0 / may≥1)
ssil = [] # silent (apr≥3 / may=0)
for tr in INSP_TREND:
    apr = int(tr["apr_count"]); may = int(tr["may_count"])
    if may >= 1 and apr == 0:
        ret.append(f"{tr['inspector_name']}（{tr['inspector_role']}，4 月 0 次 → 5 月 {may} 次）")
    if apr >= 3 and may == 0:
        ssil.append(f"{tr['inspector_name']}（{tr['inspector_role']}，4 月 {apr} 次 → 5 月前 4 日 0 次）")
if ret:
    add_para(doc, f"  • 回归 / 新增 ({len(ret)} 人)：" + "；".join(ret), size=10, indent_left=8, space_after=2)
if ssil:
    add_para(doc, f"  • 月初尚未提交 ({len(ssil)} 人活跃于 4 月)：" + "；".join(ssil) + "——属月初节奏，月底重新评估。", size=10, indent_left=8, space_after=2)
add_para(doc, "  • Yu Jiang（Senior QA Manager）：4 月 0 次 / 5 月前 4 日 0 次，HR 离职 / 转岗流程仍未关闭。",
         size=10, indent_left=8, space_after=4)

add_heading(doc, "7.5 巡检覆盖趋势（2026 年 Q1 + 4 月 + 5 月在途）", 2)
trend_data = defaultdict(lambda: {"门店自检":0,"QA审计":0,"区经检查":0})
for tr in TREND_SUMMARY:
    if tr["month"] in ("2026-01","2026-02","2026-03","2026-04","2026-05"):
        trend_data[tr["month"]][tr["inspection_type"]] = int(tr["inspection_count"])
status_for = {"2026-01":"✅ 三类齐全","2026-02":"⚠ 区经检查中断",
              "2026-03":"🔴 体系崩溃","2026-04":"✅ 全面恢复",
              "2026-05":"🟡 5月在途（4日 snapshot）"}
trd_rows = []
for ym in ("2026-01","2026-02","2026-03","2026-04","2026-05"):
    d = trend_data[ym]
    total = d["门店自检"] + d["QA审计"] + d["区经检查"]
    trd_rows.append([ym, f"{d['门店自检']} 次", f"{d['QA审计']} 次", f"{d['区经检查']} 次", str(total), status_for[ym]])
add_data_table(doc, ["月份","门店自检","QA审计","区经检查","总数","状态"], trd_rows,
               col_widths=[2.0,2.0,2.0,2.0,1.4,4.0], align_center_cols=[1,2,3,4])
add_callout(doc,
    f"5 月前 4 日轨迹：自检 {trend_data['2026-05']['门店自检']} / QA {trend_data['2026-05']['QA审计']} / 区经 {trend_data['2026-05']['区经检查']} = {sum(trend_data['2026-05'].values())} 次。按 4 月节奏（58 次/30 天 ≈ 1.93/天）线性外推，5 月预期 60 次左右；但前 4 日仅 7 次（1.75/天）略低于 4 月节奏，QA 审计尚未启动是主因。需在剩余 {DAYS_REMAINING} 天内由 Eamonn Caballar 启动 QA 审计周期以维持月度全覆盖。",
    kind="info", title="📈 趋势分析 (snapshot)")

add_heading(doc, "7.6 三类巡检发现差异分析", 2)
type_diff_rows = [
    ["核心价值", "门店日常自查、低成本高频", "QA 标准化外审、培训校准", "区经实地督导、跨店共性整改"],
    ["最常发现的问题",
     "5 月：Equipment and utensils、Clean & Sanitize、Customer Area / Food Processing Area（清洁卫生类高频；多家自检未填写描述）",
     "5 月：尚未启动 — 沿用 4 月观察（Sinks and Pipes、Cross-Contamination、Handwashing Standards）",
     "5 月：Sinks and Pipes（102 Fulton S 项）、Equipment and utensils、Standard procedures（Daniel Chu 主导）"],
    ["典型严重度分布",
     "S 项暂无；M 项 2 个（28th & 6th sani ppm、37th & Broadway sani ppm）；G/L 项为主",
     "5 月暂无样本 — 沿用 4 月观察（S 项命中率最高，air gap 类系统性 S 项主要在 QA 发现）",
     "S 项 1 个（102 Fulton Sinks and Pipes）+ M 项 2 个，与 4 月跨店共性问题高度同源"],
]
add_data_table(doc, ["维度","门店自检","QA审计","区经检查"], type_diff_rows,
               col_widths=[2.6,4.5,4.5,4.5])
add_para(doc, f"※ 5 月 4 日截至样本不足，QA 审计列沿用 4 月观察；月底完整数据时本表将重新生成。",
         size=9, color_hex="666666", space_before=2, space_after=4)

add_heading(doc, f"7.7 后续观察重点（5 月剩余 {DAYS_REMAINING} 天 + 月度收口）", 2)
follow = [
    f"【本月剩余 27 天 BD 整改 / QA 审计 / 区经检查目标】",
    f"  ① BD 整改：4 月 19 起 S 项（air gap / Sinks and Pipes 10 起、Handwashing Standards 5 起、其他 4 起）+ 5 月 102 Fulton 1 起 S 项的实际闭环验证；建立跨月 S 项追踪台账。",
    f"  ② QA 审计：Eamonn Caballar 5 月剩余 {DAYS_REMAINING} 天启动 QA 周期，目标 12+ 次（覆盖 13 家可比门店 + 含新开店 US00007/00010/00015/00019）。",
    f"  ③ 区经检查：Daniel Chu 已完成 3 次（节奏稳定），Jung Han Liang 5 月内补足 5+ 次以达 4 月节奏（每人 7 次）。",
    f"【5 月底数据完整后需补全】",
    f"  ④ §4.4 同店同期评分背离案例（5 月每店 ≥2 次巡检后形成自然样本）；",
    f"  ⑤ §7.3 自检评分一致性分析（同人多次自检案例）；",
    f"  ⑥ §1.2 跨月 ※ 申诉调整门店标记（若 5 月新增申诉立案）。",
    f"【4 月遗留事项跟踪 [接续]】",
    f"  ⑦ 申诉案例后续：54th & 8th 与 15th & 3rd 申诉获批后 air gap / plumbing 实际整改完成情况；驳回的 3 起（8th & Broadway QA 4/9、102 Fulton QA 4/27、28th & 6th QA 4/27）按原方案闭环；审批中的 1 起（28th & 6th 区经 4/23）审批 SLA。",
    f"  ⑧ QA workload sustainability：第二位 QA Manager 配置时机，避免 Eamonn 单点依赖。",
    f"  ⑨ Yu Jiang HR 流程关闭确认（4 月 0 / 5 月前 4 日 0，需走流程）。",
    f"  ⑩ 区经检查频率维持：Daniel Chu / Jung Han Liang 月度均衡负荷，避免 2-3 月断流复发。",
    f"  ⑪ 巡检员标准化培训：自检偏严（Brionna Jiles 59、Eric Park 63）与偏宽（Joselyn Pacheco Trejo 97、Juliana Li 96）双侧偏离的评分尺度统一培训。",
    f"  ⑫ 多门店 S 项系统短板（Sinks and Pipes air gap 跨月持续）：BD 整改清单完成情况。",
    f"  ⑬ 新开门店 5 月覆盖：US00007（4/30）、US00010（4/28）、US00019（4/11）首次 QA 审计 / 区经检查；US00015（4/30 开业，5/1 已完成首次自检 90 分）的后续 QA / 区经验证。",
    f"  ⑭ 自检评分一致性：US00020 同日三次自检案例的 empapp 系统校验上线情况。",
    f"  ⑮ empapp 系统优化项：① 同日重复自检 UI 提示；② S 项整改工单与申诉立案双向同步；③ 申诉审批 SLA 监控；④ 自检问题描述必填校验（5 月已暴露 22 项空描述）。",
    f"【snapshot 特有】下次 snapshot 时点建议：5/15（月中）+ 5/25（月末预演），便于及早发现节奏偏差。",
]
for ln in follow:
    if ln.startswith("【"):
        add_para(doc, ln, size=10, bold=True, color_hex=NAVY, space_before=6, space_after=2)
    elif ln.startswith("  "):
        add_para(doc, ln, size=10, indent_left=18, space_after=3)
    else:
        add_para(doc, "• " + ln, size=10, indent_left=12, space_after=4)

# ---------- Footer ----------
add_para(doc, "── 报告结束 ──", size=11, color_hex=NAVY, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=18, space_after=4)

# Save
doc.save(OUT_PATH)
print(f"\nWROTE: {OUT_PATH}")
print(f"  size = {OUT_PATH.stat().st_size:,} bytes")
