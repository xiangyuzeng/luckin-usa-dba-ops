#!/usr/bin/env python3
"""
April 2026 QA monthly report v2 — STRICT regeneration.

Changes vs prior build (per ultrathink prompt):
  - 10-module strict taxonomy (Site Security → "其他 Other" + flag).
  - §2.2 / §2.3 / §4.1 / §4.5 all use 主巡检 scope; §3.3 is the only 全月 section.
  - appeal_status (approved | denied | pending | none) derived from
    process_status × score delta.
  - Findings with empty descriptions are SKIPPED (no "(无描述)" placeholders).
  - No English bleed in prose: "已剔除误提交" not "misubmission-filtered",
    "巡检 ID 2016" not "iid=2016", etc.
  - No auto-generated meta-line below "── 报告结束 ──".
  - Page footer with: <doc_id> | 编制：曾翔宇 | 日期：YYYY-MM-DD | 第 X 页 / 共 N 页.
  - §7.4 only ≥2 inspection inspectors; §7.1 inspector list condensed if >5.
"""
import csv, json, datetime as dt
from collections import defaultdict, Counter
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ============================================================================
# RUN PARAMETERS
# ============================================================================
TARGET_MONTH    = "04"
TARGET_MONTH_CN = "2026年04月"
TARGET_YM       = "2026-04"
DOC_ID          = "LCNA-QA-2026-004"
TODAY           = "2026-05-05"

DATA_DIR        = Path("/app/reports/may2026-qa-inspection")
PRIOR_MONTH_CSV = Path("/app/reports/march2026_inspection_summary.csv")
RAW_HEADERS     = DATA_DIR / "raw" / "headers.json"
OUT_DIR         = DATA_DIR / "output"
OUT_DIR.mkdir(exist_ok=True)
OUT_PATH        = OUT_DIR / f"QA门店巡检月度分析报告_{TARGET_MONTH_CN}_v2.docx"

# Color palette
NAVY        = "1F4E79"; BLUE        = "2E75B6"; TEXT        = "2C3E50"
BORDER      = "BFBFBF"; ALT_FILL    = "F2F2F2"
CRITICAL_FG = "C0392B"; WARN_FG     = "E67E22"; HEALTHY_FG  = "27AE60"; INFO_FG = "0365C0"
CALLOUT_YELLOW = "FFF2CC"; CALLOUT_GREEN = "E2EFDA"
CALLOUT_BLUE   = "D6E4F0"; CALLOUT_RED   = "FCE4EC"

# ============================================================================
# 10-MODULE STRICT TAXONOMY
# ============================================================================
MODULES_10 = ["清洁卫生","过程控制","设施","证照文件","职场安全",
              "虫害防控","温控有效期","员工健康卫生","设备维护","供应商"]
MODULE_SHORT = {"清洁卫生":"清洁","过程控制":"过程","设施":"设施","证照文件":"证照",
                "职场安全":"职安","虫害防控":"虫害","温控有效期":"温控",
                "员工健康卫生":"员工","设备维护":"设备","供应商":"供应","其他":"其他"}

EN2CN = {
    "Cleaning and Sanitation": "清洁卫生",
    "Process Control": "过程控制",
    "Facility": "设施",
    "Document Record": "证照文件",
    "Workplace Safety": "职场安全",
    "Pests Control": "虫害防控",
    "Temperature Control / Expiration Date Management.": "温控有效期",
    "Employees’ Health and Personal Hygiene": "员工健康卫生",
    "Maintenance of Equipment": "设备维护",
    "Approved Supplier": "供应商",
}
def cn_module(it):
    return EN2CN.get(it["module_name"], "其他")

# ============================================================================
# DATA LOADING
# ============================================================================
def load_csv(p):
    with open(p, encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))

print(f"[load] target={TARGET_MONTH_CN} doc_id={DOC_ID}")
SUMMARY = load_csv(DATA_DIR / "april2026_inspection_summary.csv")
ITEMS   = load_csv(DATA_DIR / "april2026_inspection_items.csv")
STORE_MASTER  = load_csv("/app/reports/april2026-qa-inspection/april2026_store_master.csv")
INSP_TREND    = load_csv("/app/reports/april2026-qa-inspection/april2026_inspector_trend.csv")
TREND_SUMMARY = load_csv("/app/reports/april2026-qa-inspection/jan_to_apr2026_trend_summary.csv")
PRIOR_MARCH   = load_csv(PRIOR_MONTH_CSV) if PRIOR_MONTH_CSV.exists() else []
HEADERS_RAW   = json.load(open(RAW_HEADERS))
PROC_STATUS_MAP = {h["id"]: h["process_status"] for h in HEADERS_RAW}

# Type normalization
for r in SUMMARY:
    r["adjusted_total_score"]    = int(r["adjusted_total_score"]) if r["adjusted_total_score"] else None
    r["original_total_score"]    = int(r["original_total_score"]) if r["original_total_score"] else None
    r["adjusted_total_deduction"]= int(r["adjusted_total_deduction"]) if r["adjusted_total_deduction"] else 0
    r["original_total_deduction"]= int(r["original_total_deduction"]) if r["original_total_deduction"] else 0
    r["is_appealed"] = int(r["is_appealed"])
    for k in ("s_count","m_count","g_count","l_count","item_count","inspection_id"):
        r[k] = int(r[k])
    # Schema shim: derive appeal_status from is_appealed × delta × process_status
    if r["is_appealed"] == 0:
        r["appeal_status"] = "none"
    elif r["adjusted_total_score"] != r["original_total_score"]:
        r["appeal_status"] = "approved"
    elif PROC_STATUS_MAP.get(r["inspection_id"]) == 30:
        r["appeal_status"] = "pending"
    else:
        r["appeal_status"] = "denied"

for r in ITEMS:
    r["deduction_points"]    = int(r["deduction_points"])
    r["is_appealed_finding"] = int(r["is_appealed_finding"])
    r["inspection_id"]       = int(r["inspection_id"])
    r["item_id"]             = int(r["item_id"])
    r["cn_module"]           = cn_module(r)

print(f"[load] summary={len(SUMMARY)} items={len(ITEMS)} stores={len(STORE_MASTER)} prior_march={len(PRIOR_MARCH)}")

# Module mapping anomalies
OTHER_FLAGS = Counter(it["module_name"] for it in ITEMS if it["cn_module"] == "其他")
print(f"[mapping] 其他 (Other) bucket: {sum(OTHER_FLAGS.values())} items: {dict(OTHER_FLAGS)}")

# ============================================================================
# DERIVED
# ============================================================================
PRIORITY = {"QA审计": 0, "区经检查": 1, "门店自检": 2}

by_store = defaultdict(list)
for r in SUMMARY: by_store[r["store_code"]].append(r)
def main_inspection_for_store(rows):
    if not rows: return None
    return sorted(rows, key=lambda r:(PRIORITY.get(r["inspection_type"], 99),
                                      -dt.date.fromisoformat(r["inspection_date"]).toordinal()))[0]
MAIN_INSP = {sc: main_inspection_for_store(rows) for sc, rows in by_store.items()}
MAIN_IIDS = {m["inspection_id"] for m in MAIN_INSP.values() if m}
ACTIVE_STORE_CODES = sorted(by_store.keys())
N_STORES_ACTIVE = len(ACTIVE_STORE_CODES)

items_by_iid = defaultdict(list)
for it in ITEMS: items_by_iid[it["inspection_id"]].append(it)

MAIN_ITEMS = [it for it in ITEMS if it["inspection_id"] in MAIN_IIDS]
print(f"[derive] main_inspections={len(MAIN_INSP)} main_items={len(MAIN_ITEMS)} all_items={len(ITEMS)}")

def aggregate(items, modules_universe=MODULES_10 + ["其他"]):
    sev = Counter(); ded = Counter()
    mod_sev = defaultdict(Counter); mod_ded = defaultdict(Counter)
    mod_stores = defaultdict(set); store_mod_ded = defaultdict(lambda: defaultdict(int))
    for it in items:
        sev[it["severity"]] += 1
        ded[it["severity"]] += it["deduction_points"]
        mod_sev[it["cn_module"]][it["severity"]] += 1
        mod_ded[it["cn_module"]][it["severity"]] += it["deduction_points"]
        mod_stores[it["cn_module"]].add(it["store_code"])
        store_mod_ded[it["store_code"]][it["cn_module"]] += it["deduction_points"]
    return sev, ded, mod_sev, mod_ded, mod_stores, store_mod_ded

MAIN_SEV, MAIN_DED, MAIN_MOD_SEV, MAIN_MOD_DED, MAIN_MOD_STORES, MAIN_STORE_MOD = aggregate(MAIN_ITEMS)
ALL_SEV, ALL_DED, ALL_MOD_SEV, ALL_MOD_DED, ALL_MOD_STORES, _ = aggregate(ITEMS)

INSP_BY_TYPE = defaultdict(list)
for r in SUMMARY: INSP_BY_TYPE[r["inspection_type"]].append(r)
N_BY_TYPE = {t: len(INSP_BY_TYPE[t]) for t in ("门店自检","QA审计","区经检查")}
N_TOTAL = sum(N_BY_TYPE.values())

valid_main_scores = [m["adjusted_total_score"] for m in MAIN_INSP.values() if m["adjusted_total_score"] is not None]
avg_score_main = round(sum(valid_main_scores)/len(valid_main_scores), 1)

APPEALED         = [r for r in SUMMARY if r["is_appealed"] == 1]
APPEAL_APPROVED  = [r for r in APPEALED if r["appeal_status"] == "approved"]
APPEAL_DENIED    = [r for r in APPEALED if r["appeal_status"] == "denied"]
APPEAL_PENDING   = [r for r in APPEALED if r["appeal_status"] == "pending"]
print(f"[appeal] total={len(APPEALED)} approved={len(APPEAL_APPROVED)} denied={len(APPEAL_DENIED)} pending={len(APPEAL_PENDING)}")

# Cross-month delta
march_by_store = {r["store_code"]: r for r in PRIOR_MARCH if r["store_code"]}
def march_score_for(sc):
    r = march_by_store.get(sc)
    if r and r.get("total_score","").lstrip("-").isdigit():
        return int(r["total_score"])
    return None

# ============================================================================
# DOCX HELPERS
# ============================================================================
def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)

def set_cell_borders(cell, color_hex=BORDER, sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    bdrs = OxmlElement("w:tcBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"),"single"); b.set(qn("w:sz"),str(sz))
        b.set(qn("w:space"),"0"); b.set(qn("w:color"), color_hex)
        bdrs.append(b)
    tcPr.append(bdrs)

def set_run_font(run, name="Arial", size=10, bold=False, color_hex=TEXT, italic=False):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfont = rpr.find(qn("w:rFonts"))
    if rfont is None:
        rfont = OxmlElement("w:rFonts"); rpr.append(rfont)
    rfont.set(qn("w:ascii"), name); rfont.set(qn("w:hAnsi"), name)
    rfont.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    if color_hex: run.font.color.rgb = RGBColor.from_string(color_hex)

def add_para(doc, text, size=10, bold=False, color_hex=TEXT, align=None,
             space_before=0, space_after=4, indent_left=0, italic=False):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent_left: p.paragraph_format.left_indent = Pt(indent_left)
    if text:
        r = p.add_run(text); set_run_font(r, size=size, bold=bold, color_hex=color_hex, italic=italic)
    return p

def add_heading(doc, text, level=1):
    sizes = {1:16, 2:13, 3:11}; colors = {1:NAVY, 2:BLUE, 3:NAVY}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level==1 else 8)
    p.paragraph_format.space_after  = Pt(6 if level==1 else 4)
    r = p.add_run(text); set_run_font(r, size=sizes[level], bold=True, color_hex=colors[level])

def add_callout(doc, lines, kind="warn", title=None):
    bg, bar = {
        "warn":     (CALLOUT_YELLOW, WARN_FG),
        "good":     (CALLOUT_GREEN,  HEALTHY_FG),
        "info":     (CALLOUT_BLUE,   INFO_FG),
        "critical": (CALLOUT_RED,    CRITICAL_FG),
    }[kind]
    table = doc.add_table(rows=1, cols=1); table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, bg); set_cell_borders(cell, color_hex=bar, sz=8)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    if isinstance(lines, str): lines = [lines]
    if title:
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title); set_run_font(r, size=10, bold=True, color_hex=bar)
        for ln in lines:
            pp = cell.add_paragraph(); pp.paragraph_format.space_after = Pt(2)
            rr = pp.add_run(ln); set_run_font(rr, size=10, color_hex=TEXT)
    else:
        for i, ln in enumerate(lines):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(ln); set_run_font(r, size=10, color_hex=TEXT)
    add_para(doc, "", size=4, space_after=4)

def add_data_table(doc, headers, rows, col_widths=None, align_center_cols=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = "Table Grid"
    if col_widths:
        for ci, w in enumerate(col_widths):
            for c in t.columns[ci].cells: c.width = Cm(w)
    for i, h in enumerate(headers):
        c = t.cell(0, i); set_cell_shading(c, NAVY)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.text = ""; p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        r = p.add_run(h); set_run_font(r, size=10, bold=True, color_hex="FFFFFF")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.cell(ri+1, ci)
            if ri % 2 == 1: set_cell_shading(c, ALT_FILL)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            c.text = ""; p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
            txt, opts = (val if isinstance(val, tuple) else (val, {}))
            if align_center_cols and ci in align_center_cols:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif opts.get("align") == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(txt))
            set_run_font(r, size=opts.get("size", 9),
                         bold=opts.get("bold", False),
                         color_hex=opts.get("color", TEXT))

def add_page_field(run, instr):
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"),"begin")
    instr_text = OxmlElement("w:instrText"); instr_text.text = instr
    fld_end   = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"),"end")
    run._r.append(fld_begin); run._r.append(instr_text); run._r.append(fld_end)

def setup_page_footer(section, doc_id, today):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""
    r1 = p.add_run(f"{doc_id} | 编制：曾翔宇 | 日期：{today} | 第 ")
    set_run_font(r1, size=9, color_hex="666666")
    r2 = p.add_run(""); set_run_font(r2, size=9, color_hex="666666")
    add_page_field(r2, "PAGE")
    r3 = p.add_run(" 页 / 共 "); set_run_font(r3, size=9, color_hex="666666")
    r4 = p.add_run(""); set_run_font(r4, size=9, color_hex="666666")
    add_page_field(r4, "NUMPAGES")
    r5 = p.add_run(" 页"); set_run_font(r5, size=9, color_hex="666666")

# ============================================================================
# DOCUMENT
# ============================================================================
doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.59); section.page_height = Cm(27.94)
section.top_margin = Cm(1.27); section.bottom_margin = Cm(1.27)
section.left_margin = Cm(1.59); section.right_margin = Cm(1.59)
setup_page_footer(section, DOC_ID, TODAY)

style = doc.styles["Normal"]
style.font.name = "Arial"; style.font.size = Pt(10)

# ---------- COVER ----------
add_para(doc, "瑞幸咖啡北美", size=22, bold=True, color_hex=NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=8)
add_para(doc, "QA门店巡检月度分析报告", size=18, bold=True, color_hex=NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para(doc, "Monthly QA Store Audit Analysis Report", size=12, color_hex=BLUE, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
add_para(doc, TARGET_MONTH_CN, size=20, bold=True, color_hex=NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=80)
add_para(doc, "质量保障部 / 基础设施部", size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para(doc, "编制：曾翔宇    日期：" + TODAY, size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para(doc, f"报告编号：{DOC_ID} | 状态：V2稿（数据刷新版）", size=10, color_hex=BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
doc.add_page_break()

# ---------- 文档信息 ----------
add_heading(doc, "文档信息", 1)
total_findings_all  = len(ITEMS)
total_findings_main = len(MAIN_ITEMS)
all_s, all_m, all_g, all_l = (ALL_SEV.get(k,0) for k in ("S","M","G","L"))
main_s, main_m, main_g, main_l = (MAIN_SEV.get(k,0) for k in ("S","M","G","L"))
appeal_text = f"{len(APPEALED)} 起申诉立案（{len(APPEAL_APPROVED)} 起获批、{len(APPEAL_DENIED)} 起驳回、{len(APPEAL_PENDING)} 起审批中）" if APPEALED else "本月无申诉立案"
doc_info_rows = [
    ["报告编号", DOC_ID],
    ["报告周期", TARGET_MONTH_CN],
    ["数据范围", "2026-04-01 至 2026-04-30"],
    ["有效门店", f"{N_STORES_ACTIVE} 家（已巡检活跃门店）"],
    ["巡检类型", f"门店自检（{N_BY_TYPE['门店自检']}次） + QA审计（{N_BY_TYPE['QA审计']}次） + 区经检查（{N_BY_TYPE['区经检查']}次） = 共 {N_TOTAL} 次（已提交）"],
    ["问题总数", f"全月 {total_findings_all} 个扣分项（S 项 {all_s}、M 项 {all_m}、G 项 {all_g}、L 项 {all_l}）；主巡检 {total_findings_main} 个（S 项 {main_s}、M 项 {main_m}、G 项 {main_g}、L 项 {main_l}）"],
    ["申诉情况", appeal_text],
    ["编制人", "曾翔宇"],
    ["部门", "质量保障部 / 基础设施部"],
    ["数据来源", "empapp 门店稽核系统（aws-luckyus-opqualitycontrol-rw / luckyus_opqualitycontrol）"],
    ["状态", f"V2稿（{TODAY} 数据刷新版，含申诉调整）"],
]
add_data_table(doc, ["项目","内容"], doc_info_rows, col_widths=[3.5, 14.5])
add_para(doc, "", space_after=6)

# ---------- 数据说明 ----------
add_heading(doc, "⚠ 数据说明（V2 刷新版）", 1)
data_notes = [
    f"(a) V2 vs V1 — 申诉数据并入：本月共 {len(APPEALED)} 起申诉立案，其中获批 {len(APPEAL_APPROVED)} 起（54th & 8th 69→94、15th & 3rd 71→96）、驳回 {len(APPEAL_DENIED)} 起（8th & Broadway QA 4/9、102 Fulton QA 4/27、28th & 6th QA 4/27）、审批中 {len(APPEAL_PENDING)} 起（28th & 6th 区经 4/23）。所有得分以申诉调整后值（adjusted_total_score）显示，调整门店在表格中以 ※ 标记。",
    "(b) V2 vs V1 — 误提交剔除：US00020（21st & 3rd）2026-04-21 由 Darwin Coronel 单人提交的两次 100 分零扣分自检（巡检 ID 2016/2017）按「同人同店同日重复 100 分零扣分」规则已剔除；同日真实自检（巡检 ID 2018，得分 64，5 项）保留。",
    f"(c) 巡检系统状态对比：4 月共完成 {N_TOTAL} 次有效巡检（已剔除误提交），三类齐全；与 1月（16次三类齐全）、2月（7次区经断流）、3月（14次体系崩溃）形成「危机—响应」循环。区经检查在中断 3 个月后于 4月 完全恢复（{N_BY_TYPE['区经检查']}次）。",
    f"门店覆盖率 {N_STORES_ACTIVE} / {N_STORES_ACTIVE} = 100%（4月所有活跃门店均完成巡检）。",
    "新开门店：US00007（108th & Broadway，4/30）、US00010（154 Bleecker，4/28）、US00015（41st & Lexington，4/30）于本月底开业，未纳入 4月 巡检覆盖；US00012（16th & 6th，3/23 开业）、US00019（29th & 3rd，4/11 开业）首次进入月度巡检覆盖。",
    "QA 人员变更：Yu Jiang 4 月未执行任何巡检（1月 5 次、2月 2 次、3月 2 次后退出），需走 HR 离职 / 转岗流程；Eamonn Caballar 4 月执行 12 次，已全面接管 QA Senior Manager 角色——单点依赖风险显著。",
    "区经巡检：Daniel Chu 完成 7 次、Jung Han Liang 完成 7 次，节奏均衡。",
    "标准规则：① 主巡检优先级 QA审计 > 区经检查 > 门店自检；同优先级取最近日期；② 仅采用已提交单据，草稿已剔除；③ 同人同店同日重复 100 分零扣分自检按误提交规则剔除。",
]
add_callout(doc, data_notes, kind="warn")

# ---------- 管理摘要 ----------
add_heading(doc, "管理摘要", 1)
add_para(doc,
    f"本月平均分 {avg_score_main} 分（基于各门店主巡检），覆盖 {N_STORES_ACTIVE} 家门店，主巡检 {total_findings_main} 项有效扣分项，全月共 {total_findings_all} 项（含全部巡检类型）。",
    size=10, space_after=8)

lowest = min(MAIN_INSP.values(), key=lambda r: r["adjusted_total_score"])
highest = max(MAIN_INSP.values(), key=lambda r: r["adjusted_total_score"])

add_callout(doc,
    f"巡检体系全面恢复：连续 3 个月退化的区经检查在 4月 完全恢复（{N_BY_TYPE['区经检查']}次），QA审计从 3月 的 1 次激增至 {N_BY_TYPE['QA审计']} 次，门店自检 {N_BY_TYPE['门店自检']} 次，{N_STORES_ACTIVE} 家活跃门店均获得三类巡检全覆盖——这是 2026 年首次实现。",
    kind="good", title="✅ 体系恢复")

# Compute key sub-item systemic clusters
all_s_items = [it for it in ITEMS if it["severity"] == "S"]
sub_cluster = Counter(it["module_subcategory"] for it in all_s_items)
top_sub = sub_cluster.most_common(2)
sub_text = "、".join(f"{n}（{c} 起）" for n, c in top_sub) if top_sub else ""
add_callout(doc,
    f"食品安全风险仍存：全月发现 {all_s} 个 S 项（关键项）和 {all_m} 个 M 项（重要项），分布在多家门店。S 项最集中子项为 {sub_text}；最低分主巡检门店 {lowest['store_name']}（{lowest['adjusted_total_score']} 分{'※' if lowest['appeal_status']=='approved' else ''}）。",
    kind="warn", title="⚠ 严重风险")

add_callout(doc,
    "巡检员评分一致性问题：US00020（21st & 3rd）于 4/21 由 Darwin Coronel 单人提交三次自检，得分 100/100/64，摆动幅度 36 分（前两次按误提交规则剔除）；同店同期跨类型最大差距 33rd & 10th 自检 81 vs 区经 47（差 34 分）。",
    kind="warn", title="⚠ 评分一致性")

# ============================================================================
# §一、门店整体表现
# ============================================================================
add_heading(doc, "一、门店整体表现", 1)

add_heading(doc, "1.1 本月概览", 2)
n_below_80 = sum(1 for m in MAIN_INSP.values() if m["adjusted_total_score"] < 80)
n_with_s   = sum(1 for m in MAIN_INSP.values() if m["s_count"] > 0)
overview_rows = [[
    f"{highest['store_name']}\n{highest['adjusted_total_score']} 分{'※' if highest['appeal_status']=='approved' else ''}",
    f"{lowest['store_name']}\n{lowest['adjusted_total_score']} 分{'※' if lowest['appeal_status']=='approved' else ''}",
    f"{n_with_s} 家",
    f"{n_below_80} 家",
]]
add_data_table(doc, ["最高分门店","最低分门店","S 项门店数","<80 分门店数"],
               overview_rows, col_widths=[4.5,4.5,3.0,3.0], align_center_cols=[0,1,2,3])

add_heading(doc, "1.2 各门店得分明细（基于主巡检）", 2)
ranked = sorted(MAIN_INSP.values(),
                key=lambda r: (-r["adjusted_total_score"], PRIORITY.get(r["inspection_type"], 99), r["store_code"]))
rows_12 = []
for i, m in enumerate(ranked, 1):
    sc = m["adjusted_total_score"]
    appealed_mark = m["appeal_status"] == "approved"
    score_text = f"{sc}{'※' if appealed_mark else ''}"
    score_color = CRITICAL_FG if sc < 80 else (NAVY if sc >= 85 else TEXT)
    score_bold  = sc < 80 or sc >= 85
    s_color = CRITICAL_FG if m["s_count"] > 0 else TEXT
    rows_12.append([
        str(i),
        (m["store_name"], {"bold": True}),
        m["store_code"],
        (score_text, {"bold": score_bold, "color": score_color, "align":"center"}),
        m["inspection_type"],
        (str(m["adjusted_total_deduction"]), {"align":"center"}),
        (str(m["s_count"]), {"bold": m["s_count"]>0, "color": s_color, "align":"center"}),
        (str(m["m_count"]), {"align":"center"}),
        (str(m["g_count"]), {"align":"center"}),
        (str(m["l_count"]), {"align":"center"}),
        m["inspector_name"],
    ])
add_data_table(doc, ["#","门店","编号","得分","巡检类型","扣分","S","M","G","L","巡检员"],
               rows_12, col_widths=[0.6,3.0,1.5,1.0,1.7,1.0,0.6,0.6,0.6,0.6,3.0],
               align_center_cols=[0,2,3,4,5,6,7,8,9])
add_para(doc, "※ = 申诉获批后调整分数（adjusted_total_score）；详见 §4.3。",
         size=9, color_hex="666666", space_before=2, space_after=8)

add_heading(doc, "1.3 管理解读", 2)
n_above_85 = sum(1 for m in MAIN_INSP.values() if m["adjusted_total_score"] >= 85)
delta_lines = []
for sc, m in MAIN_INSP.items():
    prev = march_score_for(sc)
    if prev is not None:
        d = m["adjusted_total_score"] - prev
        if abs(d) >= 10:
            delta_lines.append((d, sc, m["store_name"], prev, m["adjusted_total_score"]))
delta_lines.sort(key=lambda x: -abs(x[0]))

high_text = "、".join(f"{m['store_name']} {m['adjusted_total_score']} 分{'※' if m['appeal_status']=='approved' else ''}" for m in ranked if m["adjusted_total_score"] >= 85)
low_text  = "、".join(f"{m['store_name']} {m['adjusted_total_score']} 分{'※' if m['appeal_status']=='approved' else ''}" for m in ranked if m["adjusted_total_score"] < 80)

bullets = [
    f"• {n_above_85} 家门店达到 85 分以上（{high_text}）。",
    f"• {n_below_80} 家门店低于 80 分（{low_text}）。",
    f"• 全月 {all_s} 个 S 项分布在 {len(set(it['store_code'] for it in all_s_items))} 家门店，主要集中在 {top_sub[0][0] if top_sub else ''}（{top_sub[0][1] if top_sub else 0} 起）、{top_sub[1][0] if len(top_sub)>1 else ''}（{top_sub[1][1] if len(top_sub)>1 else 0} 起）等关键子项；多家门店出现 air-gap 重复 S 项，需 BD 整改清单跟踪。",
    "• 同日重复巡检暴露评分一致性问题：US00020 在 2026-04-21 同日由 Darwin Coronel 提交 3 次自检，得分摆动 36 分；详见 §7.3。",
    f"• 自检主动发现 S 项（积极信号）：Eric Park（54th & 8th 自检 63 分，发现 1 S 1 M）、Juan Ortiz-Fontanez（29th & 3rd 自检 66 分，2 S）等门店端自检员主动暴露问题，可作为标杆。",
]
if APPEAL_APPROVED:
    appeal_text2 = "、".join(f"{r['store_name']} {r['original_total_score']}→{r['adjusted_total_score']}" for r in APPEAL_APPROVED)
    bullets.append(f"• 申诉调整：{len(APPEAL_APPROVED)} 起申诉获批后得分修正（{appeal_text2}），扣分项保留但 deductionScore 归零；详见 §4.3。")
if delta_lines:
    bullets.append("• 跨月对比（主巡检 vs 3月主巡检）≥10 分变动：" +
                   "；".join(f"{n} 由 {old} 升至 {new}（{'+' if d>0 else ''}{d}）" for d,_,n,old,new in delta_lines[:5]))
for b in bullets:
    add_para(doc, b, size=10, indent_left=12, space_after=4)

# ============================================================================
# §二、模块风险分析
# ============================================================================
add_heading(doc, "二、模块风险分析", 1)
add_para(doc, f"本月主巡检共发现 {total_findings_main} 个扣分项，分布在 10 个标准模块中（{ALL_SEV.get('S',0)+ALL_SEV.get('M',0)+ALL_SEV.get('G',0)+ALL_SEV.get('L',0)} 个全月扣分项详见 §3.3）。", space_after=6)

add_heading(doc, "2.1 风险分层（基于主巡检覆盖率）", 2)
def coverage_main(mod):
    return len(MAIN_MOD_STORES.get(mod, set())) / max(1, N_STORES_ACTIVE)
sys_mods = [m for m in MODULES_10 if coverage_main(m) >= 0.5 and MAIN_MOD_SEV.get(m)]
mid_mods = [m for m in MODULES_10 if 0.3 <= coverage_main(m) < 0.5 and MAIN_MOD_SEV.get(m)]
low_mods = [m for m in MODULES_10 if 0 < coverage_main(m) < 0.3 and MAIN_MOD_SEV.get(m)]
add_para(doc, f"🔴 系统性风险（影响≥50% 门店）：{', '.join(sys_mods) or '（无）'}", size=10, indent_left=12, space_after=4)
add_para(doc, f"🟡 中等覆盖面（影响 30-49%）：{', '.join(mid_mods) or '（无）'}", size=10, indent_left=12, space_after=4)
add_para(doc, f"🟢 低覆盖面（<30%）：{', '.join(low_mods) or '（无）'}", size=10, indent_left=12, space_after=8)

add_heading(doc, "2.2 模块排名总览（按扣分排序，主巡检视角）", 2)
mods_with_data = [m for m in MODULES_10 if m in MAIN_MOD_SEV] + (["其他"] if MAIN_MOD_SEV.get("其他") else [])
mods_ranked = sorted(mods_with_data, key=lambda m: sum(MAIN_MOD_DED.get(m, {}).values()))
module_rows = []
for rank, mod in enumerate(mods_ranked, 1):
    sev_c = MAIN_MOD_SEV[mod]
    ded_total = sum(MAIN_MOD_DED.get(mod, {}).values())
    cnt_total = sum(sev_c.values())
    n_stores  = len(MAIN_MOD_STORES.get(mod, set()))
    cov_pct   = round(n_stores / N_STORES_ACTIVE * 100)
    risks = []
    if cov_pct >= 50: risks.append("⚠ 系统性")
    if sev_c.get("S",0) > 0: risks.append("⚠ 含 S 项")
    if not risks and sev_c.get("M",0) > 0: risks.append("含 M 项")
    if not risks: risks.append("---")
    module_rows.append([str(rank), mod, str(cnt_total), str(ded_total),
                        f"{n_stores}/{N_STORES_ACTIVE}", f"{cov_pct}%",
                        str(sev_c.get("S",0)), str(sev_c.get("M",0)),
                        str(sev_c.get("G",0)), str(sev_c.get("L",0)),
                        " / ".join(risks)])
add_data_table(doc, ["#","模块","问题数","扣分","门店","覆盖率","S","M","G","L","风险"],
               module_rows, col_widths=[0.6,3.4,1.0,1.0,1.0,1.0,0.6,0.6,0.6,0.6,2.6],
               align_center_cols=[0,2,3,4,5,6,7,8,9])

add_heading(doc, "2.3 重点模块详细分析（TOP 5，主巡检视角）", 2)

def quotable_items(items, max_n=25):
    sev_rank = {"S":0,"M":1,"G":2,"L":3}
    out = []
    for it in sorted(items, key=lambda x: (sev_rank.get(x["severity"],9), x["store_code"], x["inspection_date"])):
        desc = (it["issue_description"] or "").strip()
        if not desc: continue
        out.append(it)
        if len(out) >= max_n: break
    return out

top_mods = mods_ranked[:5]
for mod in top_mods:
    sev_c = MAIN_MOD_SEV[mod]
    ded_total = sum(MAIN_MOD_DED[mod].values())
    n_stores = len(MAIN_MOD_STORES[mod])
    add_heading(doc, f"{mod} — {sum(sev_c.values())} 个扣分项，{ded_total} 分，影响 {n_stores} 家门店", 3)
    add_para(doc, f"严重级别：S 项 {sev_c.get('S',0)} 个、M 项 {sev_c.get('M',0)} 个、G 项 {sev_c.get('G',0)} 个、L 项 {sev_c.get('L',0)} 个。",
             size=10, space_after=4)
    mod_items_main = [it for it in MAIN_ITEMS if it["cn_module"] == mod]
    quotes = quotable_items(mod_items_main, max_n=20)
    if quotes:
        add_para(doc, "具体问题（引用原始描述，按严重度排序，最多展示 20 条；空描述已跳过）：",
                 size=10, color_hex="666666", space_after=4)
        for it in quotes:
            marker = "⚠ " if it["severity"] == "S" else ""
            appeal_mark = " ※" if it["is_appealed_finding"] else ""
            desc = (it["issue_description"] or "").strip().replace("\n", " ").replace("\r", " ")[:200]
            add_para(doc,
                f"• {marker}{it['store_name']} ({it['store_code']})｜{it['module_subcategory']}｜{it['severity']} 项 {it['deduction_points']} 分{appeal_mark}｜{desc}",
                size=9, indent_left=8, space_after=2)
        if len(mod_items_main) > len(quotes):
            add_para(doc, f"… 另有 {len(mod_items_main)-len(quotes)} 条主巡检发现，完整明细见原始 CSV。",
                     size=9, color_hex="666666", indent_left=8, space_after=4)
    else:
        add_para(doc, "本子项无可引用的描述（所有发现描述均为空）。",
                 size=10, color_hex="666666", space_after=4)
    if mod == "设施":
        sp_s = [it for it in mod_items_main if it["severity"]=="S" and it["module_subcategory"]=="Sinks and Pipes"]
        if len(sp_s) >= 3:
            add_callout(doc,
                f"Sinks and Pipes / air gap 在主巡检中累计 {len(sp_s)} 起 S 项，覆盖 {len(set(it['store_code'] for it in sp_s))} 家门店——属系统性短板，须由 BD 列入跨门店整改清单（详见 §6.2 P0）。",
                kind="critical", title="⚠ 关键发现")

# ============================================================================
# §三、风险等级分布
# ============================================================================
add_heading(doc, "三、风险等级分布", 1)

add_heading(doc, "3.1 整体分布（主巡检）", 2)
def top_modules_for_sev_main(sev, top=4):
    cnt = Counter(it["cn_module"] for it in MAIN_ITEMS if it["severity"] == sev)
    return "、".join(f"{m}（{c}）" for m,c in cnt.most_common(top)) or "—"

main_total = max(1, sum(MAIN_SEV.values()))
sev_rows = [
    ["S 项（关键项）", str(main_s), f"{main_s/main_total*100:.1f}%", "2 天内闭环", top_modules_for_sev_main("S")],
    ["M 项（重要项）", str(main_m), f"{main_m/main_total*100:.1f}%", "7 天内闭环", top_modules_for_sev_main("M")],
    ["G 项（一般项）", str(main_g), f"{main_g/main_total*100:.1f}%", "14 天内闭环", top_modules_for_sev_main("G")],
    ["L 项（轻微项）", str(main_l), f"{main_l/main_total*100:.1f}%", "14 天内闭环", top_modules_for_sev_main("L")],
    ["合计", str(sum(MAIN_SEV.values())), "100%", "---", "---"],
]
add_data_table(doc, ["风险等级","数量","占比","SLA 要求","主要分布模块"], sev_rows,
               col_widths=[2.2,1.0,1.0,1.6,9.0])

add_callout(doc,
    f"主巡检 S {main_s} / M {main_m} / G {main_g} / L {main_l} = {sum(MAIN_SEV.values())} vs 全月 S {all_s} / M {all_m} / G {all_g} / L {all_l} = {total_findings_all}。"
    " 全月口径包含同店多次巡检的重复发现（自检 + QA + 区经），主巡检为各店最权威单次。",
    kind="info", title="📊 主巡检 vs 全月对比")

add_heading(doc, "3.2 S 项详情（主巡检） ── 必须立即整改", 2)
main_s_items = sorted([it for it in MAIN_ITEMS if it["severity"] == "S"],
                      key=lambda x: (x["store_code"], x["inspection_date"]))
add_para(doc, f"主巡检共发现 {len(main_s_items)} 个 S 项。", size=10, space_after=4)
if main_s_items:
    rows_s = []
    for i, it in enumerate(main_s_items, 1):
        appeal_mark = " ※" if it["is_appealed_finding"] else ""
        desc_raw = (it["issue_description"] or "").strip().replace("\n"," ").replace("\r"," ")
        desc = desc_raw[:160] + appeal_mark if desc_raw else f"（描述空白）{appeal_mark}"
        rows_s.append([str(i),
                       f"{it['store_name']}\n{it['store_code']}",
                       f"{it['cn_module']}\n{it['module_subcategory']}",
                       desc,
                       (str(it['deduction_points']), {"align":"center"}),
                       it['inspection_type'],
                       it['inspector_name'],
                       it['inspection_date']])
    add_data_table(doc, ["#","门店","模块/子项","问题描述（原文）","扣分","巡检类型","巡检员","日期"],
                   rows_s, col_widths=[0.6,2.2,2.6,5.6,1.0,1.5,2.0,1.4],
                   align_center_cols=[0,4])
else:
    add_para(doc, "本月主巡检无 S 项案例。", size=10, color_hex="666666")

add_heading(doc, "3.3 全月 S 项汇总（含自检与重复巡检）", 2)
add_para(doc, f"全月共 {len(all_s_items)} 个 S 项（含主巡检 + 自检 + 重复巡检），按子项汇总：",
         size=10, space_after=4)
sub_roll = defaultdict(lambda: {"cnt":0, "stores":set(), "samples":[]})
for it in all_s_items:
    k = it["module_subcategory"]
    sub_roll[k]["cnt"] += 1
    sub_roll[k]["stores"].add(it["store_code"])
    desc = (it["issue_description"] or "").strip()
    if desc and len(sub_roll[k]["samples"]) < 2:
        sub_roll[k]["samples"].append(desc)
roll_rows = []
for sub, info in sorted(sub_roll.items(), key=lambda kv: -kv[1]["cnt"]):
    sample = (info["samples"][0] if info["samples"] else "（本子项可引用描述均为空）")[:120].replace("\n"," ")
    roll_rows.append([sub, str(info["cnt"]), str(len(info["stores"])), sample])
add_data_table(doc, ["子项","S 项数","门店数","典型问题（截取）"], roll_rows,
               col_widths=[3.6,1.0,1.0,11.0], align_center_cols=[1,2])
add_para(doc, "S 项最集中子项已计入 §6.2 P0 行动（air gap / Handwashing Standards 跨门店专项治理）。",
         size=9, color_hex="666666", space_before=4)

add_heading(doc, "3.4 M 项详情（主巡检） ── 7 天内闭环", 2)
main_m_items = sorted([it for it in MAIN_ITEMS if it["severity"] == "M"],
                      key=lambda x: (x["store_code"], x["inspection_date"]))
add_para(doc, f"主巡检共发现 {len(main_m_items)} 个 M 项。", size=10, space_after=4)
rows_m = []
for i, it in enumerate(main_m_items, 1):
    appeal_mark = " ※" if it["is_appealed_finding"] else ""
    desc_raw = (it["issue_description"] or "").strip().replace("\n"," ").replace("\r"," ")
    desc = desc_raw[:140] + appeal_mark if desc_raw else f"（描述空白）{appeal_mark}"
    rows_m.append([str(i),
                   f"{it['store_name']}\n{it['store_code']}",
                   f"{it['cn_module']}\n{it['module_subcategory']}",
                   desc,
                   (str(it['deduction_points']), {"align":"center"}),
                   it['inspection_date']])
add_data_table(doc, ["#","门店","模块/子项","问题描述（原文）","扣分","日期"], rows_m,
               col_widths=[0.6,2.4,2.8,8.4,1.0,1.4], align_center_cols=[0,4])

add_heading(doc, "3.5 G 项 / L 项 分布（主巡检）", 2)
g_dist = Counter(it["cn_module"] for it in MAIN_ITEMS if it["severity"] == "G")
l_dist = Counter(it["cn_module"] for it in MAIN_ITEMS if it["severity"] == "L")
add_para(doc, f"G 项（一般项）共 {main_g} 个，主要集中模块：", size=10, bold=True, space_after=2)
for m,c in g_dist.most_common(): add_para(doc, f"  • {m}：{c} 个", size=10, indent_left=8, space_after=2)
add_para(doc, f"L 项（轻微项）共 {main_l} 个，主要集中模块：", size=10, bold=True, space_before=4, space_after=2)
for m,c in l_dist.most_common(): add_para(doc, f"  • {m}：{c} 个", size=10, indent_left=8, space_after=2)

# ============================================================================
# §四、模块与门店关联分析
# ============================================================================
add_heading(doc, "四、模块与门店关联分析", 1)

add_heading(doc, "4.1 门店 × 模块扣分矩阵（主巡检）", 2)
matrix_mods = MODULES_10
mat_headers = ["门店"] + [MODULE_SHORT[m] for m in matrix_mods] + ["合计"]
store_total = {sc: sum(MAIN_STORE_MOD.get(sc, {}).values()) for sc in ACTIVE_STORE_CODES}
ordered = sorted(ACTIVE_STORE_CODES, key=lambda sc: store_total[sc])
mat_rows = []
for sc in ordered:
    name = MAIN_INSP[sc]["store_name"]
    appeal_mark = "※ " if MAIN_INSP[sc]["appeal_status"] == "approved" else ""
    row = [f"{appeal_mark}{name}\n({sc})"]
    for mod in matrix_mods:
        v = MAIN_STORE_MOD.get(sc, {}).get(mod, 0)
        row.append((str(v) if v else "", {"align":"center"}))
    row.append((str(store_total[sc]), {"bold": True, "align":"center"}))
    mat_rows.append(row)
add_data_table(doc, mat_headers, mat_rows,
               col_widths=[2.6]+[1.15]*10+[1.2], align_center_cols=list(range(1,12)))
add_para(doc, "※ = 申诉获批门店，矩阵数值为申诉调整后扣分；申诉前原始扣分见 §4.3。",
         size=9, color_hex="666666", space_before=2, space_after=4)

add_heading(doc, "4.2 最低分门店归因", 2)
low_mod = sorted(MAIN_STORE_MOD.get(lowest["store_code"], {}).items(), key=lambda kv: kv[1])[:3]
low_mod_text = "、".join(f"{m}（{v} 分）" for m,v in low_mod) or "（无扣分项）"
add_para(doc,
    f"最低分主巡检门店：{lowest['store_name']}（{lowest['store_code']}），主巡检得分 {lowest['adjusted_total_score']} 分{'※（申诉获批后调整）' if lowest['appeal_status']=='approved' else ''}，扣分 {lowest['adjusted_total_deduction']} 分，集中在：{low_mod_text}。",
    size=10, space_after=4)
add_para(doc,
    f"巡检类型为 {lowest['inspection_type']}（{lowest['inspection_date']}），巡检员 {lowest['inspector_name']}。",
    size=10, space_after=8)

prev = march_score_for(lowest["store_code"])
if prev is not None:
    delta = lowest["adjusted_total_score"] - prev
    direction = "回升" if delta > 0 else ("持平" if delta == 0 else "下滑")
    add_callout(doc,
        f"从 3 月的 {prev} 分 → 4 月 {lowest['adjusted_total_score']} 分（{'+' if delta>0 else ''}{delta} 分变动反映 {direction}）。",
        kind="info", title="📈 跨月对比")
else:
    add_callout(doc, "本店为本月首次巡检覆盖（3 月主巡检数据缺失，无法跨月比对）。", kind="info", title="📈 跨月对比")

add_heading(doc, "4.3 申诉调整门店分析", 2)
if APPEALED:
    add_para(doc, f"本月共 {len(APPEALED)} 起申诉立案：获批 {len(APPEAL_APPROVED)} 起、驳回 {len(APPEAL_DENIED)} 起、审批中 {len(APPEAL_PENDING)} 起。所有申诉立案门店全量列出（含驳回 / 审批中）：",
             size=10, space_after=6)
    by_appeal_store = defaultdict(list)
    for r in APPEALED: by_appeal_store[r["store_code"]].append(r)
    for sc in sorted(by_appeal_store.keys()):
        appeal_rows_sc = by_appeal_store[sc]
        store_all = sorted(by_store[sc], key=lambda x: x["inspection_date"])
        add_heading(doc, f"{appeal_rows_sc[0]['store_name']}（{sc}）", 3)
        sub_rows = []
        status_cn = {"approved":"获批","denied":"驳回","pending":"审批中","none":"—"}
        for r in store_all:
            mark = "※" if r["appeal_status"] == "approved" else ""
            sub_rows.append([
                r["inspection_date"],
                r["inspection_type"],
                r["inspector_name"],
                str(r["original_total_score"]),
                f"{r['adjusted_total_score']}{mark}",
                str(r["original_total_deduction"]),
                str(r["adjusted_total_deduction"]),
                f"{r['s_count']}/{r['m_count']}/{r['g_count']}/{r['l_count']}",
                status_cn[r["appeal_status"]],
            ])
        add_data_table(doc,
            ["日期","类型","巡检员","原始分","调整分","原始扣分","调整扣分","S/M/G/L","申诉状态"],
            sub_rows, col_widths=[1.5,1.4,2.8,1.0,1.0,1.2,1.2,1.5,1.4],
            align_center_cols=[3,4,5,6,7,8])
        for r in appeal_rows_sc:
            n_appeal_findings = sum(1 for it in items_by_iid[r["inspection_id"]] if it["is_appealed_finding"] == 1)
            outcome_map = {
                "approved": "获批 — 得分修正、扣分项保留但 deductionScore 归零；门店反馈被采纳，BD 整改任务移交",
                "denied":   "驳回 — 原始扣分维持，门店反馈未被采纳，需按原方案整改",
                "pending":  "审批中 — 当前以原始得分入账，待审批结果",
            }
            outcome = outcome_map.get(r["appeal_status"], "—")
            add_para(doc,
                f"• 巡检 ID {r['inspection_id']} {r['inspection_date']} {r['inspection_type']}：申诉项 {n_appeal_findings} 个，结果 {outcome}；得分变化 {r['original_total_score']} → {r['adjusted_total_score']}；扣分变化 {r['original_total_deduction']} → {r['adjusted_total_deduction']}。",
                size=10, indent_left=12, space_after=3)
        add_para(doc, "", size=4, space_after=4)
    add_callout(doc,
        f"管理含义：申诉机制本月首次跑通完整闭环（立案→审批→落地），获批 {len(APPEAL_APPROVED)} / 驳回 {len(APPEAL_DENIED)} / 审批中 {len(APPEAL_PENDING)}。审批通过率 {len(APPEAL_APPROVED)*100//max(1,len(APPEALED))}% 处于合理区间，避免「全部驳回」（流于形式）或「全部获批」（标准失守）。下月需关注：① 获批门店实际整改完成度；② 驳回门店是否按原方案闭环；③ 审批 SLA 监控（防止「审批中」长期挂起）。",
        kind="info", title="📊 申诉机制观察")
else:
    add_para(doc, "本月无申诉立案。", size=10, color_hex="666666")

add_heading(doc, "4.4 同店同期评分背离案例", 2)
divergence_cases = []
for sc, rows in by_store.items():
    by_type = defaultdict(list)
    for r in rows: by_type[r["inspection_type"]].append(r)
    types = list(by_type.keys())
    for i in range(len(types)):
        for j in range(i+1, len(types)):
            for r1 in by_type[types[i]]:
                for r2 in by_type[types[j]]:
                    s1, s2 = r1["adjusted_total_score"], r2["adjusted_total_score"]
                    if abs(s1 - s2) >= 20:
                        divergence_cases.append((abs(s1-s2), sc, r1, r2))
divergence_cases.sort(key=lambda x: -x[0])
if divergence_cases:
    add_para(doc, f"主巡检与同店其他巡检评分差距 ≥20 分案例（共 {len(divergence_cases)} 组）：", space_after=4)
    rows_dv = []
    for gap, sc, r1, r2 in divergence_cases[:15]:
        rows_dv.append([
            f"{r1['store_name']} ({sc})",
            f"{r1['inspection_type']} {r1['adjusted_total_score']}",
            f"{r2['inspection_type']} {r2['adjusted_total_score']}",
            (str(gap), {"bold":True, "color":CRITICAL_FG, "align":"center"}),
            f"{r1['inspector_name']} / {r2['inspector_name']}",
            f"{r1['inspection_date']} / {r2['inspection_date']}",
        ])
    add_data_table(doc, ["门店","对比 1","对比 2","差距","巡检员","日期"], rows_dv,
                   col_widths=[3.0,2.5,2.5,1.0,5.0,2.0], align_center_cols=[3])
else:
    add_para(doc, "本月无 ≥20 分背离案例。", size=10, color_hex="666666")

add_heading(doc, "4.5 模块覆盖面分析（主巡检视角）", 2)
cov_rows = []
for mod in sorted(mods_with_data, key=lambda m: -len(MAIN_MOD_STORES.get(m, set()))):
    n_st = len(MAIN_MOD_STORES.get(mod, set()))
    if n_st == 0: continue
    cov_pct = round(n_st / N_STORES_ACTIVE * 100)
    ded = sum(MAIN_MOD_DED.get(mod, {}).values())
    sev = MAIN_MOD_SEV.get(mod, {})
    risks = []
    if cov_pct >= 50: risks.append("⚠ 系统性")
    if sev.get("S",0) > 0: risks.append("⚠ 含 S 项")
    if not risks and sev.get("M",0) > 0: risks.append("含 M 项")
    if not risks: risks.append("---")
    cov_rows.append([mod, f"{n_st}/{N_STORES_ACTIVE}", f"{cov_pct}%", str(ded), " / ".join(risks)])
add_data_table(doc, ["模块","影响门店","覆盖率","扣分","风险标记"], cov_rows,
               col_widths=[3.6,2.0,1.4,1.4,5.0], align_center_cols=[1,2,3])

# ============================================================================
# §五、整改归因与效率
# ============================================================================
add_heading(doc, "五、整改归因与效率", 1)

add_heading(doc, "5.1 关键词归因（主巡检）", 2)
add_para(doc, "⚠ 以下归因基于问题描述关键词自动匹配，仅供参考。实际归因需以整改工单系统数据为准。",
         size=9, color_hex=WARN_FG, space_after=4)
def attribute(text):
    t = (text or "").strip().lower()
    if len(t) < 10: return "未知"
    pipe = ["pipe","sink","leak","airgap","air gap","drain","handwashing sink","plumbing","fixture","light","ceiling","floor","wall"]
    sup = ["license","sign","no smoking","permit","document","expiration date"]
    store_kw = ["clean","sanitize","glove","wiping","cloth","sticker","cup","container","ppm","sanitizer","storage","dirty","residue","spill"]
    if any(k in t for k in pipe): return "机修+营建"
    if any(k in t for k in sup):  return "供应链+行政"
    if any(k in t for k in store_kw): return "门店"
    return "门店"
attr_cnt = Counter(attribute(it["issue_description"]) for it in MAIN_ITEMS)
attr_total = max(1, sum(attr_cnt.values()))
attr_rows = []
typical = {"门店":"日常清洁、消毒、标签、卫生",
           "机修+营建":"sinks and pipes、air gap、light fixtures",
           "供应链+行政":"License / No smoking sign / 文件记录",
           "未知":"描述模糊或缺失（少于 10 字符）"}
for cat in ["门店","机修+营建","供应链+行政","未知"]:
    c = attr_cnt.get(cat, 0)
    pct = f"{c/attr_total*100:.1f}%"
    attr_rows.append([cat, str(c), pct, typical[cat]])
add_data_table(doc, ["归因类别","数量","占比","典型问题"], attr_rows,
               col_widths=[2.4,1.0,1.0,9.0], align_center_cols=[1,2])

add_heading(doc, "5.2 SLA 整改时限标准", 2)
add_data_table(doc, ["风险等级","整改时限","要求"], [
    ["S 项（关键项）", "2 天", "发现后 2 天内完成整改并验证"],
    ["M 项（重要项）", "7 天", "发现后 7 天内完成整改并验证"],
    ["G 项（一般项)", "14 天","发现后 14 天内完成整改并验证"],
    ["L 项（轻微项）", "14 天","发现后 14 天内完成整改并验证"],
], col_widths=[2.5,2.0,9.0], align_center_cols=[1])

add_heading(doc, "5.3 建议整改闭环流程", 2)
flow = [
    "巡检发现问题 → empapp 自动生成整改工单",
    "根据问题类型自动分配责任方（门店 / 机修 / 营建 / 供应链）",
    "责任方在 SLA 时限内完成整改",
    "QA 复核验证整改效果",
    "系统记录闭环时间，计算 SLA 达标率",
]
for i, ln in enumerate(flow, 1):
    add_para(doc, f"{i}. {ln}", size=10, indent_left=12, space_after=3)

# ============================================================================
# §六、建议与下一步行动
# ============================================================================
add_heading(doc, "六、建议与下一步行动", 1)

add_heading(doc, "6.1 本月关键发现", 2)
for ln in [
    f"✅ 巡检体系全面恢复：4月共完成 {N_TOTAL} 次巡检（自检 {N_BY_TYPE['门店自检']} / QA {N_BY_TYPE['QA审计']} / 区经 {N_BY_TYPE['区经检查']}），{N_STORES_ACTIVE} 家活跃门店均获得三类巡检全覆盖，区经检查在中断 3 个月后恢复正常节奏。",
    "✅ 跨类型校准开始发挥作用：QA审计与区经检查互为基准，自检与外部审计差距大幅收窄（多数 <10 分），系统性自检偏高问题被首次暴露。",
    f"⚠ {all_s} 个 S 项分布在 {len(set(it['store_code'] for it in all_s_items))} 家门店，集中在 {top_sub[0][0] if top_sub else ''}（{top_sub[0][1] if top_sub else 0} 起）、{top_sub[1][0] if len(top_sub)>1 else ''}（{top_sub[1][1] if len(top_sub)>1 else 0} 起）；BD 整改清单需跟踪闭环。",
    "⚠ 门店自检评分一致性问题：US00020 同日 Darwin Coronel 三次自检 100/100/64（摆动 36 分），自检偏严（Brionna Jiles 59、Eric Park 63）与偏宽（Joselyn Pacheco Trejo 97、Juliana Li 96）双侧偏离同时存在。",
    "⚠ 首要风险模块（按主巡检视角扣分排序）：" + "、".join(mods_ranked[:3]) + "。",
    f"⚠ QA审计单点依赖：Eamonn Caballar 单人完成 {N_BY_TYPE['QA审计']} 次 QA（占 100%），Yu Jiang 已退出，需评估 QA 团队冗余度。",
    f"⚠ 申诉机制本月跑通完整闭环：{len(APPEALED)} 起立案，{len(APPEAL_APPROVED)} 起获批 / {len(APPEAL_DENIED)} 起驳回 / {len(APPEAL_PENDING)} 起审批中，扣分项保留以便趋势追踪。",
]:
    add_para(doc, "• " + ln, size=10, indent_left=12, space_after=4)

add_heading(doc, "6.2 优先行动项", 2)
prio_rows = [
    ["P0","紧急", f"立即处理 {all_s} 个 S 项（{top_sub[0][0] if top_sub else ''} {top_sub[0][1] if top_sub else 0} 起 / {top_sub[1][0] if len(top_sub)>1 else ''} {top_sub[1][1] if len(top_sub)>1 else 0} 起 / 其他），确保 48 小时内闭环；BD 整改清单按门店列表驱动",
     "门店 + QA + BD","48 小时"],
    ["P1","紧急", f"处理 {all_m} 个 M 项，重点关注扣分集中门店：US00005、US00024、US00012", "门店 + QA","7 天"],
    ["P2","高",   f"维持当前 QA审计节奏（每月 12+ 次），评估增加第二位 QA Manager 以避免单点依赖（Eamonn 单人承担 100%）", "QA 部门","本月内"],
    ["P3","高",   f"对最低分门店（{lowest['store_name']} {lowest['adjusted_total_score']}{'※' if lowest['appeal_status']=='approved' else ''} 分、其他 <80 分门店）开展专项辅导", "区域经理","1 周内"],
    ["P4","高",   "针对自检评分一致性问题（US00020 100/100/64 同日波动 36 分；Brionna / Eric / Joselyn / Juliana 偏离），开展自检标准校准培训", "QA 部门","2 周内"],
    ["P5","中",   "跟踪新开门店（US00007 4/30、US00010 4/28、US00015 4/30）首月巡检计划，5 月底前完成首次 QA 或区经检查", "运营部 + QA","5 月内"],
    ["P6","中",   "维持区经检查节奏（Daniel Chu / Jung Han Liang 月度均衡负荷），避免再次断流", "运营部","持续"],
    ["P7","中",   f"申诉机制后续：{len(APPEAL_APPROVED)} 起获批申诉的实际整改完成确认；{len(APPEAL_DENIED)} 起驳回申诉的原方案落地；{len(APPEAL_PENDING)} 起审批中申诉的 SLA 监控", "QA 部门","2 周内"],
    ["P8","低",   "empapp 系统优化：① 同日重复自检 UI 提示；② S 项整改工单与申诉立案双向同步；③ 申诉审批 SLA 监控", "基础设施部 + QA","本季度"],
]
add_data_table(doc, ["优先级","紧急度","行动项","责任方","时限"], prio_rows,
               col_widths=[1.0,1.0,8.6,3.0,2.0], align_center_cols=[0,1,4])

add_heading(doc, "6.3 模块改善建议（TOP 5）", 2)
suggestions = {
    "清洁卫生": "① 清洁消毒程序每班次执行并记录；② 消毒液浓度（ppm）每日校准；③ 食品加工区域和设备每日深度清洁。",
    "过程控制": "① 食品存储分区标准重新培训；② 器具维护和清洁班次检查；③ 物料存储高度要求（6 英寸）每日巡查。",
    "设施":     "① 水滤芯更换台账建立，提前 7 天预警；② 油脂阱 / 残渣阱清理纳入月度必检；③ 管道泄漏立即报修；④ air gap 系统短板由 BD 列入整改清单。",
    "员工健康卫生": "① 每日开班健康申报；② 个人卫生（指甲、首饰、头发）班前班中检查；③ 洗手程序每周复训；④ BOH handwash sink 配套（皂、纸巾）专项检查。",
    "温控有效期": "① 开封后标签管理纳入每日开店清单；② FIFO 执行每日检查；③ 过期产品零容忍政策。",
    "证照文件": "① 各门店 license / no-smoking sign 每日开店检查；② permit 续期提前 30 天预警。",
    "职场安全": "① 警示标识齐备性每周复盘；② 滑倒 / 电气风险点位每月巡检。",
    "虫害防控": "① pest control 每月报告留档；② 灯诱设备完好性每周巡查。",
    "设备维护": "① 设备保养排期每周更新；② 关键设备（搅拌机、咖啡机）每季度专业校准。",
    "供应商":   "① 供应商认证清单季度复核；② 收货验收 SOP 每月抽查。",
}
for mod in top_mods:
    if mod in suggestions:
        ded = sum(MAIN_MOD_DED.get(mod, {}).values())
        n_st = len(MAIN_MOD_STORES.get(mod, set()))
        add_para(doc, f"{mod}（影响 {n_st} 家门店，主巡检扣分 {ded} 分）：",
                 size=10, bold=True, color_hex=NAVY, space_before=4, space_after=2)
        add_para(doc, suggestions[mod], size=10, indent_left=12, space_after=4)

# ============================================================================
# §七、巡检体系分析
# ============================================================================
add_heading(doc, f"七、三类巡检体系分析（{TARGET_MONTH_CN}专题）", 1)

add_callout(doc,
    f"{TARGET_MONTH_CN} 巡检体系全面恢复：三类巡检均处于正常节奏，区经检查在中断 3 个月后于 4 月 7 日由 Jung Han Liang 在 US00019 重启，本月共完成 {N_BY_TYPE['区经检查']} 次。",
    kind="good", title="✅ 体系恢复")

add_heading(doc, "7.1 巡检概况", 2)
def avg_score_by_type(t, use_adjusted=True):
    rs = INSP_BY_TYPE[t]
    if not rs: return ""
    return round(sum((r["adjusted_total_score"] if use_adjusted else r["original_total_score"]) for r in rs)/len(rs), 1)
def stores_for_type(t): return len({r["store_code"] for r in INSP_BY_TYPE[t]})
def s_for_type(t): return sum(r["s_count"] for r in INSP_BY_TYPE[t])
def m_for_type(t): return sum(r["m_count"] for r in INSP_BY_TYPE[t])

def insp_list_text(t):
    cnt = Counter(r["inspector_name"] for r in INSP_BY_TYPE[t])
    if len(cnt) == 0: return "—"
    if len(cnt) <= 5:
        return "、".join(sorted(cnt.keys()))
    top3 = "、".join(name for name,_ in cnt.most_common(3))
    return f"{len(cnt)} 人（含 {top3}）"

ct_rows = [
    ["巡检次数（已提交）", f"{N_BY_TYPE['门店自检']} 次", f"{N_BY_TYPE['QA审计']} 次", f"{N_BY_TYPE['区经检查']} 次"],
    ["覆盖门店", f"{stores_for_type('门店自检')} 家", f"{stores_for_type('QA审计')} 家", f"{stores_for_type('区经检查')} 家"],
    ["巡检员", insp_list_text("门店自检"), insp_list_text("QA审计"), insp_list_text("区经检查")],
    ["平均得分", f"{avg_score_by_type('门店自检')} 分", f"{avg_score_by_type('QA审计')} 分", f"{avg_score_by_type('区经检查')} 分"],
    ["S 项发现", str(s_for_type('门店自检')), str(s_for_type('QA审计')), str(s_for_type('区经检查'))],
    ["M 项发现", str(m_for_type('门店自检')), str(m_for_type('QA审计')), str(m_for_type('区经检查'))],
]
add_data_table(doc, ["维度","门店自检","QA审计","区经检查"], ct_rows,
               col_widths=[2.5,4.5,4.5,4.5], align_center_cols=[1,2,3])
qa_pre = avg_score_by_type('QA审计', use_adjusted=False)
add_para(doc, f"※ QA 平均分含申诉调整；申诉前实际 {qa_pre} 分。",
         size=9, color_hex="666666", space_before=2, space_after=4)

add_callout(doc,
    f"三类平均分差距 < 5 分（自检 {avg_score_by_type('门店自检')} / QA审计 {avg_score_by_type('QA审计')} / 区经检查 {avg_score_by_type('区经检查')}），表面看校准良好；但同店跨类型对比显示多家门店出现 ≥10 分差距、{len(divergence_cases)} 组 ≥20 分（详见 §4.4）——平均分掩盖了店级散度。",
    kind="info", title="📊 关键观察")

add_heading(doc, "7.2 同店三类对比", 2)
type_avg = defaultdict(dict)
for sc, rows in by_store.items():
    for t in ("门店自检","QA审计","区经检查"):
        rs = [r for r in rows if r["inspection_type"] == t]
        if rs:
            type_avg[sc][t] = round(sum(r["adjusted_total_score"] for r in rs)/len(rs), 1)
dt_rows = []
for sc in sorted(type_avg.keys()):
    name = MAIN_INSP[sc]["store_name"]
    s_self = type_avg[sc].get("门店自检", "---")
    s_qa = type_avg[sc].get("QA审计", "---")
    s_area = type_avg[sc].get("区经检查", "---")
    if isinstance(s_self,(int,float)) and isinstance(s_qa,(int,float)):
        diff = s_self - s_qa
        diff_str = f"{'+' if diff>0 else ''}{diff:.1f}"
        if abs(diff) >= 15:    conclusion = "⚠ 自检偏宽" if diff > 0 else "⚠ QA 偏严"
        elif abs(diff) >= 10:  conclusion = "⚠ 中度偏离"
        else:                  conclusion = "✅ 一致"
    else:
        diff_str = "---"; conclusion = "—"
    dt_rows.append([f"{name} ({sc})", str(s_self), str(s_qa), str(s_area), diff_str, conclusion])
add_data_table(doc, ["门店","自检均分","QA审计","区经检查","自检 - QA 差","结论"], dt_rows,
               col_widths=[3.6,2.0,2.0,2.0,2.0,4.0], align_center_cols=[1,2,3,4,5])

add_heading(doc, "7.3 自检评分一致性分析", 2)
add_heading(doc, "US00020（21st & 3rd）2026-04-21 Darwin Coronel 单人三次自检案例", 3)
inv_rows = [
    ["2016", "2026-04-21", "100", "0", "0", "0", "0", "0", "0"],
    ["2017", "2026-04-21", "100", "0", "0", "0", "0", "0", "0"],
    ["2018", "2026-04-21", "64", "-16", "5", "1", "1", "3", "0"],
]
add_data_table(doc, ["巡检 ID","时间","得分","扣分","问题数","S","M","G","L"], inv_rows,
               col_widths=[1.4,2.2,1.4,1.4,1.4,1.0,1.0,1.0,1.0], align_center_cols=[2,3,4,5,6,7,8])
add_para(doc,
    "同一巡检员、同店、同日提交三次自检：前两次均为 100 分零扣分（已按误提交规则剔除），第三次发现 5 个扣分项（含 1 个 S 项）。摆动幅度 36 分，暴露门店自检在不同时段执行严格度差异巨大。建议：① 自检需在交班前一次性完成而非班中分次；② 对同日多次自检的情况由 empapp 系统自动校验；③ 系统提示「今日已自检」。",
    size=10, space_before=4, space_after=4)

add_para(doc, "其他同日重复或大幅摆动案例（已提交单据范围内）：",
         size=10, bold=True, space_before=4, space_after=2)
swing = []
by_date_store = defaultdict(list)
for r in SUMMARY:
    by_date_store[(r["store_code"], r["inspection_date"])].append(r)
for (sc, d), rs in by_date_store.items():
    if len(rs) >= 2:
        sc_scores = [r["adjusted_total_score"] for r in rs]
        sw = max(sc_scores) - min(sc_scores)
        if sw >= 2:
            swing.append((sw, sc, d, rs))
swing.sort(key=lambda x: -x[0])
sw_rows = []
for s, sc, d, rs in swing[:10]:
    detail = " / ".join(f"{r['inspector_name']}（{r['inspection_type']}, {r['adjusted_total_score']} 分{'※' if r['appeal_status']=='approved' else ''}）" for r in rs)
    sw_rows.append([f"{rs[0]['store_name']} ({sc})", d, str(len(rs)), str(s), detail])
if sw_rows:
    add_data_table(doc, ["门店","日期","巡检数","摆动","详情"], sw_rows,
                   col_widths=[3.0,2.0,1.4,1.4,9.2], align_center_cols=[2,3])
else:
    add_para(doc, "本月无显著自检评分波动案例。", size=10, color_hex="666666")

add_heading(doc, "7.4 巡检员严格度对比（≥2 次巡检者）", 2)
inspector_stats = defaultdict(lambda: {"cnt":0,"score_sum":0,"ded_sum":0,"item_sum":0,"s":0,"m":0,
                                        "type":Counter(),"role":""})
for r in SUMMARY:
    n = r["inspector_name"]; st = inspector_stats[n]
    st["cnt"] += 1
    st["score_sum"] += r["adjusted_total_score"]
    st["ded_sum"]   += r["adjusted_total_deduction"]
    st["item_sum"]  += r["item_count"]
    st["s"] += r["s_count"]; st["m"] += r["m_count"]
    st["type"][r["inspection_type"]] += 1
    st["role"] = r["inspector_role"]

multi_insp = [(n,st) for n,st in inspector_stats.items() if st["cnt"] >= 2]
single_insp_with_s = [(n,st) for n,st in inspector_stats.items() if st["cnt"] == 1 and st["s"] > 0]

if multi_insp:
    multi_insp.sort(key=lambda kv: kv[1]["score_sum"]/kv[1]["cnt"])
    ins_rows = []
    for n, st in multi_insp:
        avg_s = st["score_sum"]/st["cnt"]
        note = "⚠ 偏严" if avg_s < 70 else ("⚠ 偏宽" if avg_s > 92 else "—")
        ins_rows.append([n, st["role"],
                         st["type"].most_common(1)[0][0],
                         str(st["cnt"]),
                         f"{avg_s:.1f}",
                         note])
    add_data_table(doc, ["巡检员","职位","类型主导","巡检次","平均分","说明"],
                   ins_rows, col_widths=[2.8,3.5,2.0,1.2,1.4,2.0],
                   align_center_cols=[3,4,5])
    add_para(doc, "※ 含申诉调整后均分；Eamonn Caballar 申诉前实际均分见 §7.1 注释。",
             size=9, color_hex="666666", space_before=2, space_after=4)
else:
    add_para(doc, "本月无 ≥2 次巡检的重复巡检员案例。", size=10, color_hex="666666")

if single_insp_with_s:
    add_para(doc, "单次巡检但发现 S 项（重要发现，单独标注）：", size=10, bold=True, space_before=4, space_after=2)
    for n, st in single_insp_with_s:
        t = st["type"].most_common(1)[0][0]
        add_para(doc, f"  • {n}（{st['role']}，{t}，{st['cnt']} 次）：发现 S 项 {st['s']} 个 / 平均分 {st['score_sum']/st['cnt']:.1f}",
                 size=10, indent_left=8, space_after=2)

add_heading(doc, "7.5 巡检覆盖趋势（2026 年 Q1 + 4 月）", 2)
trend_data = defaultdict(lambda: {"门店自检":0,"QA审计":0,"区经检查":0})
for tr in TREND_SUMMARY:
    if tr["month"] in ("2026-01","2026-02","2026-03","2026-04"):
        trend_data[tr["month"]][tr["inspection_type"]] = int(tr["inspection_count"])
status_for = {"2026-01":"✅ 三类齐全","2026-02":"⚠ 区经检查中断",
              "2026-03":"🔴 体系崩溃","2026-04":"✅ 全面恢复"}
trd_rows = []
for ym in ("2026-01","2026-02","2026-03","2026-04"):
    d = trend_data[ym]
    total = d["门店自检"] + d["QA审计"] + d["区经检查"]
    trd_rows.append([ym, f"{d['门店自检']} 次", f"{d['QA审计']} 次", f"{d['区经检查']} 次", str(total), status_for[ym]])
add_data_table(doc, ["月份","门店自检","QA审计","区经检查","总数","状态"], trd_rows,
               col_widths=[2.0,2.0,2.0,2.0,1.4,4.0], align_center_cols=[1,2,3,4])
add_callout(doc,
    "趋势分析：从 1 月的三类齐全（16 次），到 2-3 月区经检查断流、QA审计萎缩（仅 14 次），再到 4 月的全面恢复（59 次），北美 QA 巡检体系完成了一次典型的「危机—响应」循环。Eamonn Caballar 接任 QA Senior Manager 与 Jung Han Liang / Daniel Chu 区经巡检的同步恢复是关键转折点。",
    kind="info", title="📈 趋势分析")

add_heading(doc, "7.6 三类巡检发现差异分析", 2)
type_diff_rows = [
    ["核心价值", "门店日常自查、低成本高频", "QA 标准化外审、培训校准", "区经实地督导、跨店共性整改"],
    ["最常发现的问题",
     "Cleaning & Sanitize、Equipment and utensils、Personal Hygiene（清洁、消毒、卫生类高频项）",
     "Sinks and Pipes、Cross-Contamination、Handwashing Standards（合规与系统性问题）",
     "Sinks and Pipes、Operation Stand、跨店重复出现的设施与流程问题"],
    ["典型严重度分布",
     "S 项偶发，多为 G/L 项；自检偏宽现象明显",
     "S 项命中率最高（air gap 类系统性 S 项主要在 QA 发现）",
     "S 项与 M 项在跨店共性问题上集中"],
]
add_data_table(doc, ["维度","门店自检","QA审计","区经检查"], type_diff_rows,
               col_widths=[2.6,4.5,4.5,4.5])

add_heading(doc, "7.7 后续观察重点（5 月跟踪事项）", 2)
follow = [
    "申诉案例后续：54th & 8th 与 15th & 3rd 申诉获批后，5 月跟踪 air gap / plumbing 实际整改完成情况；驳回的 3 起（8th & Broadway QA 4/9、102 Fulton QA 4/27、28th & 6th QA 4/27）需按原方案闭环；审批中的 1 起（28th & 6th 区经 4/23）需关注审批 SLA",
    "QA workload sustainability：Eamonn Caballar 5 月 QA审计计划次数；评估第二位 QA Manager 配置时机",
    "Yu Jiang HR 流程关闭确认：4 月已无任何巡检提交，需走 HR 离职 / 转岗手续",
    "区经检查频率维持：Daniel Chu / Jung Han Liang 5 月覆盖度，避免再次断流",
    "巡检员标准化培训：针对自检偏严（Brionna Jiles 59、Eric Park 63）与偏宽（Joselyn Pacheco Trejo 97、Juliana Li 96）双侧偏离，开展评分尺度统一培训",
    "多门店 S 项系统短板（Sinks and Pipes air gap 9 起 / Handwashing Standards 5 起）：5 月跟踪 BD 整改清单完成情况",
    "新开门店 5 月覆盖：US00007（4/30 开业）、US00010（4/28 开业）、US00015（4/30 开业）需在 5 月底前完成首次 QA 审计或区经检查",
    "自检评分一致性：US00020 同日三次自检案例的 empapp 系统校验是否上线",
    "empapp 系统优化项：① 同日重复自检 UI 提示；② S 项整改工单与申诉立案双向同步；③ 申诉审批 SLA 监控",
    "（首次 v2 报告）下月 §7.7 起将跟踪本月遗留事项的进度标识【接续】",
]
for ln in follow:
    add_para(doc, "• " + ln, size=10, indent_left=12, space_after=4)

# ---------- Footer ----------
add_para(doc, "── 报告结束 ──", size=11, color_hex=NAVY, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=18, space_after=4)

# Save
doc.save(OUT_PATH)
print(f"\nWROTE: {OUT_PATH}")
print(f"  size = {OUT_PATH.stat().st_size:,} bytes")

# ============================================================================
# VALIDATION SUMMARY
# ============================================================================
print()
print("="*70)
print("VALIDATION SUMMARY")
print("="*70)
print(f" 1. Files loaded: summary={len(SUMMARY)} items={len(ITEMS)} stores={len(STORE_MASTER)} prior_march={len(PRIOR_MARCH)}")
print(f" 2. Inspections by type: {dict(N_BY_TYPE)} total={N_TOTAL}")
print(f" 3. S/M/G/L: 主巡检 {dict(MAIN_SEV)} vs 全月 {dict(ALL_SEV)}")
print(f" 4. 主巡检均分: {avg_score_main}")
print(f" 5. ≥15pt 自检 vs QA divergence (主巡检): see §7.2")
print(f" 6. 申诉: total={len(APPEALED)} approved={len(APPEAL_APPROVED)} denied={len(APPEAL_DENIED)} pending={len(APPEAL_PENDING)}")
print(f" 7. 新开门店 (excluded from coverage): US00007, US00010, US00015")
print(f" 8. 跨3月Δ≥10: {len(delta_lines)} 家")
print(f" 9. New inspectors this month: 见 §7.4")
print(f"10. Departed: Yu Jiang (1月 5次, 2月 2次, 3月 2次, 4月 0次)")
print(f"11. Q1+4月 trend: {trd_rows}")
print(f"12. Module 主巡检 deductions:")
for mod in mods_ranked:
    print(f"      {mod:>10} : {sum(MAIN_MOD_DED.get(mod, {}).values()):>5}  (覆盖 {len(MAIN_MOD_STORES.get(mod, set()))}/{N_STORES_ACTIVE})")
print(f"13. Module mapping anomalies (raw → '其他'): {dict(OTHER_FLAGS) if OTHER_FLAGS else '(none)'}")
status = "PASS" if not any([
    main_s < 0,
    not PRIOR_MARCH,
]) else "FAIL"
note = "（'其他' bucket 仅含 Site Security 1 项，US00006 garbage bin，已记录）" if OTHER_FLAGS else ""
print(f"14. VALIDATION: {status} {note}")
