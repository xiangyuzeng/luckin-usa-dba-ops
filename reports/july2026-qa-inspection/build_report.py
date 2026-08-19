#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
July 2026 QA monthly report — V1 正式月度版.
Consumes july2026_qa_datapack.json + core CSVs, emits the docx using the
June (LCNA-QA-2026-006) structural + visual baseline.
"""
import csv, json
from collections import defaultdict, Counter
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ============================================================================
# RUN PARAMETERS
# ============================================================================
TARGET_MONTH_CN = "2026年07月"
DOC_ID          = "LCNA-QA-2026-007"
TODAY           = "2026-08-03"
RANGE_FROM, RANGE_TO, DAYS = "2026-07-01", "2026-07-31", 31

HERE     = Path(__file__).resolve().parent
OUT_DIR  = HERE / "output"; OUT_DIR.mkdir(exist_ok=True)
OUT_PATH = OUT_DIR / f"QA门店巡检月度分析报告_{TARGET_MONTH_CN}_v1.docx"

NAVY="1F4E79"; BLUE="2E75B6"; TEXT="2C3E50"
BORDER="BFBFBF"; ALT_FILL="F2F2F2"; PINK_FILL="FCE4EC"
CRITICAL_FG="C0392B"; WARN_FG="E67E22"; HEALTHY_FG="27AE60"; INFO_FG="0365C0"
CALLOUT_YELLOW="FFF2CC"; CALLOUT_GREEN="E2EFDA"; CALLOUT_BLUE="D6E4F0"; CALLOUT_RED="FCE4EC"

MODULES_10 = ["清洁卫生","过程控制","设施","证照","职业安全",
              "虫害防控","温控有效期","员工健康卫生","设备维护","供应链"]
MOD_SHORT = {"清洁卫生":"清洁","过程控制":"过程","设施":"设施","证照":"证照","职业安全":"职安",
             "虫害防控":"虫害","温控有效期":"温控","员工健康卫生":"员工","设备维护":"设备","供应链":"供应"}

# ============================================================================
# DATA
# ============================================================================
PACK = json.loads((HERE/"july2026_qa_datapack.json").read_text(encoding="utf-8"))
DER  = json.loads((HERE/"derived.json").read_text(encoding="utf-8"))

def load_csv(p):
    with open(p, encoding="utf-8-sig") as f: return list(csv.DictReader(f))
SUMMARY = load_csv(HERE/"july2026_inspection_summary.csv")
ITEMS   = load_csv(HERE/"july2026_inspection_items.csv")
for r in SUMMARY:
    for k in ("inspection_id","adjusted_total_score","original_total_score",
              "adjusted_total_deduction","original_total_deduction","is_appealed",
              "S_count","M_count","G_count","L_count"):
        r[k] = int(r[k]) if r[k] not in ("","None",None) else 0
for r in ITEMS:
    r["deduction"] = int(r["deduction"]); r["inspection_id"] = int(r["inspection_id"])
    r["is_appealed_finding"] = int(r["is_appealed_finding"])

COVER   = PACK["cover"]
MGMT    = PACK["mgmt_summary"]
PRIMARY = DER["primary_by_store"]          # store_code -> dict
JUNP    = DER["june_primary"]
NPRIM   = len(PRIMARY)                      # 21
NAME    = {r["store_code"]: r["store_name"] for r in SUMMARY}
def disp(c): return f"{NAME.get(c,c)}（{c}）"
def sn(c):   return NAME.get(c, c)

PRIM_IIDS = set(DER["primary_iids"])
PITEMS = [it for it in ITEMS if it["inspection_id"] in PRIM_IIDS]

FM = COVER["findings_full_month"]; PM = COVER["findings_primary"]
TC = COVER["type_counts"]
AP = COVER["appeals"]
PRIM_AVG = COVER["primary_avg"]                   # 89.4
LL_AVG   = MGMT["like_for_like_avg"]              # 92.1
PRIOR_AVG, PRIOR_N = 88.7, 18
# pre-appeal (原始) primary averages
ORIG_AVG_JUL = round(sum(p["orig"] for p in PRIMARY.values())/NPRIM, 1)          # 76.2
ORIG_AVG_JUN = 79.9
N_APPEAL_ADJ = sum(1 for p in PRIMARY.values() if p["appeal_status"]=="approved")  # 13
AVG_UPLIFT   = round(sum(PRIMARY[c]["score"]-PRIMARY[c]["orig"]
                         for c in PRIMARY if PRIMARY[c]["appeal_status"]=="approved")/N_APPEAL_ADJ, 1)

REPEAT = PACK["sX_repeat_offenders"]

print(f"[load] summary={len(SUMMARY)} items={len(ITEMS)} primary={NPRIM} "
      f"avg={PRIM_AVG} orig={ORIG_AVG_JUL}")

# ============================================================================
# DOCX HELPERS  (June/April visual baseline)
# ============================================================================
def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)

def set_cell_borders(cell, color_hex=BORDER, sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    bdrs = OxmlElement("w:tcBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"),"single"); b.set(qn("w:sz"),str(sz))
        b.set(qn("w:space"),"0"); b.set(qn("w:color"), color_hex)
        bdrs.append(b)
    tcPr.append(bdrs)

def set_run_font(run, name="Arial", size=10, bold=False, color_hex=TEXT, italic=False):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfont = rpr.find(qn("w:rFonts"))
    if rfont is None:
        rfont = OxmlElement("w:rFonts"); rpr.append(rfont)
    rfont.set(qn("w:ascii"), name); rfont.set(qn("w:hAnsi"), name)
    rfont.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    if color_hex: run.font.color.rgb = RGBColor.from_string(color_hex)

def add_para(doc, text, size=10, bold=False, color_hex=TEXT, align=None,
             space_before=0, space_after=4, indent_left=0, italic=False):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent_left: p.paragraph_format.left_indent = Pt(indent_left)
    if text:
        r = p.add_run(text); set_run_font(r, size=size, bold=bold, color_hex=color_hex, italic=italic)
    return p

def add_bullet(doc, text, size=10, color_hex=TEXT, bold=False):
    return add_para(doc, "• " + text, size=size, color_hex=color_hex, bold=bold,
                    indent_left=12, space_after=3)

def add_heading(doc, text, level=1):
    sizes = {1:16, 2:13, 3:11}; colors = {1:NAVY, 2:BLUE, 3:NAVY}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level==1 else 8)
    p.paragraph_format.space_after  = Pt(6 if level==1 else 4)
    r = p.add_run(text); set_run_font(r, size=sizes[level], bold=True, color_hex=colors[level])

def add_callout(doc, lines, kind="warn", title=None):
    bg, bar = {"warn":(CALLOUT_YELLOW,WARN_FG), "good":(CALLOUT_GREEN,HEALTHY_FG),
               "info":(CALLOUT_BLUE,INFO_FG), "critical":(CALLOUT_RED,CRITICAL_FG)}[kind]
    table = doc.add_table(rows=1, cols=1); table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, bg); set_cell_borders(cell, color_hex=bar, sz=8)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    if isinstance(lines, str): lines = [lines]
    if title:
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title); set_run_font(r, size=10, bold=True, color_hex=bar)
        for ln in lines:
            pp = cell.add_paragraph(); pp.paragraph_format.space_after = Pt(2)
            rr = pp.add_run(ln); set_run_font(rr, size=10, color_hex=TEXT)
    else:
        for i, ln in enumerate(lines):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(ln); set_run_font(r, size=10, color_hex=TEXT)
    add_para(doc, "", size=4, space_after=4)

def add_data_table(doc, headers, rows, col_widths=None, align_center_cols=None,
                   row_fills=None, fsize=9):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = "Table Grid"
    if col_widths:
        for ci, w in enumerate(col_widths):
            for c in t.columns[ci].cells: c.width = Cm(w)
    for i, h in enumerate(headers):
        c = t.cell(0, i); set_cell_shading(c, NAVY)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.text = ""; p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        r = p.add_run(h); set_run_font(r, size=fsize+1, bold=True, color_hex="FFFFFF")
    for ri, row in enumerate(rows):
        fill = (row_fills or {}).get(ri)
        for ci, val in enumerate(row):
            c = t.cell(ri+1, ci)
            if fill: set_cell_shading(c, fill)
            elif ri % 2 == 1: set_cell_shading(c, ALT_FILL)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            c.text = ""; p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
            txt, opts = (val if isinstance(val, tuple) else (val, {}))
            if align_center_cols and ci in align_center_cols:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif opts.get("align") == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(txt))
            set_run_font(r, size=opts.get("size", fsize), bold=opts.get("bold", False),
                         color_hex=opts.get("color", TEXT))
    add_para(doc, "", size=4, space_after=4)
    return t

def add_page_field(run, instr):
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"),"begin")
    it = OxmlElement("w:instrText"); it.text = instr
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"),"end")
    run._r.append(fb); run._r.append(it); run._r.append(fe)

def setup_page_footer(section, doc_id, today):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.text = ""
    r1 = p.add_run(f"{doc_id} | 编制：曾翔宇 | 日期：{today} | 第 "); set_run_font(r1, size=9, color_hex="666666")
    r2 = p.add_run(""); set_run_font(r2, size=9, color_hex="666666"); add_page_field(r2, "PAGE")
    r3 = p.add_run(" 页 / 共 "); set_run_font(r3, size=9, color_hex="666666")
    r4 = p.add_run(""); set_run_font(r4, size=9, color_hex="666666"); add_page_field(r4, "NUMPAGES")
    r5 = p.add_run(" 页"); set_run_font(r5, size=9, color_hex="666666")

def clean(s):
    return (s or "").replace("\r\n"," / ").replace("\n"," / ").replace("\r"," / ").strip()

def sdelta(x, dec=0):
    """Signed delta with a plain '0' instead of '+0'."""
    if x == 0: return "0"
    return f"{x:+.{dec}f}" if dec else f"{x:+d}"

# ============================================================================
# DOCUMENT
# ============================================================================
doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.59); section.page_height = Cm(27.94)
section.top_margin = Cm(1.27); section.bottom_margin = Cm(1.27)
section.left_margin = Cm(1.59); section.right_margin = Cm(1.59)
setup_page_footer(section, DOC_ID, TODAY)
doc.styles["Normal"].font.name = "Arial"; doc.styles["Normal"].font.size = Pt(10)

# ---------- COVER ----------
add_para(doc, "瑞幸咖啡北美", size=22, bold=True, color_hex=NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=8)
add_para(doc, "QA门店巡检月度分析报告", size=18, bold=True, color_hex=NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para(doc, "Monthly QA Store Audit Analysis Report", size=12, color_hex=BLUE, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
add_para(doc, TARGET_MONTH_CN, size=20, bold=True, color_hex=NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=80)
add_para(doc, "质量保障部 / 基础设施部", size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para(doc, f"编制：曾翔宇      日期：{TODAY}", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para(doc, f"报告编号：{DOC_ID}    |    状态：V1 正式月度版", size=10, color_hex=BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
doc.add_page_break()

# ---------- 文档信息 ----------
add_heading(doc, "文档信息", 1)
new_txt = "、".join(f"{sn(s['store'])} {s['store']}（开业 {s['open_date']}）" for s in COVER["new_stores_this_month"])
doc_info_rows = [
    ["报告编号", DOC_ID],
    ["报告周期", TARGET_MONTH_CN],
    ["数据范围", f"{RANGE_FROM} 至 {RANGE_TO}（共 {DAYS} 天）"],
    ["有效门店", f"{NPRIM} 家已巡检 / {NPRIM} 家在营运营门店（覆盖率 100%；6 月可比基线 {PRIOR_N} 家）；"
               f"本月新纳管 3 家：{new_txt}"],
    ["巡检类型", f"门店自检（{TC['门店自检']} 次）+ QA审计（{TC['QA审计']} 次）+ 区经检查（{TC['区经检查']} 次）"
               f"= 共 {TC['total']} 次（已提交；本月无误提交剔除）"],
    ["问题总数", f"全月 {FM['total']} 个扣分项（S 项 {FM['S']}、M 项 {FM['M']}、G 项 {FM['G']}、L 项 {FM['L']}）；"
               f"主巡检 {PM['total']} 个（S 项 {PM['S']}、M 项 {PM['M']}、G 项 {PM['G']}、L 项 {PM['L']}）"],
    ["编制人", "曾翔宇"],
    ["部门", "质量保障部 / 基础设施部"],
    ["数据来源", "empapp 门店稽核系统（aws-luckyus-opqualitycontrol-rw / luckyus_opqualitycontrol）"],
    ["状态", "V1 正式月度版（基于完整 7 月数据）"],
]
add_data_table(doc, ["项目","内容"], doc_info_rows, col_widths=[3.0, 15.0])

# ---------- 数据说明 ----------
add_heading(doc, "⚠ 数据说明（正式月度版）", 1)
data_notes = [
    f"(a) 本报告基于 2026 年 7 月完整月度数据（{TC['total']} 次提交，本月未发现误提交单据），为正式月度版，可作为 7 月月度结论使用。",
    f"(b) 申诉口径（双轨制）：本月共 {AP['total']} 起申诉立案（{AP['approved']} 获批 / {AP['denied']} 驳回 / {AP['pending']} 审批中），"
    f"涉及 {PACK['s4_3']['appealed_findings']} 条 finding。获批申诉已按调整后分数计入主巡检得分（详见 §4.3），"
    f"但 S/M/G/L 严重度计数一律保留（严重度不因申诉撤销而改变）。因此 §1.2 中「得分」与「扣分」不构成 100 − 扣分 的关系。",
    f"(c) ⚠ 双口径提示：本月主巡检官方（申诉调整后）均分 {PRIM_AVG}，而申诉前原始均分仅 {ORIG_AVG_JUL}（6 月分别为 {PRIOR_AVG} / {ORIG_AVG_JUN}）。"
    f"两者差 {round(PRIM_AVG-ORIG_AVG_JUL,1)} 分，全部来自 {N_APPEAL_ADJ} 家门店的申诉获批（平均提分 {AVG_UPLIFT} 分）。"
    f"解读月度趋势时应同时参考两个口径，详见「管理摘要」与 §4.3。",
    f"(d) 巡检覆盖：{NPRIM} 家在营运营门店全部完成至少 1 次巡检，覆盖率 {NPRIM}/{NPRIM} = 100%。6 月遗留的 2 家 6/30 新开业门店"
    f"（US00009 48th & 3rd、US00013 Grand Central Terminal）已于本月完成首检；US00021 128 W 32nd St 于 7/16 开业并于 7/28 完成首检。"
    f"3 家新店本月均未纳入 QA 审计（QA 覆盖 18/21），其主巡检为区经检查。",
    f"(e) QA 人员状态：Eamonn Caballar 独立完成全部 {TC['QA审计']} 次 QA 审计（覆盖 18 家门店，月均分 92.5，节奏 7/7–7/24），"
    f"为连续第四个月单点执行；Yu Jiang HR 离职 / 转岗流程仍未关闭（7 月无任何巡检）。",
    f"(f) 区经巡检：Jung Han Liang 完成 {TC['区经检查']} 次（覆盖 {NPRIM}/{NPRIM} 全部门店，月均分 78.9，节奏 7/8–7/31），"
    f"仍为单人执行（无 Daniel Chu），需关注负荷与断流风险。",
    "(g) 标准规则：① 主巡检优先级 QA审计 > 区经检查 > 门店自检，同优先级取最近日期；② 仅采用已提交单据（status=1），草稿与软删除已剔除；"
    "③ 同人同店同日重复 100 分零扣分自检按误提交规则剔除（本月剔除 0 条）；④ 严重度映射 deduction_type 1→S(−5)、3→M(−5)、2→G(−2)、4→L(−1)，"
    "含 S 项者另计 −20 惩罚项；本月 94 张单据分数公式校验 100% 通过（0 处偏差）。",
    "(h) 模块分类：10 模块严格分类，本月无 UNMAPPED 模块（「Site Security」→ 职业安全 映射沿用 2026-07-01 确认口径）。"
    "本月主巡检 证照、供应链 两模块无扣分，未在 §2.2 / §4.5 列示。",
]
add_callout(doc, data_notes, kind="warn")

# ---------- 管理摘要 ----------
add_heading(doc, "管理摘要", 1)
add_para(doc,
    f"本月主巡检官方平均分 {PRIM_AVG} 分（较 6 月 {PRIOR_AVG} 分 ↑{round(PRIM_AVG-PRIOR_AVG,1)} 分），覆盖 {NPRIM} 家门店"
    f"（覆盖率 100%，6 月可比基线 {PRIOR_N} 家），主巡检 {PM['total']} 项有效扣分项，全月共 {FM['total']} 项。"
    f"巡检总量 {TC['total']} 次创有记录以来新高（6 月 85 次，+9 次）。",
    size=10, space_after=8)

add_callout(doc,
    [f"7 月共 {TC['total']} 次巡检（自检 {TC['门店自检']} / QA {TC['QA审计']} / 区经 {TC['区经检查']}），较 6 月 85 次 +9 次，为有记录以来最高。"
     f"{NPRIM} 家在营门店实现 100% 覆盖，含 3 家新开业门店首检落地（6 月遗留的 US00009、US00013 已闭环）。"
     f"52nd & Madison（US00027）由 6 月最大跌幅店（64）反弹至 90（+26），54th & 8th（US00005）71→84（+13）结束连续两月 <80。"],
    kind="good", title="🟢 巡检体量与覆盖创新高")

add_callout(doc,
    [f"官方均分 {PRIOR_AVG}→{PRIM_AVG}（+{round(PRIM_AVG-PRIOR_AVG,1)}）呈改善，但申诉前原始均分 {ORIG_AVG_JUN}→{ORIG_AVG_JUL}"
     f"（{round(ORIG_AVG_JUL-ORIG_AVG_JUN,1)}）实为下滑：申诉获批门店由 6 月 7 家增至 7 月 {N_APPEAL_ADJ} 家，平均提分 {AVG_UPLIFT} 分。",
     f"本月 QA 审计 19 次原始均分仅 77.9，经申诉调整后升至 92.5（+14.6）；已决申诉 {AP['approved']} 起全部获批、0 起驳回，"
     f"连同 6 月合计连续两月 23 获批 / 0 驳回。分数改善主要由申诉机制驱动，而非现场问题减少——全月发现项由 498 增至 {FM['total']}（+{FM['total']-498}）。",
     f"建议：申诉审批引入第二审核人与「整改证据留存」强制项，并在月度看板同时披露原始分与调整分。"],
    kind="critical", title="🔴 分数改善主要来自申诉调整，非现场改善")

add_callout(doc,
    [f"Sinks and Pipes（air gap / 管道）全月 {PACK['s3_3']['by_sub_item'][0]['S_count']} 起 S 项、涉及 "
     f"{PACK['s3_3']['by_sub_item'][0]['stores']} 家门店（6 月 13 起 / 10 家，↑4 起），为连续第三个月最大系统性短板且不降反升。"
     f"其中 {len(REPEAT['repeat'])} 家连续两月复现：{'、'.join(sn(c) for c in REPEAT['repeat'])}。",
     f"主巡检 13 个 S 项中，10 个为 QA 审计发现的 Sinks and Pipes，且全部经申诉获批而扣分归零——分数已修复，物理问题未根治。"
     f"6 月 §6.2 P0（air gap 跨月治理）本月未见成效，需升级为 BD 层级专项工程（详见 §3.6 跨月台账与 §6.2 P0）。"],
    kind="critical", title="🔴 air gap 系统性短板连续三月未根治且扩大")

add_callout(doc,
    [f"3 家新纳管门店中 2 家首检低于 80 分：48th & 3rd（US00009）60 分（2 个 S 项：无擦手纸、变质牛奶未处置于拖把池）、"
     f"128 W 32nd St（US00021）64 分（1 个 S 项：奇亚籽混入奶昔粉容器）；Grand Central Terminal（US00013）96 分表现良好。"
     f"3 家新店本月均未获 QA 审计覆盖，首检全部由区经完成。建议新店开业 30 天内强制 QA 首检 + 开业前基础规范培训验收。",
     f"另：虫害防控模块本月首次进入主巡检（4 项 / 3 家），15th & 3rd 发现活蟑螂（M 项）与捕虫器满载，"
     f"21st & 3rd 排水系统见蝇、风幕机失效，23rd & 8th 6 月虫控服务报告缺失——为新增风险面，需纳入 8 月重点。"],
    kind="warn", title="⚠ 新店首检与虫害防控为新增风险面")

# ============================================================================
# 一、门店整体表现
# ============================================================================
add_heading(doc, "一、门店整体表现", 1)
add_heading(doc, "1.1 本月概览", 2)
s11 = PACK["s1_1"]
add_para(doc, f"覆盖率（vs 6 月可比基线）：{NPRIM}/{NPRIM} = 100%（6 月覆盖 18 家，本月净增 3 家新开业门店并全部完成首检）。",
         size=10, space_after=6)
add_data_table(doc, ["最高分门店","最低分门店","S 项门店数","<80 分门店数"],
    [[f"{sn(s11['highest']['store'])} {s11['highest']['score']} 分",
      f"{sn(s11['lowest']['store'])} {s11['lowest']['score']} 分",
      f"{s11['n_stores_with_S']} 家", f"{len(s11['stores_below_80'])} 家"]],
    col_widths=[4.5,4.5,4.5,4.5], align_center_cols=[0,1,2,3])

add_heading(doc, "1.2 各门店得分明细（基于主巡检）", 2)
add_para(doc, "※ = 申诉获批后调整分数门店（调整后 ≠ 原始分，详见 §4.3）；S 项门店行以粉色底纹标注。"
              "主巡检以 QA 审计为主（18/21），US00009 48th & 3rd、US00013 Grand Central Terminal、US00021 128 W 32nd St "
              "三家新店本月无 QA 审计，主巡检为区经检查。", size=9, color_hex="666666", space_after=6)
prim_sorted = sorted(PRIMARY.items(), key=lambda kv: -kv[1]["score"])
rows12=[]; fills={}
NOM = {r["store_code"]: r["original_total_deduction"] for r in SUMMARY if r["inspection_id"] in PRIM_IIDS}
for i,(c,p_) in enumerate(prim_sorted, 1):
    star = "※" if p_["appeal_status"]=="approved" else ""
    delta = "—" if c not in JUNP else sdelta(p_["score"]-JUNP[c])
    rows12.append([i, sn(c), c, p_["score"], p_["orig"], delta, p_["type"], NOM.get(c, ""),
                   p_["S"], p_["M"], p_["G"], p_["L"], p_["inspector"], star])
    if p_["S"] > 0: fills[i-1] = PINK_FILL
add_data_table(doc,
    ["#","门店","编号","得分","原始分","环比","巡检类型","扣分","S","M","G","L","巡检员","※"],
    rows12,
    col_widths=[0.65,2.85,1.35,0.95,0.95,0.95,1.55,0.85,0.55,0.55,0.55,0.55,2.55,0.55],
    align_center_cols=[0,3,4,5,7,8,9,10,11,13], row_fills=fills, fsize=8)
add_para(doc, "注：得分 = 官方调整后分（申诉获批已反映，含 S 项 −20 分惩罚项）；原始分 = 申诉前分数；"
              "环比 = 相对 6 月主巡检得分（新店无 6 月基准记为「—」）；扣分 = 名义扣分（Σ score_config，与 S/M/G/L 计数一致，不随申诉变动），"
              "故得分 ≠ 100 − 扣分。", size=8.5, color_hex="666666", space_after=6)

add_heading(doc, "1.3 管理解读", 2)
s13 = PACK["s1_3"]
b = s13["bands"]
add_bullet(doc, f"覆盖率：{NPRIM}/{NPRIM} = 100%（6 月可比基线 18 家）。6 月遗留的 2 家 6/30 新开业门店本月完成首检，"
                f"7/16 新开业的 US00021 128 W 32nd St 亦于 7/28 完成首检，新店纳管闭环。")
add_bullet(doc, f"分数带：≥85 分 {b['≥85']} 家 / 80–84 分 {b['80-84']} 家 / <80 分 {b['<80']} 家。"
                f"两家 <80 门店均为本月新纳管门店：{sn('US00009')} 60 分、{sn('US00021')} 64 分；"
                f"6 月的两家 <80 门店（54th & 8th、52nd & Madison）本月均已回升至 84 / 90 分。")
add_bullet(doc, f"主巡检 {PM['S']} 个 S 项中，10 个为设施 Sinks and Pipes（air gap / 管道），分布在 10 家门店；"
                f"另 3 个来自新店区经检查（US00009 洗手规范、US00009 与 US00021 交叉污染）。"
                f"10 个 Sinks and Pipes S 项全部申诉获批、扣分归零，但严重度计数保留。")
imp = s13["improvers"]; dec = s13["decliners"]
add_bullet(doc, "跨月改善（vs 6 月主巡检，含申诉调整，共 {} 家）：{}。".format(
    len(imp), "；".join(f"{sn(x['store'])} {x['jun']}→{x['jul']}（+{x['delta']}）" for x in imp[:6])))
add_bullet(doc, "跨月下滑（vs 6 月主巡检，含申诉调整，共 {} 家）：{}。下滑幅度均在 3 分以内，属正常波动。".format(
    len(dec), "；".join(f"{sn(x['store'])} {x['jun']}→{x['jul']}（{x['delta']}）" for x in dec)))
add_bullet(doc, f"⚠ 需注意：上述环比均基于申诉调整后分数。若按申诉前原始分比较，18 家同口径门店均分由 6 月 {ORIG_AVG_JUN} 降至 "
                f"{round(sum(PRIMARY[c]['orig'] for c in PRIMARY if c in JUNP)/len(JUNP),1)}，改善结论不成立（详见 §4.3）。")
add_bullet(doc, f"自检主动发现 S 项：本月门店自检暴露 {len(s13['selfcheck_S_discoveries'])} 个 S 项，"
                f"覆盖过期标签 / 交叉污染 / air gap / 冷藏超温 / 异物控制等一线问题，较 6 月（10 个）进一步提升，"
                f"一线自查敏感度改善。")
add_bullet(doc, f"同店跨类型背离：本月 {len(s13['cross_type_divergence'])} 家门店出现 ≥20 分的自检 / 区经 / QA 跨类型评分背离"
                f"（6 月仅 4 家，背离面显著扩大），其中 41st & Lexington 自检 54.0 vs QA 94（差 40）最为极端（详见 §4.4）。")
add_bullet(doc, f"申诉机制：{AP['total']} 起立案（{AP['approved']} 获批 / {AP['denied']} 驳回 / {AP['pending']} 审批中），"
                f"立案量较 6 月（10 起）近乎翻倍，已决申诉获批率 100%，连续两月无驳回，需评估审批独立性与整改证据要求。")

# ============================================================================
# 二、模块风险分析
# ============================================================================
add_heading(doc, "二、模块风险分析", 1)
add_para(doc, f"本月主巡检共发现 {PM['total']} 个扣分项，分布在 8 个标准模块中（主巡检口径；全月 {FM['total']} 项另见 §3.3 / §7.x）。",
         size=10, space_after=6)
add_heading(doc, "2.1 风险分层（基于主巡检覆盖率）", 2)
MA = PACK["module_agg_primary"]
_act = sorted([m for m in MODULES_10 if MA[m]["problems"]>0], key=lambda m: -MA[m]["coverage"])
red = [m for m in _act if MA[m]["coverage"]>=50]
yel = [m for m in _act if 30<=MA[m]["coverage"]<50]
grn = [m for m in _act if MA[m]["coverage"]<30]
add_bullet(doc, "🔴 系统性风险（影响 ≥50% 门店）：" + ("、".join(f"{m}（{MA[m]['coverage']}%）" for m in red) or "无"),
           color_hex=CRITICAL_FG)
add_bullet(doc, "🟡 中等覆盖面（影响 30–49%）：" + ("、".join(f"{m}（{MA[m]['coverage']}%）" for m in yel) or "无"))
add_bullet(doc, "🟢 低覆盖面（<30%）：" + ("、".join(f"{m}（{MA[m]['coverage']}%）" for m in grn) or "无"))
add_para(doc, f"※ 覆盖率分母为已巡检 {NPRIM} 家门店（主巡检口径）。证照 / 供应链 本月主巡检无扣分，未列示。"
              f"虫害防控为本月新增进入主巡检的模块（6 月主巡检 0 项）。", size=9, color_hex="666666", space_after=6)

add_heading(doc, "2.2 模块排名总览（按扣分排序，主巡检视角）", 2)
mods_by_ded = sorted([m for m in MODULES_10 if MA[m]["problems"]>0], key=lambda m: MA[m]["deduction"])
rows22=[]
for i,m in enumerate(mods_by_ded,1):
    a=MA[m]
    risk = "🔴 系统性" + ("／含 S 项" if a["S"]>0 else "") if a["coverage"]>=50 else \
           ("🟡 中等" if a["coverage"]>=30 else ("🟢 低" + ("／含 S 项" if a["S"]>0 else "")))
    rows22.append([i,m,a["problems"],a["deduction"],f"{a['stores']}/{NPRIM}",f"{a['coverage']}%",
                   a["S"],a["M"],a["G"],a["L"],risk])
add_data_table(doc, ["#","模块","问题数","扣分","门店","覆盖率","S","M","G","L","风险"], rows22,
    col_widths=[0.7,2.6,1.2,1.0,1.2,1.2,0.6,0.6,0.6,0.6,3.5],
    align_center_cols=[0,2,3,4,5,6,7,8,9])
add_para(doc, f"※ 覆盖率分母为已巡检 {NPRIM} 家门店。Σ 主巡检模块扣分 = {PACK['s4_1']['grand_total']} 分。",
         size=9, color_hex="666666", space_after=6)

add_heading(doc, "2.3 重点模块详细分析（TOP 5，主巡检视角）", 2)
for m, info in PACK["s2_3"].items():
    a = MA[m]
    add_para(doc, f"{m}  — {a['problems']} 个扣分项，{a['deduction']} 分，影响 {a['stores']} 家门店",
             size=11, bold=True, color_hex=NAVY, space_before=6, space_after=3)
    add_para(doc, f"严重级别：S 项 {a['S']} 个、M 项 {a['M']} 个、G 项 {a['G']} 个、L 项 {a['L']} 个。",
             size=9.5, space_after=3)
    add_para(doc, "具体问题（引用原始描述，按严重度排序，最多展示 20 条；空描述已跳过）：",
             size=9.5, color_hex="666666", space_after=3)
    for f_ in info["findings"]:
        add_para(doc, f"• {sn(f_['store'])}（{f_['store']}）｜{f_['sub_item']}｜{f_['severity']} 项 {f_['deduction']} 分｜{clean(f_['description'])}",
                 size=9, indent_left=12, space_after=2)
    if info["omitted"] or info["skipped_empty"]:
        add_para(doc, f"… 另有 {info['omitted']} 条{m}主巡检发现未展示"
                      f"（含 {info['skipped_empty']} 条空描述），完整明细见原始 CSV。",
                 size=9, color_hex="666666", space_after=4)

# ============================================================================
# 三、风险等级分布
# ============================================================================
add_heading(doc, "三、风险等级分布", 1)
add_heading(doc, "3.1 整体分布（主巡检）", 2)
s31 = PACK["s3_1"]
LBL = {"S":"S 项（关键项）","M":"M 项（重要项）","G":"G 项（一般项）","L":"L 项（轻微项）"}
rows31=[]
for sv in ["S","M","G","L"]:
    d=s31[sv]
    rows31.append([LBL[sv], d["count"], f"{d['pct']}%", f"{d['sla']}内闭环",
                   "、".join(f"{k}（{v}）" for k,v in d["main_modules"].items())])
rows31.append([("合计",{"bold":True}), (PM["total"],{"bold":True}), ("100%",{"bold":True}), "—", "—"])
add_data_table(doc, ["风险等级","数量","占比","SLA 要求","主要分布模块"], rows31,
    col_widths=[3.0,1.3,1.3,2.4,9.0], align_center_cols=[1,2,3])

add_heading(doc, "3.2 S 项详情（主巡检） ── 必须立即整改", 2)
s32 = PACK["s3_2"]
add_para(doc, f"主巡检共发现 {len(s32)} 个 S 项：10 个为设施 Sinks and Pipes（air gap / 管道，全部由 QA 审计发现且全部申诉获批），"
              f"2 个为新店交叉污染，1 个为新店洗手规范。S 项已计入 §6.2 P0 行动。", size=10, space_after=6)
rows32=[]
for i,it in enumerate(s32,1):
    rows32.append([i, f"{sn(it['store'])} {it['store']}", f"{it['module']} {it['sub_item']}",
                   clean(it["description"]), it["deduction"], it["type"], it["inspector"],
                   "已获批" if it["is_appealed_finding"] and it["opp_status"]==0 else
                   ("申诉中" if it["is_appealed_finding"] else "未申诉")])
add_data_table(doc, ["#","门店","模块/子项","问题描述（原文）","扣分","巡检类型","巡检员","申诉"], rows32,
    col_widths=[0.6,2.9,2.6,5.6,0.8,1.3,2.5,1.1],
    align_center_cols=[0,4,5,7], fsize=8)

add_heading(doc, "3.3 全月 S 项汇总（含自检与重复巡检）", 2)
add_para(doc, f"全月共 {FM['S']} 个 S 项（6 月 23 个，↑{FM['S']-23}），按子项汇总：", size=10, space_after=4)
rows33=[[f"{x['sub_item']}［{x['module']}］", x["S_count"], x["stores"], clean(x["typical"]) or "—"]
        for x in PACK["s3_3"]["by_sub_item"]]
add_data_table(doc, ["子项［模块］","S 项数","门店数","典型问题（截取）"], rows33,
    col_widths=[4.6,1.3,1.3,10.8], align_center_cols=[1,2])
add_para(doc, f"Sinks and Pipes（{PACK['s3_3']['by_sub_item'][0]['S_count']} 起 / "
              f"{PACK['s3_3']['by_sub_item'][0]['stores']} 家）为全月最大系统性 S 项，已计入 §6.2 P0。"
              f"全月 S/M/G/L 合计：S {FM['S']} / M {FM['M']} / G {FM['G']} / L {FM['L']} = {FM['total']}；"
              f"主巡检 vs 全月：S 项 {PM['S']}/{FM['S']}、全部发现 {PM['total']}/{FM['total']}。",
         size=9, color_hex="666666", space_after=6)

add_heading(doc, "3.4 M 项详情（主巡检） ── 7 天内闭环", 2)
add_para(doc, f"主巡检共发现 {PM['M']} 个 M 项，其中 5 个为过期标签缺失（温控有效期），呈明显集中特征。", size=10, space_after=4)
rows34=[[i, f"{sn(x['store'])} {x['store']}", f"{x['module']} {x['sub_item']}", clean(x["description"]), x["deduction"]]
        for i,x in enumerate(PACK["s3_4"],1)]
add_data_table(doc, ["#","门店","模块/子项","问题描述（原文）","扣分"], rows34,
    col_widths=[0.7,3.4,3.4,9.5,1.4], align_center_cols=[0,4])

add_heading(doc, "3.5 G 项 / L 项 分布（主巡检）", 2)
g_by = PACK["s3_5"]["G_by_module"]; l_by = PACK["s3_5"]["L_by_module"]
add_para(doc, f"G 项（一般项）共 {PM['G']} 个，主要集中模块：", size=10, space_after=3)
for m,v in sorted(g_by.items(), key=lambda x:-x[1]):
    add_bullet(doc, f"{m}：{v} 个", size=9.5)
add_para(doc, f"L 项（轻微项）共 {PM['L']} 个，主要集中模块：", size=10, space_before=4, space_after=3)
for m,v in sorted(l_by.items(), key=lambda x:-x[1]):
    add_bullet(doc, f"{m}：{v} 个", size=9.5)

# ---------- 3.6 NEW: cross-month S tracking ledger ----------
add_heading(doc, "3.6 Sinks and Pipes 跨月追踪台账（6 月 → 7 月）", 2)
add_para(doc, "本节为 6 月报告 §6.2 P0 与 §7.7 要求建立的跨月 S 项台账，逐店记录 air gap / 管道类 S 项的月度复现状态。",
         size=10, space_after=4)
jun_set = set(REPEAT["jun_stores"]); jul_set = set(REPEAT["jul_stores"])
rows36=[]
for c in sorted(jun_set | jul_set):
    inj = "✓" if c in jun_set else "—"
    inl = "✓" if c in jul_set else "—"
    if c in jun_set and c in jul_set:   st, col = "⚠ 连续两月复现（未根治）", CRITICAL_FG
    elif c in jul_set:                  st, col = "🆕 7 月新增", WARN_FG
    else:                               st, col = "✅ 7 月已消除", HEALTHY_FG
    rows36.append([f"{sn(c)} {c}", inj, inl, (st, {"color": col, "bold": True})])
add_data_table(doc, ["门店","6 月 S 项","7 月 S 项","跨月状态"], rows36,
    col_widths=[5.0,2.5,2.5,8.4], align_center_cols=[1,2,3])
add_callout(doc,
    [f"6 月 {len(jun_set)} 家 → 7 月 {len(jul_set)} 家；连续两月复现 {len(REPEAT['repeat'])} 家、7 月新增 {len(REPEAT['new'])} 家"
     f"（{'、'.join(sn(c) for c in REPEAT['new'])}）、7 月已消除 {len(REPEAT['cleared'])} 家（{'、'.join(sn(c) for c in REPEAT['cleared'])}）。",
     f"结论：6 月 P0（48 小时闭环 + BD 整改清单）执行未见成效——{len(REPEAT['repeat'])}/{len(jun_set)} 的 6 月问题门店在 7 月再次命中同类 S 项，"
     f"且全月 S 项数由 13 起升至 {PACK['s3_3']['by_sub_item'][0]['S_count']} 起。建议将其由「整改工单」升级为「BD 专项工程」，"
     f"按门店逐一出具管道 / air gap 改造方案与完工验收单，QA 复检以实物验收为准、不接受仅申诉调分。"],
    kind="critical", title="🔴 台账结论")

# ============================================================================
# 四、模块与门店关联分析
# ============================================================================
add_heading(doc, "四、模块与门店关联分析", 1)
add_heading(doc, "4.1 门店 × 模块扣分矩阵（主巡检）", 2)
add_para(doc, f"※ 矩阵覆盖 {NPRIM} 家已巡检门店（主巡检口径，含 {N_APPEAL_ADJ} 起申诉获批后调整）；"
              f"行合计 = 该店主巡检总扣分（名义）。清洁卫生影响全部 {NPRIM} 家（100%），为最广覆盖模块。",
         size=9, color_hex="666666", space_after=4)
MX = PACK["s4_1"]["matrix"]; STOT = PACK["s4_1"]["store_total"]
order = sorted(PRIMARY.keys(), key=lambda c: STOT[c])
rows41=[]
for c in order:
    rows41.append([f"{sn(c)}（{c}）"] + [(MX.get(c,{}).get(m) or "") for m in MODULES_10] + [STOT[c]])
tot = ["合计"] + [ (sum(MX.get(c,{}).get(m,0) for c in PRIMARY) or "") for m in MODULES_10 ] + [PACK["s4_1"]["grand_total"]]
rows41.append([(x, {"bold":True}) for x in tot])
add_data_table(doc, ["门店"] + [MOD_SHORT[m] for m in MODULES_10] + ["合计"], rows41,
    col_widths=[4.3]+[1.05]*10+[1.1], align_center_cols=list(range(1,12)), fsize=8)
add_para(doc, f"100% 门店命中的模块：{'、'.join(PACK['s4_1']['full_coverage_modules'])}。",
         size=9, color_hex="666666", space_after=6)

add_heading(doc, "4.2 低分门店归因（<80 分）", 2)
for x in PACK["s4_2"]:
    c = x["store"]
    add_para(doc, f"{sn(c)}（{c}）　主巡检 {x['score']} 分（{x['type']}，{x['date']}，巡检员 {x['inspector']}）",
             size=11, bold=True, color_hex=NAVY, space_before=6, space_after=3)
    add_para(doc, "模块扣分构成：" + "、".join(f"{m}（{v}）" for m,v in sorted(x["module_breakdown"].items(), key=lambda kv: kv[1])),
             size=9.5, space_after=2)
    if x["S_items"]:
        add_para(doc, "S 项：" + "；".join(f"{s['sub_item']}「{clean(s['desc'])}」" for s in x["S_items"]),
                 size=9.5, color_hex=CRITICAL_FG, space_after=2)
    add_para(doc, f"门店状态：{'本月新纳管门店（开业 ' + x['open_date'] + '），无 6 月主巡检基准，本次为首检' if x['is_new'] else '6 月主巡检基准 ' + str(x['prior_baseline']) + ' 分'}。"
                  f"{'首检由区经完成，本月未纳入 QA 审计覆盖，建议 8 月安排 QA 首检并配套新店基础规范复训。' if x['is_new'] else ''}",
             size=9.5, space_after=4)

add_heading(doc, "4.3 申诉调整门店分析（全口径：获批 / 驳回 / 审批中）", 2)
add_para(doc, f"本月共 {AP['total']} 起申诉立案：{AP['approved']} 起获批、{AP['denied']} 起驳回、{AP['pending']} 起审批中，"
              f"涉及 {PACK['s4_3']['appealed_findings']} 条 finding。", size=10, space_after=4)
resmap={"approved":"获批 ※","denied":"驳回","pending":"审批中"}
rows43=[]
for r in PACK["s4_3"]["rows"]:
    delta = r["adj"]-r["orig"]
    rows43.append([f"{sn(r['store'])} {r['store']}", r["type"], resmap[r["result"]], r["date"],
                   f"{r['orig']}→{r['adj']}", (sdelta(delta), {"bold": delta>=20, "color": CRITICAL_FG if delta>=20 else TEXT}),
                   "主巡检" if r["is_primary"] else "非主巡检", r["inspector"]])
add_data_table(doc, ["门店","巡检类型","申诉结果","日期","分数变动","变动幅度","是否主巡检","巡检员"], rows43,
    col_widths=[3.3,1.5,1.3,1.9,1.7,1.3,1.5,3.9], align_center_cols=[1,2,3,4,5,6], fsize=8)
add_callout(doc,
    [f"获批 {AP['approved']} 起中，10 起为 QA 审计发现的 air gap / 管道 S 项撤销，单笔提分普遍达 +25 至 +31 分"
     f"（S 项撤销同时解除 −20 惩罚项）；主巡检层面 {N_APPEAL_ADJ} 家门店因此平均提分 {AVG_UPLIFT} 分。",
     f"连续两月已决申诉 0 驳回（6 月 7 获批 / 7 月 {AP['approved']} 获批），累计 23 获批 0 驳回。"
     f"若申诉理由为「整改后复核」，则应以复检实物证据为准并同步关闭整改工单；本月 §3.6 台账显示 "
     f"{len(REPEAT['repeat'])} 家门店问题仍在，提示部分获批申诉未对应实际整改完成。",
     f"建议：① 申诉审批引入第二审核人（避免与原巡检员 / 单一审批人闭环）；② S 项申诉强制附整改前后照片与完工验收单；"
     f"③ 申诉获批后 14 天内安排 QA 实物复检；④ {AP['pending']} 起审批中申诉（16th & 6th、Grand Central Terminal、40th & 10th 区经检查）跟踪 SLA。"],
    kind="warn", title="⚠ 申诉治理建议")

add_heading(doc, "4.4 同店同期评分背离案例", 2)
add_para(doc, f"本月 {len(PACK['s4_4'])} 家门店出现 ≥20 分的跨类型评分背离（自检 / QA审计 / 区经检查之间），"
              f"较 6 月（4 家）显著扩大：", size=10, space_after=4)
LAB={"门店自检_avg":"自检","QA审计":"QA","区经检查":"区经"}
rows44=[]
for x in PACK["s4_4"]:
    lo_k, lo_v = list(x["lower"].items())[0]; hi_k, hi_v = list(x["higher"].items())[0]
    if hi_k=="门店自检_avg":   note="自检显著宽松，正式巡检暴露更多问题，需校准自检尺度"
    elif lo_k=="门店自检_avg": note="自检偏严于正式巡检，一线自查更保守"
    else:                      note="正式巡检间尺度差异（区经严于 QA），需尺度对齐"
    rows44.append([f"{sn(x['store'])} {x['store']}", f"{LAB[lo_k]} {lo_v}", f"{LAB[hi_k]} {hi_v}",
                   (x["gap"], {"bold": x["gap"]>=30, "color": CRITICAL_FG if x["gap"]>=30 else TEXT}), note])
add_data_table(doc, ["门店","较低类型（分）","较高类型（分）","差值","解读"], rows44,
    col_widths=[3.4,2.2,2.2,1.2,9.4], align_center_cols=[1,2,3], fsize=8.5)
add_para(doc, "41st & Lexington（自检 54.0 vs QA 94，差 40）背离最大，根因为该店自检员 Afsana Gu 评分极端偏严"
              "（3 次自检 29 / 37 / 58，月均 41.3，为全月唯一「偏严」巡检员，详见 §7.4）。"
              "28th & 6th（区经 57 vs QA 94，差 37）与 154 Bleecker（区经 62 vs QA 95，差 33）则反映区经与 QA 之间的系统性尺度差异——"
              "本月区经均分 78.9、QA 均分 92.5，差距 13.6 分为历史最大，需优先开展区经 / QA 尺度对齐校准。",
         size=9.5, space_after=6)

add_heading(doc, "4.5 模块覆盖面分析（主巡检视角）", 2)
rows45=[[x["module"], f"{x['stores']}/{NPRIM}", f"{x['coverage']}%", x["deduction"],
         "🔴 系统性" + ("／含 S 项" if MA[x["module"]]["S"]>0 else "") if x["coverage"]>=50
         else ("🟡 中等" if x["coverage"]>=30 else "🟢 低" + ("／含 S 项" if MA[x["module"]]["S"]>0 else ""))]
        for x in PACK["s4_5"]]
add_data_table(doc, ["模块","影响门店","覆盖率","扣分","风险标记"], rows45,
    col_widths=[3.4,2.2,2.2,2.0,8.6], align_center_cols=[1,2,3])
add_para(doc, f"※ 覆盖率分母为已巡检 {NPRIM} 家门店。清洁卫生 100% 覆盖为最广系统性模块（{MA['清洁卫生']['deduction']} 分），"
              f"设施（{MA['设施']['coverage']}%）次之且含 10 个 Sinks and Pipes S 项，为整改优先方向。",
         size=9, color_hex="666666", space_after=6)

# ============================================================================
# 五、整改归因与效率
# ============================================================================
add_heading(doc, "五、整改归因与效率", 1)
add_heading(doc, "5.1 关键词归因（全月发现）", 2)
add_para(doc, "⚠ 以下归因基于问题描述关键词自动匹配，仅供参考。实际归因需以整改工单系统数据为准。",
         size=9, color_hex=WARN_FG, space_after=4)
s51 = PACK["s5_1"]; A = s51["attribution"]; TOTA = s51["total"]
typdesc = {"门店":"日常清洁、消毒、PPM、标签、储存卫生、个人卫生",
           "机修+营建":"sinks and pipes、air gap、油脂阱、灯具、门 / 墙 / 地面",
           "供应链+行政":"license / certificate / 文件记录 / 洗手标识",
           "未知":"描述缺失或少于 10 字符（多来自门店自检未填明细）"}
rows51=[[k, A.get(k,0), f"{A.get(k,0)/TOTA*100:.1f}%", typdesc[k]] for k in ["门店","机修+营建","供应链+行政","未知"]]
add_data_table(doc, ["归因类别","数量","占比","典型问题"], rows51,
    col_widths=[2.6,1.5,1.5,12.8], align_center_cols=[1,2])
add_para(doc, f"全月空 / 短描述占比 {s51['empty_pct']}%（{s51['empty_short']}/{TOTA}，其中门店自检 {s51['empty_short_selfcheck']} 条），"
              f"较 6 月 27.1% 有所改善但仍偏高，是自检描述质量的持续改进重点；"
              f"机修 + 营建类（{A.get('机修+营建',0)/TOTA*100:.1f}%）集中于 Sinks and Pipes，是 BD 整改重点。",
         size=9.5, space_after=6)

add_heading(doc, "5.2 SLA 整改时限标准", 2)
add_data_table(doc, ["风险等级","整改时限","要求"], [
    ["S 项（关键项）","2 天","发现后 2 天内完成整改并验证"],
    ["M 项（重要项）","7 天","发现后 7 天内完成整改并验证"],
    ["G 项（一般项）","14 天","发现后 14 天内完成整改并验证"],
    ["L 项（轻微项）","14 天","发现后 14 天内完成整改并验证"]],
    col_widths=[3.0,2.5,12.9], align_center_cols=[1])

add_heading(doc, "5.3 建议整改闭环流程", 2)
for s in ["巡检发现问题 → empapp 自动生成整改工单",
          "根据问题类型自动分配责任方（门店 / 机修 / 营建 / 供应链）",
          "责任方在 SLA 时限内完成整改",
          "QA 复核验证整改效果（S 项须留存整改前后照片与完工验收单）",
          "系统记录闭环时间，计算 SLA 达标率",
          "【新增建议】申诉获批与整改工单双向绑定：申诉以「整改后复核」为由获批的，须同步关闭对应工单并留存验收证据"]:
    add_bullet(doc, s)

# ============================================================================
# 六、建议与下一步行动
# ============================================================================
add_heading(doc, "六、建议与下一步行动", 1)
add_heading(doc, "6.1 本月关键发现", 2)
key_findings = [
    ("🟢", f"巡检体量与覆盖创新高：7 月 {TC['total']} 次巡检（自检 {TC['门店自检']} / QA {TC['QA审计']} / 区经 {TC['区经检查']}）较 6 月 85 次 +9 次，"
           f"为有记录以来最高；{NPRIM} 家在营门店 100% 覆盖，3 家新开业门店全部完成首检，新店纳管闭环。"),
    ("🔴", f"分数改善主要来自申诉调整而非现场改善：官方均分 {PRIOR_AVG}→{PRIM_AVG}（+{round(PRIM_AVG-PRIOR_AVG,1)}），"
           f"但申诉前原始均分 {ORIG_AVG_JUN}→{ORIG_AVG_JUL}（{round(ORIG_AVG_JUL-ORIG_AVG_JUN,1)}）；"
           f"申诉获批门店 7→{N_APPEAL_ADJ} 家、平均提分 {AVG_UPLIFT} 分；全月发现项 498→{FM['total']}（+{FM['total']-498}）。"),
    ("🔴", f"air gap / Sinks and Pipes 连续三月未根治且扩大：全月 13→{PACK['s3_3']['by_sub_item'][0]['S_count']} 起、10→"
           f"{PACK['s3_3']['by_sub_item'][0]['stores']} 家；{len(REPEAT['repeat'])} 家连续两月复现。"
           f"主巡检 10 个此类 S 项全部申诉获批、扣分归零，分数已修复但物理问题仍在（§3.6 台账）。"),
    ("⚠", f"新店首检两家低于 80 分：48th & 3rd 60 分（2 个 S 项）、128 W 32nd St 64 分（1 个 S 项）；"
          f"3 家新店本月均未获 QA 审计覆盖，首检全部由区经完成，新店开业前基础规范验收存在缺口。"),
    ("⚠", f"虫害防控成为新增风险面：本月首次进入主巡检（4 项 / 3 家），15th & 3rd 发现活蟑螂（M 项）与捕虫器满载，"
          f"21st & 3rd 排水见蝇 + 风幕机失效，23rd & 8th 6 月虫控服务报告缺失。"),
    ("⚠", f"巡检尺度分化加剧：区经均分 84.3→78.9（−5.4）、QA 均分 88.2→92.5（+4.3），两者差距 13.6 分为历史最大；"
          f"≥20 分跨类型背离门店由 4 家增至 {len(PACK['s4_4'])} 家。Afsana Gu 自检月均 41.3 分为全月唯一「偏严」，"
          f"Laurel Vorhies 同店摆动 35 分（59→94→93）。"),
    ("🟡", f"申诉治理需加强：{AP['total']} 起立案较 6 月近乎翻倍，已决申诉获批率 100%、连续两月 0 驳回（累计 23 获批 / 0 驳回）；"
           f"建议引入第二审核人、S 项申诉强制附整改证据、获批后 14 天内实物复检。"),
    ("🟡", f"QA 单点依赖延续第四个月：Eamonn Caballar 独立完成全部 {TC['QA审计']} 次 QA 审计，QA 覆盖 18/21（3 家新店未覆盖）；"
           f"Yu Jiang HR 流程仍未关闭（7 月 0 次）。区经亦为单人执行（Jung Han Liang 21 次）。"),
    ("✅", f"低分门店改善明显：6 月两家 <80 门店本月均回升——52nd & Madison 64→90（+26，全月最大改善）、"
           f"54th & 8th 71→84（+13，结束连续两月 <80）；自检描述完整度由 72.9% 提升至 {100-s51['empty_pct']:.1f}%。"),
]
for icon, txt in key_findings:
    col = CRITICAL_FG if icon=="🔴" else (WARN_FG if icon=="⚠" else (HEALTHY_FG if icon in ("🟢","✅") else TEXT))
    add_bullet(doc, f"{icon} {txt}", color_hex=col)

add_heading(doc, "6.2 优先行动项", 2)
actions = [
 ["P0","紧急",
  f"air gap / 管道治理升级为 BD 专项工程：处理本月 10 个主巡检 Sinks and Pipes S 项，48 小时内出具处置方案；"
  f"对 §3.6 台账中 {len(REPEAT['repeat'])} 家连续两月复现门店（{'、'.join(sn(c) for c in REPEAT['repeat'][:5])} 等）"
  f"逐店出具管道 / air gap 改造方案 + 完工验收单，QA 复检以实物验收为准",
  "BD + 门店 + QA","48 小时出方案 / 本月内完工"],
 ["P0【接续】","紧急",
  f"申诉与整改解耦核查：对 7 月 {AP['approved']} 起获批申诉（尤其 10 起 air gap S 项撤销）逐笔核对整改工单闭环证据，"
  f"未留存整改证据的重新计入待整改清单；建立「申诉获批 ≠ 整改完成」的双轨台账",
  "QA 部门 + BD","2 周内"],
 ["P1","紧急",
  f"处理本月 {PM['M']} 个主巡检 M 项：5 起过期标签缺失（100 Maiden Ln、54th & 8th、102 Fulton、16th & 6th、128 W 32nd St）、"
  f"54th & 8th 消毒液浓度不达标、15th & 3rd 活蟑螂与捕虫器满载",
  "门店 + QA + 虫控供应商","7 天"],
 ["P2","高",
  f"新店专项：48th & 3rd（60）、128 W 32nd St（64）两家新店 8 月安排 QA 首检 + 基础规范复训；"
  f"建立「开业前 QA 验收 + 开业 30 天内强制 QA 首检」制度，避免新店裸奔",
  "QA 部门 + 运营部","2 周内"],
 ["P3","高",
  f"区经 / QA 尺度对齐校准：本月区经均分 78.9 vs QA 92.5，差距 13.6 分为历史最大，{len(PACK['s4_4'])} 家门店出现 ≥20 分跨类型背离；"
  f"组织区经与 QA 联合走店、统一评分基准与取证标准",
  "QA 部门 + 运营部","本月内"],
 ["P4","高",
  f"申诉审批机制改造：① 引入第二审核人；② S 项申诉强制附整改前后照片与完工验收单；③ 获批后 14 天内 QA 实物复检；"
  f"④ 月度看板同时披露原始分与调整分",
  "QA 部门","本月内"],
 ["P5","中",
  f"虫害防控专项：15th & 3rd 立即启动灭治并复检；21st & 3rd 风幕机报修、排水系统处理；"
  f"23rd & 8th 补齐 6 月虫控服务报告；全店排查虫控服务记录完整性",
  "门店 + 行政 + 虫控供应商","2 周内"],
 ["P6","中",
  f"巡检员评分尺度标准化培训：Afsana Gu（41st & Lexington，月均 41.3、摆动 29，全月唯一偏严，需即时干预）、"
  f"Laurel Vorhies（16th & 6th，摆动 35）、Yaqing Zuo（33rd & 10th，摆动 24）优先校准",
  "QA 部门","本月内"],
 ["P7","中",
  f"empapp 系统优化：自检问题描述必填校验（全月 {s51['empty_pct']}% 空 / 短描述，其中自检 {s51['empty_short_selfcheck']} 条）；"
  f"申诉立案与整改工单双向同步",
  "基础设施部","2 周内"],
 ["P8","中",
  f"评估 QA / 区经团队冗余度：Eamonn Caballar 连续四月独立完成全部 QA 审计、Jung Han Liang 单人完成 21 次区经检查；"
  f"推进 Yu Jiang HR 流程关闭与第二位 QA Manager 配置",
  "QA 部门 + HR","本季度"],
]
add_data_table(doc, ["优先级","紧急度","行动项","责任方","时限"], actions,
    col_widths=[1.6,1.1,9.9,2.5,2.3], align_center_cols=[0,1], fsize=8.5)

add_heading(doc, "6.3 模块改善建议（TOP 5）", 2)
suggestions = [
 (f"清洁卫生（影响 {MA['清洁卫生']['stores']} 家，主巡检 {MA['清洁卫生']['deduction']} 分）",
  "① 冰箱顶 / 冰机 / 咖啡研磨机 / 搅拌机每日深度清洁（dust、coffee grounds、matcha staining、sanitizer residue 为最高频描述）；"
  "② 消毒液浓度（PPM）每日校准并记录（54th & 8th 读数不达标）；③ 清洁任务须在 15 分钟窗口内执行完毕（多店出现 tasks out of grace period）；"
  "④ 器具清洗后不得置于积水中。"),
 (f"设施（影响 {MA['设施']['stores']} 家，主巡检 {MA['设施']['deduction']} 分，含 10 个 S 项）",
  "① air gap / Sinks and Pipes 升级为 BD 专项工程，逐店出改造方案与完工验收单（§3.6 台账 9 家连续两月复现）；"
  "② 管道与滤芯间距、包裹层脱落、油脂阱排水为本月高频子问题，纳入月度必检；"
  "③ 门体缝隙 / 店招外立面清洁纳入营建例行维护。"),
 (f"温控有效期（影响 {MA['温控有效期']['stores']} 家，主巡检 {MA['温控有效期']['deduction']} 分，全部为 M 项）",
  "① 开封后过期标签管理纳入每日开店清单——本月 5 起全部为标签缺失，集中度极高；"
  "② 糖浆瓶 / 冷藏物料为重灾区，建议标签打印机点位前置到吧台；③ FIFO 每日检查、过期产品零容忍。"),
 (f"过程控制（影响 {MA['过程控制']['stores']} 家，主巡检 {MA['过程控制']['deduction']} 分，含 2 个 S 项）",
  "① 交叉污染防控纳入全员复训（新店 48th & 3rd 变质牛奶未处置、128 W 32nd St 奇亚籽混入奶昔粉为 S 项）；"
  "② 物料储位规范：多店出现 lock box 未上锁且可接触；③ 容器加盖、粉料受潮结块每日巡查。"),
 (f"虫害防控（影响 {MA['虫害防控']['stores']} 家，主巡检 {MA['虫害防控']['deduction']} 分）",
  "① 15th & 3rd 活蟑螂 + 捕虫器满载立即启动灭治并复检；② 21st & 3rd 风幕机失效立即报修、排水系统见蝇处理；"
  "③ 全店排查虫控服务报告归档完整性（23rd & 8th 6 月报告缺失）；④ 虫控服务改为月度双向确认。"),
]
for title, body in suggestions:
    add_para(doc, title, size=10.5, bold=True, color_hex=NAVY, space_before=5, space_after=2)
    add_para(doc, body, size=9.5, space_after=3)

# ============================================================================
# 七、三类巡检体系分析
# ============================================================================
add_heading(doc, "七、三类巡检体系分析", 1)
add_para(doc, "瑞幸北美门店食品安全采用三类巡检体系：门店自检（Store Self-Check）、QA审计（QA Audit）、区经检查（Area Check）。"
              "本月三类齐全且全部活跃，巡检总量创新高。", size=10, space_after=6)

add_heading(doc, "7.1 三类巡检概况", 2)
s71 = PACK["s7_1"]
add_para(doc, f"QA审计平均分最高（{s71['QA审计']['avg']}，含申诉获批调整；申诉前原始均分仅 77.9）且 S 项集中度高"
              f"（{s71['QA审计']['S']} 个 S 项全部为 air gap / 管道）；区经检查均分最低（{s71['区经检查']['avg']}）但覆盖面最广"
              f"（{NPRIM}/{NPRIM}）且尺度最严；门店自检（{s71['门店自检']['avg']}）高频暴露日常问题。",
         size=10, space_after=4)
rows71=[]
tempo={"门店自检":"全月高频（日常暴露）","QA审计":"月中集中（专业定级）","区经检查":"全月 + 月末复核"}
for t in ["门店自检","QA审计","区经检查"]:
    d=s71[t]
    rows71.append([t, d["count"], f"{d['stores']}/{NPRIM}", d["inspectors"], d["avg"], d["S"], d["M"],
                   f"{d['first'][5:]}~{d['last'][5:]}", tempo[t]])
add_data_table(doc, ["巡检类型","次数","覆盖门店","巡检员数","平均分","S 项","M 项","日期区间","节奏"], rows71,
    col_widths=[2.0,1.0,1.5,1.3,1.2,0.9,0.9,2.2,5.4], align_center_cols=[1,2,3,4,5,6,7])

add_heading(doc, "7.2 同店三类巡检对比", 2)
add_para(doc, "下表列示各门店三类巡检平均分（「—」表示本月无该类型巡检）：", size=10, space_after=4)
rows72=[]
for x in PACK["s7_2"]:
    d = x["QA_minus_self"]
    rows72.append([f"{sn(x['store'])} {x['store']}",
                   x["selfcheck_avg"] if x["selfcheck_avg"] is not None else "—",
                   x["QA"] if x["QA"] is not None else "—",
                   x["area"] if x["area"] is not None else "—",
                   (sdelta(d, 1) if d != "" else "—",
                    {"color": CRITICAL_FG if isinstance(d,(int,float)) and abs(d)>=20 else TEXT,
                     "bold": isinstance(d,(int,float)) and abs(d)>=20})])
add_data_table(doc, ["门店","自检均分","QA审计","区经检查","QA−自检差"], rows72,
    col_widths=[5.0,2.8,2.8,2.8,2.9], align_center_cols=[1,2,3,4], fsize=8.5)
add_para(doc, "「QA−自检差」为正表示 QA 审计分高于自检（自检偏严或整改后改善），为负表示自检高于 QA（自检偏宽，需关注）。"
              "41st & Lexington（+40.0）、29th & 3rd（+25.5）、21st & 3rd（+23.0）自检显著偏严；"
              "100 Maiden Ln（−8.0）自检偏宽。需注意 QA 分含申诉获批调整成分，跨类型比较应同时参考原始分。",
         size=9.5, space_after=6)

add_heading(doc, "7.3 自检评分一致性分析", 2)
add_para(doc, "针对同一巡检员对同一门店的多次自检，评估其评分稳定性（摆动 = 最高分 − 最低分），列示摆动 ≥10 分者：",
         size=10, space_after=4)
rows73=[[x["inspector"], f"{sn(x['store'])} {x['store']}", x["n"], " → ".join(str(s) for s in x["scores"]),
         (x["swing"], {"bold": x["swing"]>=25, "color": CRITICAL_FG if x["swing"]>=25 else TEXT})]
        for x in PACK["s7_3"] if x["swing"]>=10]
add_data_table(doc, ["巡检员","门店","次数","历次得分","摆动"], rows73,
    col_widths=[3.6,4.2,1.2,6.0,1.4], align_center_cols=[2,4], fsize=8.5)
add_para(doc, "Laurel Vorhies（16th & 6th，摆动 35：59→94→93）与 Afsana Gu（41st & Lexington，摆动 29：29→37→58）摆动最大；"
              "Afsana Gu 同时为全月唯一「偏严」巡检员（月均 41.3），其评分尺度与其余自检员差距过大，"
              "直接造成该店 40 分的跨类型背离。Yaqing Zuo（33rd & 10th，摆动 24：68→81→92）呈单向上升，"
              "可能反映真实整改改善。已纳入 §6.2 P6 标准化培训。", size=9.5, space_after=6)

add_heading(doc, "7.4 巡检员严格度分析（≥2 次巡检）", 2)
add_para(doc, "仅纳入本月完成 ≥2 次巡检的巡检员，避免单次样本误判：", size=10, space_after=4)
rows74=[]
for x in PACK["s7_4"]:
    col = CRITICAL_FG if x["scale"].startswith("偏严") else (WARN_FG if x["scale"].startswith("偏宽") else TEXT)
    lab = {"偏严(<70)":"偏严","偏宽(>92)":"偏宽"}.get(x["scale"], x["scale"])
    rows74.append([x["inspector"], x["role"], x["type"], x["n"],
                   (f"{x['avg']}（{lab}）", {"color": col, "bold": col!=TEXT})])
add_data_table(doc, ["巡检员","角色","巡检类型","次数","均分（尺度）"], rows74,
    col_widths=[4.0,4.2,2.2,1.2,4.8], align_center_cols=[3,4], fsize=8.5)
add_para(doc, f"Eamonn Caballar（QA，{TC['QA审计']} 次，92.5）本月被标记为「偏宽」——但其原始（申诉前）均分仅 77.9，"
              f"该标记实为申诉获批调整的结果，而非其现场尺度放松，解读时须注意。"
              f"Jung Han Liang（区经，{TC['区经检查']} 次，78.9）尺度正常但较 6 月（84.3）明显收紧。"
              f"自检员整体处于 70–92 正常带，仅 Afsana Gu（41.3）严重偏离，为唯一需要即时干预的个案。",
         size=9.5, space_after=6)

add_heading(doc, "7.5 巡检覆盖率趋势", 2)
rows75=[]
notes={"2026-01":"体系初建","2026-02":"断流低谷（区经 0）","2026-03":"自检为主、QA 仅 1 次",
       "2026-04":"危机后恢复，三类齐全","2026-05":"历史新高，三类齐全","2026-06":"高位运行，三类齐全",
       "2026-07":"创历史新高，三类齐全，覆盖 21 家"}
for x in PACK["s7_5"]:
    mlabel = f"{int(x['month'][5:])} 月"
    rows75.append([mlabel, x["self"], x["QA"], x["area"], x["total"],
                   x["self_avg"] or "—", x["QA_avg"] or "—", x["area_avg"] or "—", notes[x["month"]]])
add_data_table(doc, ["月份","门店自检","QA审计","区经检查","合计","自检均分","QA均分","区经均分","说明"], rows75,
    col_widths=[1.3,1.5,1.4,1.5,1.1,1.5,1.4,1.5,6.2], align_center_cols=[0,1,2,3,4,5,6,7], fsize=8.5)
add_para(doc, f"7 月 {TC['total']} 次较 6 月 85 次 +9 次，为有记录以来最高，且三类巡检齐全、全部活跃。"
              f"三类平均分（7 月）：门店自检 {s71['门店自检']['avg']}、QA审计 {s71['QA审计']['avg']}、区经检查 {s71['区经检查']['avg']}，"
              f"较 6 月（82.8 / 88.2 / 84.3）呈现「QA 上升、自检与区经下降」的分化格局——"
              f"QA 上升主要由申诉获批驱动，自检与区经下降则反映一线与区域层面尺度收紧、暴露问题增多。",
         size=9.5, space_after=6)

add_heading(doc, "7.6 三类巡检发现差异", 2)
s76 = {x["type"]: x for x in PACK["s7_6"]}
add_bullet(doc, f"门店自检（{TC['门店自检']} 次）：S 项 {s76['门店自检']['S']} / M 项 {s76['门店自检']['M']}，"
                f"高频暴露过期标签 / 交叉污染 / 冷藏超温 / 异物控制等一线问题；价值在于发现频次高、贴近现场，"
                f"但描述完整度仍偏低（全月 {s51['empty_pct']}% 空 / 短描述，其中自检 {s51['empty_short_selfcheck']} 条）。")
add_bullet(doc, f"QA审计（{TC['QA审计']} 次）：S 项 {s76['QA审计']['S']} / M 项 {s76['QA审计']['M']}，专业度最高，"
                f"S 项全部集中于 Sinks and Pipes（air gap / 管道）等结构性、设施性风险；是 BD 整改清单的主要来源，"
                f"但本月其 S 项全部经申诉撤销，需与整改闭环绑定。")
add_bullet(doc, f"区经检查（{TC['区经检查']} 次）：S 项 {s76['区经检查']['S']} / M 项 {s76['区经检查']['M']}，覆盖面最广（{NPRIM} 家），"
                f"本月尺度最严（均分 78.9），在新店（48th & 3rd、128 W 32nd St）发现全部 3 个非 air gap 类主巡检 S 项，"
                f"是自检与 QA 之间的有效补充。")
add_para(doc, "三类互补：自检高频暴露、QA 专业定级、区经全覆盖复核。本月三类同时指向 Sinks and Pipes 系统性短板，"
              "交叉验证了整改优先级；同时区经在新店的严格发现，弥补了 QA 未覆盖新店的空白。", size=9.5, space_after=6)

add_heading(doc, "7.7 后续观察重点", 2)
watch = [
 f"air gap 跨月台账（§3.6）：{len(REPEAT['repeat'])} 家连续两月复现门店须在 8 月前完成实物整改验收；"
 f"7 月新增 2 家（28th & 6th、40th & 10th）纳入台账跟踪；已消除 1 家（16th & 6th）保持复检。",
 f"申诉与整改解耦：连续两月 23 获批 / 0 驳回，须核查获批申诉对应的整改工单闭环证据；"
 f"{AP['pending']} 起审批中申诉（16th & 6th、Grand Central Terminal、40th & 10th）跟踪审批 SLA。",
 f"新店质量爬坡：48th & 3rd（60）、128 W 32nd St（64）8 月 QA 首检与复训效果；"
 f"Grand Central Terminal（96）作为新店标杆经验沉淀。",
 f"虫害防控：15th & 3rd 灭治复检结果；全店虫控服务报告归档完整性排查；8 月观察该模块是否延续上升。",
 f"区经 / QA 尺度对齐：本月差距 13.6 分为历史最大，8 月观察联合走店校准后的收敛情况；"
 f"≥20 分背离门店数（本月 {len(PACK['s4_4'])} 家）应作为尺度一致性的月度指标跟踪。",
 f"巡检员标准化：Afsana Gu（月均 41.3）为即时干预对象；Laurel Vorhies（摆动 35）等纳入尺度校准。",
 f"自检描述质量：空 / 短描述由 27.1% 降至 {s51['empty_pct']}%，继续推进 empapp 必填校验以进一步压降。",
 f"团队冗余度：QA 单点依赖已延续四个月、区经单人执行两个月，须评估断流风险与第二人选配置时点。",
]
for w in watch:
    add_bullet(doc, w)

add_para(doc, "", space_after=10)
add_para(doc, "── 报告结束 ──", size=10, color_hex="666666",
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12)

doc.save(OUT_PATH)
print(f"[write] {OUT_PATH}")
print(f"[write] size={OUT_PATH.stat().st_size:,} bytes")
